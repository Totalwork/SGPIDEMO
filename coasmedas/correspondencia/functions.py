# -*- coding: utf-8 -*- 
import sys
import subprocess
import re
import json
from django.http import HttpResponse,JsonResponse
from rest_framework.parsers import MultiPartParser, FormParser

import pdfkit ,uuid

from django.conf import settings
from docx import Document
from docx.shared import Inches , Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from rest_framework import status

from .models import CorrespondenciaEnviada, CorrespondenciaPlantilla

from coasmedas.functions import functions
from datetime import date
import datetime
import shutil
import time
import os
from io import StringIO
from django.db.models import Max

from babel.dates import format_date, format_datetime, format_time
from babel.numbers import format_number, format_decimal, format_percent
from numbertoletters import number_to_letters

import mimetypes
from django.http import StreamingHttpResponse
from wsgiref.util import FileWrapper

tagsHtml=['<p>'
	,'<p style="text-align: justify;">'
	,'<p style="text-align: center;">'
	,'<p style="text-align: left;">'
	,'<p style="text-align: right;">'
	,'<strong>'
	,'<em>'
	,'&nbsp;']

def buscaEtiqueta(etiqueta,contenidoHtml,indice):
	sw=0
	while sw==0:
		etiqueta= etiqueta+contenidoHtml[indice+1]
		# print(e)tiqueta+"==contador ="+str(indice)
		if (etiqueta in tagsHtml):
			sw =1
		indice+=1
	return indice

def buscaEtiqueta2(etiqueta,contenidoHtml,indice):
	sw=0
	while sw==0:
		etiqueta= etiqueta+contenidoHtml[indice+1]

		if (etiqueta in tagsHtml):
			sw =1
		indice+=1
	return etiqueta	

array_span =[
	'<span'
	,'<span>'
	,'<span style="mso-spacerun: yes;">'
]

array_open =[
	'<strong>'
	,'<span>'
	,'<em>'
]

def replaceCaracter(html):
	lista = []
	
	lista.append({'code': '&Aacute;' , 'letra' : unicode('Á','utf-8') })
	lista.append({'code': '&aacute;' , 'letra' : unicode('á','utf-8') })
	lista.append({'code': '&Agrave;' , 'letra' : unicode('À','utf-8') })
	lista.append({'code': '&Acirc;' , 'letra' : unicode('Â','utf-8') })
	lista.append({'code': '&agrave;' , 'letra' : unicode('à','utf-8') })
	lista.append({'code': '&Acirc;' , 'letra' : unicode('Â','utf-8') })
	lista.append({'code': '&acirc;' , 'letra' : unicode('â','utf-8') })
	lista.append({'code': '&Auml;' , 'letra' : unicode('Ä','utf-8') })
	lista.append({'code': '&auml;' , 'letra' : unicode('ä','utf-8') })
	lista.append({'code': '&Atilde;' , 'letra' : unicode('Ã','utf-8') })
	lista.append({'code': '&atilde;' , 'letra' : unicode('ã','utf-8') })
	lista.append({'code': '&Aring;' , 'letra' : unicode('Å','utf-8') })
	lista.append({'code': '&aring;' , 'letra' : unicode('å','utf-8') })
	lista.append({'code': '&Aelig;' , 'letra' : unicode('Æ','utf-8') })
	lista.append({'code': '&aelig;' , 'letra' : unicode('æ','utf-8') })
	lista.append({'code': '&Ccedil;' , 'letra' : unicode('Ç','utf-8') })
	lista.append({'code': '&ccedil;' , 'letra' : unicode('ç','utf-8') })
	lista.append({'code': '&Eth;' , 'letra' : unicode('Ð','utf-8') })
	lista.append({'code': '&eth;' , 'letra' : unicode('ð','utf-8') })
	lista.append({'code': '&Eacute;' , 'letra' : unicode('É','utf-8') })
	lista.append({'code': '&eacute;' , 'letra' : unicode('é','utf-8') })
	lista.append({'code': '&Egrave;' , 'letra' : unicode('È','utf-8') })
	lista.append({'code': '&egrave;' , 'letra' : unicode('è','utf-8') })
	lista.append({'code': '&Ecirc;' , 'letra' : unicode('Ê','utf-8') })
	lista.append({'code': '&ecirc;' , 'letra' : unicode('ê','utf-8') })
	lista.append({'code': '&Euml;' , 'letra' : unicode('Ë','utf-8') })
	lista.append({'code': '&euml;' , 'letra' : unicode('ë','utf-8') })
	lista.append({'code': '&Iacute;' , 'letra' : unicode('Í','utf-8') })
	lista.append({'code': '&iacute;' , 'letra' : unicode('í','utf-8') })
	lista.append({'code': '&Igrave;' , 'letra' : unicode('Ì','utf-8') })
	lista.append({'code': '&igrave;' , 'letra' : unicode('ì','utf-8') }) 
	lista.append({'code': '&Icirc;&icirc;' , 'letra' : unicode('Î î','utf-8') })
	lista.append({'code': '&Iuml;' , 'letra' : unicode('Ï','utf-8') })
	lista.append({'code': '&iuml;' , 'letra' : unicode('ï','utf-8') })
	lista.append({'code': '&Ntilde;' , 'letra' : unicode('Ñ','utf-8') })
	lista.append({'code': '&ntilde;' , 'letra' : unicode('ñ','utf-8') })
	lista.append({'code': '&Oacute;' , 'letra' : unicode('Ó','utf-8') })
	lista.append({'code': '&oacute;' , 'letra' : unicode('ó','utf-8') })
	lista.append({'code': '&Ograve;' , 'letra' : unicode('Ò','utf-8') })

	lista.append({'code': '&ograve;' , 'letra' : unicode('ò','utf-8') })
	lista.append({'code': '&Ocirc;' , 'letra' : unicode('Ô','utf-8') })
	lista.append({'code': '&ocirc;' , 'letra' : unicode('ô','utf-8') })
	lista.append({'code': '&Ouml;' , 'letra' : unicode('Ö','utf-8') })
	lista.append({'code': '&ouml;' , 'letra' : unicode('ö','utf-8') })
	lista.append({'code': '&Otilde;' , 'letra' : unicode('Õ','utf-8') })
	lista.append({'code': '&otilde;' , 'letra' : unicode('õ','utf-8') })
	lista.append({'code': '&Oslash;' , 'letra' : unicode('Ø','utf-8') })
	lista.append({'code': '&oslash;' , 'letra' : unicode('ø','utf-8') })
	lista.append({'code': '&szlig;' , 'letra' : unicode('ß','utf-8') })
	lista.append({'code': '&Thorn;' , 'letra' : unicode('Þ','utf-8') })
	lista.append({'code': '&thorn;' , 'letra' : unicode('þ','utf-8') })
	lista.append({'code': '&Uacute;' , 'letra' : unicode('Ú','utf-8') })
	lista.append({'code': '&uacute;' , 'letra' : unicode('ú','utf-8') })
	lista.append({'code': '&Ugrave;' , 'letra' : unicode('Ù','utf-8') })
	lista.append({'code': '&ugrave;' , 'letra' : unicode('ù','utf-8') })
	lista.append({'code': '&Ucirc;' , 'letra' : unicode('Û','utf-8') })
	lista.append({'code': '&ucirc;' , 'letra' : unicode('û','utf-8') })
	lista.append({'code': '&Uuml;' , 'letra' : unicode('Ü','utf-8') })
	lista.append({'code': '&uuml;' , 'letra' : unicode('ü','utf-8') })
	lista.append({'code': '&Yacute;' , 'letra' : unicode('Ý','utf-8') })
	lista.append({'code': '&yacute;' , 'letra' : unicode('ý','utf-8') })
	lista.append({'code': '&yuml;' , 'letra' : unicode('ÿ','utf-8') })
	lista.append({'code': '&copy;' , 'letra' : unicode('©','utf-8') }) 
	lista.append({'code': '&reg;' , 'letra' : unicode('®','utf-8') })
	lista.append({'code': '&trade;' , 'letra' : unicode('™','utf-8') })
	lista.append({'code': '&amp;' , 'letra' : unicode('&','utf-8') })
	lista.append({'code': '&lt;' , 'letra' : unicode('<','utf-8') })
	lista.append({'code': '&gt;' , 'letra' : unicode('>','utf-8') }) 
	lista.append({'code': '&euro;' , 'letra' : unicode('€','utf-8') })
	lista.append({'code': '&cent;' , 'letra' : unicode('¢','utf-8') })
	lista.append({'code': '&pound;' , 'letra' : unicode('£','utf-8') })
	lista.append({'code': '&quot;' , 'letra' : unicode('"','utf-8') })
	lista.append({'code': '&lsquo;' , 'letra' : unicode('‘','utf-8') })
	lista.append({'code': '&rsquo;' , 'letra' : unicode('’','utf-8') })

	lista.append({'code': '&ldquo;', 'letra' : unicode('"','utf-8') })
	lista.append({'code': '&rdquo;' , 'letra' : unicode('"','utf-8') })
	# lista.append({'code': '&laquo;' , 'letra' : unicode('«','utf-8') })
	# lista.append({'code': '&raquo;' , 'letra' : unicode('»','utf-8') })
	lista.append({'code': '&mdash;' , 'letra' : unicode('—','utf-8') })
	lista.append({'code': '&ndash;' , 'letra' : unicode('-','utf-8') })
	
	lista.append({'code': '&deg;' , 'letra' : unicode('°','utf-8') })
	lista.append({'code': '&plusmn;' , 'letra' : unicode('±','utf-8') })
	lista.append({'code': '&frac14;' , 'letra' : unicode('¼','utf-8') })
	lista.append({'code': '&frac12;' , 'letra' : unicode('½','utf-8') })
	lista.append({'code': '&frac34;' , 'letra' : unicode('¾','utf-8') })
	lista.append({'code': '&times;' , 'letra' : unicode('×','utf-8') })
	lista.append({'code': '&divide;' , 'letra' : unicode('÷','utf-8') })
	lista.append({'code': '&alpha;' , 'letra' : unicode('α','utf-8') })
	lista.append({'code': '&beta;' , 'letra' : unicode('β','utf-8') })
	lista.append({'code': '&infin;' , 'letra' : unicode('∞','utf-8') })
	lista.append({'code': '&nbsp;' , 'letra' : unicode(' ','utf-8') })

	# LISTA DE HTML
	lista.append({'code': ' class="MsoNoSpacing"' , 'letra' : '' })
	lista.append({'code': ' class="MsoNormal"' , 'letra' : '' })
	lista.append({'code': ' lang="ES"' , 'letra' : '' })
	lista.append({'code': ' mso-ascii-theme-font: major-latin;' , 'letra' : ''})
	lista.append({'code': ' mso-hansi-theme-font: major-latin;' , 'letra' : ''})
	lista.append({'code': ' mso-ansi-language: ES;' , 'letra' : ''})
	lista.append({'code': ' mso-bidi-font-family: Arial;' , 'letra' : ''})
	lista.append({'code': " font-family: 'Calibri Light',sans-serif;" , 'letra' : ''})
	lista.append({'code': ' name="OLE_LINK3"' , 'letra' : ''})
	lista.append({'code': ' mso-fareast-font-family: Calibri;' , 'letra' : ''})
	lista.append({'code': ' lang="ES-TRAD"' , 'letra' : ''})
	lista.append({'code': " mso-bidi-font-family: 'Times New Roman';" , 'letra' : ''})
	lista.append({'code': " mso-bidi-font-family: Symbol;" , 'letra' : ''})
	lista.append({'code': " mso-fareast-font-family: Symbol;" , 'letra' : ''})
	lista.append({'code': ' class="MsoListParagraph"' , 'letra' : ''})



	replace_text = html
	for value in lista:
		replace_text = replace_text.replace(value["code"], value["letra"])	

	return replace_text	

def reemplaceParagragh(linea):

	texto_buscado = re.findall(r'text-indent: .*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")	

	texto_buscado = re.findall(r'text-indent: .*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(" "+value,"")		

	texto_buscado = re.findall(r'mso-list: .*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")	

	texto_buscado = re.findall(r'font-size: .*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")	

	texto_buscado = re.findall(r'text-indent: .*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")	

	texto_buscado = re.findall(r'\<!--\[endif\]--\>', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	texto_buscado = re.findall(r'\<!-- \[if !supportLists\]--\>', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	texto_buscado = re.findall(r'mso-list: Ignore;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")	

	texto_buscado = re.findall(r'line-height: .*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")	

	texto_buscado = re.findall(r'font: .*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")	

	texto_buscado = re.findall(r'font-family:.*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")	

	texto_buscado = re.findall(r'mso-bidi-font-weight:.*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	texto_buscado = re.findall(r'margin-top:.*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	texto_buscado = re.findall(r'border-collapse:.*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	texto_buscado = re.findall(r'start=".*?"', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"") 

	texto_buscado = re.findall(r'type=".*?"', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"") 

	texto_buscado = re.findall(r'margin-left:.*?;', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"") 

	texto_buscado = re.findall(r'\<span .*?>', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	texto_buscado = re.findall(r'\</span>', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	texto_buscado = re.findall(r'\<div.*?>', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")

	texto_buscado = re.findall(r'\</div>', linea)
	if texto_buscado:
		for value in texto_buscado:
			linea = linea.replace(value,"")  

	linea = linea.replace("'Times New Roman'",'')
	linea = linea.replace('font-family: Symbol;','')
	linea = linea.replace('&middot;','')
	linea = linea.replace('<p style="text-align: justify; ">','<p style="text-align: justify;">')
	linea = linea.replace('<p style="text-align: justify;  ">','<p style="text-align: justify;">')	
	linea = linea.replace(' style=""',"")
	linea = linea.replace(' style=" "',"")
	linea = linea.replace(' style="  "',"")	

	return linea


# REGEX
def converToDocx(request):
	if request.method == 'GET':
		nombre_modulo = "correspondencia.converToDocx"
		try:		
			regex_paragraph = r"\<p\>(.+?)\n*(.*?)\</p\>"
			regex_paragraph = r"\<p\>((.*?)\n*)*\</p\>"
			regex_paragraph = r"\<p\>((.+?)\n*(.*))\</p\>"
			regex_paragraph = r"\<p\>(.*)\n*(.*)\</p\>"
			regex_paragraph = r"\<p\>\n*(.*)\n*(.*)\n*(.*)\n*\</p\>"
			regex_paragraph = r"#\<p\>(.*)\</p\>#sU"
			regex_text = r".*"
			regex_table = r"\<table\>(.*?)\n*(.*?)\n*(((.*?)\n*)*)\n*(.*?)\n*</table\>"
			regex_tbody = r"\<tbody\>(((.*?)\n*)*)\</tbody\>"
			regex_tr = r"\<tr\>(((.*?)\n*)*)\</tr\>"
			regex_td = r"\<td\>(.)*\n*(.*)\n*(.)*\n*(.*)\</td\>"
			regex_width = r'\swidth=".*?"'
			regex_class = r'\swidth=".*?"' 

			correspondencia_id = request.GET['correspondencia_id']
			documento = CorrespondenciaEnviada.objects.get(id=correspondencia_id)
			try:
				plantilla = CorrespondenciaPlantilla.objects.get(empresa_id = documento.empresa.id)
			except CorrespondenciaPlantilla.DoesNotExist:
				return JsonResponse({'message':'La empresa no ha cargado plantilla para las cartas de correspondencia','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		
			
			ruta = settings.STATICFILES_DIRS[0]
			newpath = ruta + '/papelera/'
			filename = str(plantilla.soporte)
			extension = filename[filename.rfind('.'):]
			nombre = 'empresa'+str(documento.empresa.id)+str(documento.consecutivo)+extension
			functions.descargarArchivoS3(str(filename), str(newpath) , nombre )	
			plantilla = settings.STATICFILES_DIRS[0] + '/papelera/'+nombre

			try:
				f = open(plantilla,'rb')
				doc = Document(f)

				font = doc.styles['Normal'].font
				font.name=documento.empresa.tipo_letra

				styles = doc.styles
				table_styles = [s for s in styles ]

				for parrafo in doc.paragraphs:
					texto = parrafo.text

					if texto == '<fechaenvio>':
						fecha_envio = format_date(documento.fechaEnvio, format='long' , locale='es')
						ciudad = documento.ciudad.nombre
						parrafo.text = ''
						run = parrafo.add_run(str(ciudad.capitalize())+' , '+str(fecha_envio))
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)
					if int(texto.find("<consecutivo>"))>=0:
						ano = str(documento.anoEnvio)
						ano = ano[2:4]
						parrafo.text = ''

						consecutivo = str(documento.prefijo.nombre)+'-'+str(documento.consecutivo)
						if documento.prefijo.mostrar_ano:
							consecutivo = consecutivo +'-'+str(ano)

						run = parrafo.add_run('CONSECUTIVO:'+consecutivo)
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_consecutivo)
					if texto == '<referencia>':
						parrafo.text = ''
						run = parrafo.add_run('REFERENCIA :')
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_referencia)
						run = parrafo.add_run(documento.referencia)
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_referencia)
						run.add_break()
					if texto == '<iniciales>':
						parrafo.text = ''	
						run = parrafo.add_run('INICIALES: '+str(documento.usuarioSolicitante.iniciales))
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_iniciales)
					if texto == '<destinatario>':
						parrafo.text = ''
						run = parrafo.add_run(unicode('Señores', 'utf-8'))
						run.add_break()
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)							
						run = parrafo.add_run(documento.empresa_destino)
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)
						run.bold = True
						run.add_break()	
						run = parrafo.add_run('Atn: '+documento.persona_destino)
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)
						run.add_break()	
						run = parrafo.add_run(documento.cargo_persona)
						run.add_break()
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)	
						if documento.direccion:
							run = parrafo.add_run(documento.direccion)
							run.add_break()
							font = run.font
							font.size = Pt(documento.empresa.tam_letra_contenido)	
						if documento.telefono:
							run = parrafo.add_run('Tel.'+documento.telefono+'.')
							run.add_break()
							font = run.font
							font.size = Pt(documento.empresa.tam_letra_contenido)
						if documento.municipioEmpresa:
							run = parrafo.add_run(documento.municipioEmpresa.nombre+'.')
							font = run.font
							font.size = Pt(documento.empresa.tam_letra_contenido)

					if int(texto.find("<asunto>"))>=0:
						parrafo.text = ''	
						run = parrafo.add_run('ASUNTO :')
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)
						run.bold = True	
						run.add_tab()
						run = parrafo.add_run(documento.asunto)
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)
						run.bold = True	
						run.add_break()	
						run = parrafo.add_run()
					if texto == '<contenido>':
						parrafo.text = ''					
						html = str(documento.contenidoHtml)
						html = unicode(html, 'utf-8')					
						# reemplaza caracteres
						html = replaceCaracter(html)					
						# total lineas
						count =len(html.split('\n'))					
						lista = []
						# se divide el parrafo por lineas
						paragraphs = html.split('\n')
						texto = ''
						# se rrecorren las lineas o parrafos
						count_linea = 0
						
						for linea in paragraphs:
							
							linea = reemplaceParagragh(linea)
							# print str(count_linea)+ " -- "+ linea											
							if re.findall(r'\<p\n*(.*)\</p\>', linea):

								etiqueta_paragragh = re.findall(r'\<p.*?>', linea)
								if etiqueta_paragragh:
									texto = ''
									for text in etiqueta_paragragh:
										texto = texto+text
									etiqueta_paragragh = texto

								etiqueta_paragragh= etiqueta_paragragh.replace("<p","")
								etiqueta_paragragh= etiqueta_paragragh.replace(">","")

								etiqueta_p=re.findall(r'\<p'+re.escape(etiqueta_paragragh)+r'\>\n*(.*)\</p\>', linea)
								parrafo_p = parrafo.insert_paragraph_before(" ")
								texto = ''
								for text in etiqueta_p:
									texto = texto+text	
								texto = reemplaceParagragh(texto)

								i = 0
								while  i<len(texto):
									
									if texto[i:(i+8)] == "<strong>":
										i+=8					
															
										while texto[i:(i+9)] != "</strong>":
											# print "<p> -- <strong>"
											if texto[i:(i+4)] == '<em>':
												i+=4
												while texto[i:(i+5)] != "</em>":

													if texto[i:(i+8)] == "<strong>":
														i+=8
														while texto[i:(i+9)] != "</strong>":
															run = parrafo_p.add_run(texto[i])
															run.bold = True
															run.italic = True
															i+=1
														i+=8	
													else:										
														run = parrafo_p.add_run(texto[i])
														run.italic = True														
														i+=1											
												i+=5
											else:																							
												run = parrafo_p.add_run(texto[i])
												run.bold = True													
												i+=1										
										i+=9

									elif text[i:(i+4)] == "<em>" :
										i+=4								
										while texto[i:(i+5)] != "</em>":

											if texto[i:(i+8)] == "<strong>":
												i+=8
												while texto[i:(i+9)] != "</strong>":
													run = parrafo_p.add_run(texto[i])
													run.bold = True
													run.italic = True
													i+=1
												i+=8	
											else:										
												run = parrafo_p.add_run(texto[i])
												run.italic = True														
												i+=1
										i+=6

									else:
										run = parrafo_p.add_run(texto[i])
										i+=1
										
								run.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

							# crear lista
							elif re.findall(r'\<ol .*>', linea):
								etiqueta_lista = re.findall(r'\<ol .*>', linea)
								# texto = ''
								# for text in etiqueta_lista:
								# 	texto = texto+text							

							# elif re.findall(r'<ul style="list-style-type: disc;">', linea):
							# 	etiqueta_lista = re.findall(r'\<ol .*>', linea)
							# 	run = parrafo.add_run("",style= styles["ListaDisco"].name)
							# 	run.add_run(" prueba")

							# elif re.findall(r'<ul style="list-style-type: circle;">', linea):
							# 	etiqueta_lista = re.findall(r'\<ol .*>', linea)
							# 	run = parrafo.add_run("",style= styles["ListaCirculo"].name)
							# 	run.add_run(" prueba")

							# elif re.findall(r'<ul style="list-style-type: square;">', linea):
							# 	etiqueta_lista = re.findall(r'\<ol .*>', linea)
							# 	run = parrafo.add_run("",style= styles["ListaSquare"].name)
							# 	run.add_run(" prueba")

							# elif re.findall(r'<ul>', linea):
							# 	etiqueta_lista = re.findall(r'\<ol .*>', linea)
							# 	run = parrafo.add_run("",style= styles["ListaCirculo"].name)
							# 	run.add_run(" prueba")

							# crear lista
							elif re.findall(r'\</ol\>', linea):
								pass

							elif re.findall(r'\<ul.*>', linea):
								pass	

							elif re.findall(r'\</ul\>', linea):
								pass					
							# <LI>
							elif re.findall(r'\<li .*>', linea):
								# 
								etiqueta_lista = re.findall(r'\<li .*?>', linea)
								if etiqueta_lista:
									texto = ''
									for text in etiqueta_lista:
										texto = texto+text
									etiqueta_li = texto
									tam=len(etiqueta_li)
									etiqueta_lista_contenido = re.findall(r''+re.escape(texto)+r'(.*?)\</li\>', linea)
									if etiqueta_lista_contenido:
										for text in etiqueta_lista_contenido:
											texto = texto+text
										texto=texto+'</li>'
										i=0

										if texto[:(i+tam)] == etiqueta_li:
											i=tam
											li = parrafo.insert_paragraph_before("",style= styles["ListaCirculo"].name)
											while texto[i:(i+5)]!= "</li>":

																							
												if texto[i:(i+8)] == "<strong>":
													i+=8	

													while texto[i:(i+9)] != "</strong>":
											
														if texto[i:(i+4)] == '<em>':
															i+=4
															while texto[i:(i+5)] != "</em>":

																if texto[i:(i+8)] == "<strong>":
																	i+=8
																	while texto[i:(i+9)] != "</strong>":
																		run = li.add_run(texto[i])
																		run.bold = True
																		run.italic = True
																		i+=1
																	i+=8	
																else:										
																	run = li.add_run(texto[i])
																	run.italic = True	
																	run.bold = True													
																	i+=1											
															i+=5

														else:																							
															run = li.add_run(texto[i])
															run.bold = True													
															i+=1										
													i+=9

												elif text[i:(i+4)] == "<em>" :
													i+=4								
													while texto[i:(i+5)] != "</em>":

														if texto[i:(i+8)] == "<strong>":
															i+=8
															while texto[i:(i+9)] != "</strong>":
																run = li.add_run(texto[i])
																run.bold = True
																run.italic = True
																i+=1
															i+=8	
														else:										
															run = li.add_run(texto[i])
															run.italic = True														
															i+=1
													i+=6

												else:
													run = li.add_run(texto[i])
													i+=1
											i+=6

							# RECORRE TODA LA FILA
							elif re.findall(r"""(.*)""", linea):
								# print "todo"
								etiqueta_texto = re.findall(r"""(.*)""", linea)
								texto = ''
								for text in etiqueta_texto:
									texto = texto+text
								
								# se buscan los estilo de tamaño de ltra y se reemplazan por blanco
								texto = reemplaceParagragh(texto)
							
								run = parrafo.add_run(texto)

							count_linea+=1

					if texto == '<autor>':
						parrafo.text = ''
						run = parrafo.add_run(unicode('_____________________________________', 'utf-8'))
						run.add_break()	
						run = parrafo.add_run(documento.firma.persona.nombres+' '+documento.firma.persona.apellidos)
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)
						run.bold = True
						run.add_break()	
						run = parrafo.add_run(documento.firma.cargo.nombre)
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)
						run.add_break()
						run = parrafo.add_run(documento.firma.empresa.nombre)
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_contenido)
						run.add_break()	

					if texto == '<elaboro>':
						parrafo.text = ''
						run = parrafo.add_run()
						run.add_break()	
						run = parrafo.add_run(str('Elaboró - ')+documento.usuarioSolicitante.persona.nombres+' '+documento.usuarioSolicitante.persona.apellidos)
						run.add_break()
						font = run.font
						font.size = Pt(documento.empresa.tam_letra_elaboro)			

				nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\c' + str(documento.id) + '.docx'
				doc.save(nombreArchivo )
				f.close

			except Exception as e:
				raise e

			chunk_size = 108192
			nombreDescarga = 'documento'+str(documento.prefijo.nombre)+'-'+str(documento.consecutivo)+'.docx'
			response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
				content_type=mimetypes.guess_type(nombreArchivo)[0])
			response['Content-Length'] = os.path.getsize(nombreArchivo)
			response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
			
			return response 

		except Exception as e:
			print(e)
			functions.toLog(e,nombre_modulo)
			return JsonResponse({'message':'Se presentaron errores al generar la carta','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)
