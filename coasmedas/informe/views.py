# -*- coding: utf-8 -*- 
from coasmedas.functions import functions
from datetime import date, datetime, timedelta
import datetime
from datetime import *
import shutil
import time
import os
from io import StringIO
import locale
from django.db.models import Max
import math
from django.db.models import Q

from django.shortcuts import render, render_to_response
from django.urls import reverse
from django.http import HttpResponse,JsonResponse

from django.conf import settings
from docx import Document
from docx.shared import Inches , Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.text import WD_ALIGN_PARAGRAPH , WD_COLOR_INDEX

from giros.enum import enumEstados

from django.template import RequestContext

from rest_framework import status
from rest_framework.decorators import api_view

import os
import mimetypes
from django.http import StreamingHttpResponse
from wsgiref.util import FileWrapper

from factura.models import FacturaProyecto, Factura
from contrato.models import Contrato ,EmpresaContrato ,VigenciaContrato, VigenciaContrato_motivo, Contrato_CDP, \
Contrato_Vigencia_anual, Contrato_Financiacion,Contrato_financiacion_condicion, Contrato_Desembolso, \
Contrato_Desembolso_desembolsados,Contrato_Remuneracion,Contrato_Remuneracion_pagos, ActaAsignacionRecursosContrato, \
Actividad,Contrato_Administracion, Contrato_supervisor
from proyecto.models import Proyecto , Proyecto_empresas , Proyecto_proyecto_codigo
from contrato.enumeration import tipoC, estadoC, tipoV
from estado.models import Estado
from estado.views import EstadoSerializer

from .models import Proyecto_Actividad_contrato

from proyecto.views import ProyectoServidumbreSerializer

from babel.dates import format_date, format_datetime, format_time
from babel.numbers import format_number, format_decimal, format_percent

from seguridad_social.models import Empleado
from parametrizacion.models import Funcionario

from tipo.models import Tipo
from tipo.views import TipoSerializer

from contrato.enumeration import tipoC

from administrador_fotos.models import ACategoria,BSubcategoria,CFotosProyecto, DFotosSubcategoria
from giros.models import CNombreGiro,DEncabezadoGiro,DetalleGiro
from django.db.models import Sum 
from django.db.models import Count

from django.contrib.auth.decorators import login_required
from contrato.enumeration import tipoC
import xlsxwriter

from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from openpyxl.styles import Border, NamedStyle, Side, Font, colors, PatternFill
from openpyxl.worksheet import *
from openpyxl.cell import Cell
from openpyxl.utils import get_column_letter, column_index_from_string, coordinate_from_string
# from openpyxl.worksheet.copier import WorksheetCopy
import openpyxl, re

from django.db import transaction, connection
import json
import re
import boto 
from boto.s3.key import Key
import os
import shutil
from io import StringIO


from informe_ministerio.models import Planilla
from empresa.models import Empresa
from financiero.models import FinancieroCuenta, FinancieroCuentaMovimiento
from poliza.models import Poliza, VigenciaPoliza, Aseguradora, VigenciaPoliza_AprobacionMME
from django.db.models import Q, Max, Min
from servidumbre.models import Servidumbre_predio,Servidumbre_documento,Servidumbre_predio_documento
from factura.models import Factura
from ubicacion.models import Ubicacion
# Create your views here.

from rest_framework import viewsets, serializers
from rest_framework import viewsets, response, status
from rest_framework.response import Response
from contrato.views import ContratoLiteSelectSerializer

from logs.models import Logs, Acciones

from openpyxl.styles import Alignment
from avanceObraLite.models import FDetallePresupuesto, DetalleReporteTrabajo, CEsquemaCapitulosActividadesG, DetallePeriodoProgramacion

from cronogramacontrato.models import CcActividadContrato, CcActividadContratoSoporte

class ActividadSerializer(serializers.HyperlinkedModelSerializer):
	class Meta:
		model = Actividad
		fields = ('id','nombre')

class ActividadViewSet(viewsets.ModelViewSet):	
	model=Actividad
	queryset = model.objects.all()
	serializer_class = ActividadSerializer
	nombre_modulo = 'Actividad - ActividadViewSet'

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			
			serializer = self.get_serializer(instance)

			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(ActividadViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			proyecto_id = self.request.query_params.get('proyecto_id', None)

			qset=(~Q(id=0))
			if dato:
				qset = qset &(Q(nombre__icontains = dato))

			if proyecto_id:
				actividades_utilizadas = Proyecto_Actividad_contrato.objects.filter(proyecto_id=proyecto_id).values('actividad__id')
				qset = qset &(~Q(id__in=actividades_utilizadas))

			queryset = self.model.objects.filter(qset)

			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			

			if ignorePagination is None:
				page = self.paginate_queryset(queryset)

				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})
			
			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})

		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor',
				'status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	@transaction.atomic
	def destroy(self,request,*args,**kwargs):
		if request.method == 'DELETE':			
			try:
				###import pdb; pdb.set_trace()

				instance = self.get_object()
				id_aux = instance.id
				self.perform_destroy(instance)
				serializer = self.get_serializer(instance)	

				logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_borrar,
					nombre_modelo='activos.tipo_activo',id_manipulado=id_aux)
				logs_model.save()

				return Response({'message':'El registro ha sido eliminado exitosamente','success':'ok',
					'data':serializer.data},status=status.HTTP_202_ACCEPTED)

			except Exception as e:
				##respuesta=Estructura.error500_2('Este tipo de documento no se puede eliminar debido a que se encuentra en uso')				
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error','data':''},status=status.HTTP_400_BAD_REQUEST)



class Proyecto_ActividadSerializer(serializers.HyperlinkedModelSerializer):
	proyecto = ProyectoServidumbreSerializer (read_only=True)
	actividad = ActividadSerializer (read_only=True)

	class Meta:
		model = Proyecto_Actividad_contrato
		fields = ('id','proyecto','actividad','valor')



class Proyecto_ActividadWriteSerializer(serializers.HyperlinkedModelSerializer):
	proyecto_id = serializers.PrimaryKeyRelatedField(write_only=True, queryset = Proyecto.objects.all())
	actividad_id = serializers.PrimaryKeyRelatedField(write_only=True, queryset = Actividad.objects.all())

	class Meta:
		model = Proyecto_Actividad_contrato
		fields = ('id','proyecto_id','actividad_id','valor')


class Proyecto_ActividadViewSet(viewsets.ModelViewSet):	
	model=Proyecto_Actividad_contrato
	queryset = model.objects.all()
	serializer_class = Proyecto_ActividadSerializer
	nombre_modulo = 'Proyecto_Actividad - Proyecto_ActividadViewSet'

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			
			serializer = self.get_serializer(instance)

			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(Proyecto_ActividadViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			proyecto_id = self.request.query_params.get('proyecto_id', None)

			qset=(~Q(id=0))
			if dato:
				qset = qset &(Q(actividad__nombre__icontains = dato))

			if proyecto_id:
				qset = qset &(Q(proyecto_id = proyecto_id))

			queryset = self.model.objects.filter(qset).order_by('actividad__id')

			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			

			if ignorePagination is None:
				page = self.paginate_queryset(queryset)

				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})
			
			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})

		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor',
				'status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			# import pdb; pdb.set_trace()			
			try:
				serializer = Proyecto_ActividadWriteSerializer(data=request.data,context={'request': request})

				if serializer.is_valid():
					serializer.save(proyecto_id=request.data['proyecto_id'],actividad_id=request.data['actividad_id'])						

					logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,
							nombre_modelo=self.nombre_modulo,id_manipulado=serializer.data['id'])
					logs_model.save()

					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok','data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				return Response({'message':'Se presentaron errores al procesar los datos',
					'success':'error','data':''},status=status.HTTP_400_BAD_REQUEST)
	@transaction.atomic
	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			sid = transaction.savepoint()
			try:
				# partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer  = Proyecto_ActividadWriteSerializer(instance,data=request.data,context={'request': request})

				if serializer.is_valid():					
					serializer.save(proyecto_id=request.data['proyecto_id'],actividad_id=request.data['actividad_id'])

					logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_actualizar,nombre_modelo=self.nombre_modulo,id_manipulado=instance.id)
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok','data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					if serializer.errors['non_field_errors']:
						mensaje = serializer.errors['non_field_errors']
					else:
						mensaje='datos requeridos no fueron recibidos'
					return Response({'message':mensaje,'success':'fail', 'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error','data':''},status=status.HTTP_400_BAD_REQUEST)

	@transaction.atomic
	def destroy(self,request,*args,**kwargs):
		if request.method == 'DELETE':			
			try:
				# import pdb; pdb.set_trace()

				instance = self.get_object()
				id_aux = instance.id
				self.perform_destroy(instance)
				serializer = self.get_serializer(instance)	

				logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_borrar,
					nombre_modelo=self.nombre_modulo,id_manipulado=id_aux)
				logs_model.save()

				return Response({'message':'El registro ha sido eliminado exitosamente','success':'ok',
					'data':serializer.data},status=status.HTTP_202_ACCEPTED)

			except Exception as e:
				##respuesta=Estructura.error500_2('Este tipo de documento no se puede eliminar debido a que se encuentra en uso')				
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error','data':''},status=status.HTTP_400_BAD_REQUEST)

@login_required
def guardar_actividades_contrato(request):
	sid = transaction.savepoint()

	try:		
		soporte= request.FILES['archivo']
		proyecto_id= int(request.POST['proyecto_id'])
		doc = openpyxl.load_workbook(soporte,data_only = True)
		nombrehoja=doc.get_sheet_names()[0]
		hoja = doc.get_sheet_by_name(nombrehoja)
		i=0

		transaction.savepoint_commit(sid)
		revisar_detalle=Proyecto_Actividad_contrato.objects.filter(proyecto_id=proyecto_id)


		if len(revisar_detalle) > 0:
			Proyecto_Actividad_contrato.objects.filter(proyecto_id=proyecto_id).delete()


		contador=hoja.max_row - 1
		numeroFila = 1
		numeroColumna = 1
		numeroCelda = 1

		list_orden = []
		columnas=hoja.columns
		for columna in columnas:
			if numeroColumna==1:			

				for celdas in columna:
					if numeroCelda>1:
						# import pdb; pdb.set_trace()
						if not Actividad.objects.filter(pk=celdas.value).exists():
							return JsonResponse({
								'message':'No debe modificar el ID de la actividad ubicada en la fila No.' + str(numeroCelda) + \
								'. Sugerimos corregir la plantilla e ingresarla nuevamente',
								'success':'error',
								'data':''})
						
					numeroCelda+=1
			break;
			numeroColumna = numeroColumna + 1
		

		if int(contador) > 0:
			

			for fila in hoja.rows:
				if numeroFila>1:
					actividad = Actividad.objects.filter(pk=str(fila[0].value)).first()
					if actividad:
						# import pdb; pdb.set_trace()
						model_actividad_contrato=Proyecto_Actividad_contrato(
							proyecto_id=proyecto_id,
							actividad_id=actividad.id,
							valor=fila[2].value.encode("latin-1") if fila[2].value else ''
							)
						model_actividad_contrato.save()

						logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='servidumbre.predio_georeferencias',id_manipulado=model_actividad_contrato.id)
						logs_model.save()

				numeroFila = numeroFila + 1	

		
		
		return JsonResponse({'message':'El registro se ha guardado correctamente','success':'ok',
				'data':''})

		
	except Exception as e:
		print(e)
		transaction.savepoint_rollback(sid)
		functions.toLog(e,'informe.guardar_coordenadas_archivo')
		
		return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''})


@login_required
def descargar_plantilla_actividades_contrato(request, proyecto_id=None):
	try:
		proyecto_nombre = Proyecto.objects.filter(pk=proyecto_id).first().nombre

		response = HttpResponse(content_type='application/vnd.ms-excel;charset=utf-8')
		response['Content-Disposition'] = 'attachment; filename="Plantilla de actividades del proyecto: '+proyecto_nombre+'.xls"'

		workbook = xlsxwriter.Workbook(response, {'in_memory': True})
		worksheet = workbook.add_worksheet('Actividades del proyecto')

		format1=workbook.add_format({'border':0,'font_size':12,'bold':True})
		format1.set_align('center')
		format1.set_align('vcenter')
		format2=workbook.add_format({'border':0})
		format3=workbook.add_format({'border':0,'font_size':12})
		format5=workbook.add_format()
		format5.set_num_format('yyyy-mm-dd')

		worksheet.set_column('A:A', 10)
		worksheet.set_column('B:C', 60)

		row=1
		col=0

		worksheet.write('A1', 'ID', format1)
		worksheet.write('B1', 'Actividad', format1)
		worksheet.write('C1', 'Valor', format1)

		actividades = Actividad.objects.all().order_by('id')
		for act in actividades:
			worksheet.write(row, col,act.id,format2)
			worksheet.write(row, col+1,act.nombre.encode('utf-8'),format2)

			actividad_contrato = Proyecto_Actividad_contrato.objects.filter(proyecto_id=proyecto_id,actividad_id=act.id).first()
			if actividad_contrato:
				worksheet.write(row, col+2,actividad_contrato.valor.encode('utf-8'),format2)

			row +=1

		workbook.close()
		return response

	except Exception as e:
		functions.toLog(e,'informe.plantilla_actividades_contrato')
		return Response({'message':'Se presentaron errores al procesar los datos','success':'error','data':''},status=status.HTTP_400_BAD_REQUEST)


# INFORME DE INTERVENTORIA DISPAC
def informeInterventoriaDispac(request):
	if request.method == 'GET':
		try:
			tipo_v=tipoV()
			tipo_contrato = tipoC()
			document = Document()

			

			styles = document.styles
			# table_styles = [s for s in styles ]
			# for style in table_styles:
			# 	print(style.name)

			font = styles['Normal'].font
			font.name = 'Arial'

			contrato = request.GET['contrato'] if 'contrato' in request.GET else None;
			ano = request.GET['ano'] if 'ano' in request.GET else None;
			# NUMERO DEL MES DEL 1 AL 12
			mes = request.GET['fecha'] if 'fecha' in request.GET else None;
			firma = request.GET['firma'] if 'firma' in request.GET else None;


			contrato_interventoria = Contrato.objects.get(pk = contrato)
			funcionario_firma = Funcionario.objects.get(pk = firma)			

			interventoria_proyectos = Proyecto.contrato.through.objects.filter(contrato_id = contrato_interventoria.id)
			interventoria_proyectos_2 = Proyecto.objects.filter(contrato__id = contrato_interventoria.id)

			if interventoria_proyectos:
				
				contratos_obra = Proyecto.contrato.through.objects.filter(proyecto_id__in = interventoria_proyectos_2 , contrato__tipo_contrato_id = tipo_contrato.contratoProyecto ).values('contrato_id').annotate(contrato=Count('contrato_id'))
				
				# if contrato_obra.count() == 1:
				if contratos_obra:	
					#contrato_obra = Proyecto.contrato.through.objects.get(proyecto_id = interventoria_proyectos[0].proyecto_id , contrato__tipo_contrato_id = tipo_contrato.contratoProyecto	
					# contrato_obra = Contrato.objects.get(pk = contrato_obra.contrato_id)
					for contrato_obra in contratos_obra:
						contrato_obra = Contrato.objects.get(pk = contrato_obra["contrato_id"])
						# # VALOR TOTAL GIRADO A LOS CONTRATOS
						interventoria_total_girado = DetalleGiro.objects.filter(encabezado__contrato_id = contrato , estado_id = 3).aggregate(Sum('valor_girar'))
						obra_total_girado = DetalleGiro.objects.filter(encabezado__contrato_id = contrato_obra.id , estado_id = 3).aggregate(Sum('valor_girar'))

						if interventoria_total_girado["valor_girar__sum"]:
							interventoria_total_girado = format_decimal(interventoria_total_girado["valor_girar__sum"],  locale='es')
						else:
							interventoria_total_girado = 0

						if obra_total_girado["valor_girar__sum"]:
							
							obra_total_girado = format_decimal(obra_total_girado["valor_girar__sum"],  locale='es')
						else:
							obra_total_girado = 0

						interventoria_valor_contrato = 0
						for vigencia in contrato_interventoria.vigencia_contrato():
							if vigencia.tipo.id == tipo_v.contrato:
								interventoria_valor_contrato = vigencia.valor				

						interventoria_numero_contrato = contrato_interventoria.numero
						interventoria_tipo_contrato = contrato_interventoria.tipo_contrato.nombre
						interventoria_contratista = contrato_interventoria.contratista.nombre
						interventoria_nit_contratista = contrato_interventoria.contratista.nit
						interventoria_objecto = contrato_interventoria.descripcion
						interventoria_fecha_inicio = contrato_interventoria.fecha_inicio()
						interventoria_fecha_fin = contrato_interventoria.fecha_fin()
						interventoria_valor_actual = contrato_interventoria.valor_actual()							
						interventoria_fecha_inicio = str(interventoria_fecha_inicio)
						d = datetime.strptime(interventoria_fecha_inicio, '%Y-%m-%d')
						
						interventoria_fecha_inicio = format_date(d, format='full' , locale='es')
						interventoria_valor_actual = format_decimal(interventoria_valor_actual,  locale='es')
						interventoria_valor_contrato = format_decimal(interventoria_valor_contrato,  locale='es')
						# interventoria_fecha_inicio = ""
						# interventoria_valor_actual = ""
						# interventoria_valor_contrato = ""
						obra_valor_contrato = 0
						for vigencia in contrato_obra.vigencia_contrato():
							if vigencia.tipo.id == tipo_v.contrato:
								obra_valor_contrato = vigencia.valor

						obra_numero_contrato = contrato_obra.numero
						obra_tipo_contrato = contrato_obra.tipo_contrato.nombre
						obra_contratista = contrato_obra.contratista.nombre
						obra_nit_contratista = contrato_obra.contratista.nit
						obra_objecto = contrato_obra.descripcion
						obra_fecha_inicio = contrato_obra.fecha_inicio()
						obra_fecha_fin = contrato_obra.fecha_fin()
						obra_valor_actual = contrato_obra.valor_actual()
						obra_fecha_inicio = str(obra_fecha_inicio)
						obra_fecha_fin = str(obra_fecha_fin)
						d2 = datetime.strptime(obra_fecha_inicio, '%Y-%m-%d')
						d3 = datetime.strptime(obra_fecha_fin, '%Y-%m-%d')
						obra_fecha_inicio = format_date(d2, format='full' , locale='es')
						obra_fecha_fin = format_date(d3, format='full' , locale='es')
						obra_valor_contrato = format_decimal(obra_valor_contrato,  locale='es')
						obra_valor_actual = format_decimal(obra_valor_actual,  locale='es')
						# obra_fecha_inicio = ""
						# obra_fecha_fin = ""
						# obra_valor_contrato = ""
						# obra_valor_actual = ""

						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						p = document.add_paragraph()
						parrafo = p.add_run('INFORME DE AVANCE')
						parrafo.font.size = Pt(18)
						p.alignment = WD_ALIGN_PARAGRAPH.CENTER

						p = document.add_paragraph()
						parrafo = p.add_run('CONTRATO INTERVENTORIA No '+ interventoria_numero_contrato)
						parrafo.bold = True
						parrafo.font.size = Pt(18)
						p.alignment = WD_ALIGN_PARAGRAPH.CENTER

						p = document.add_paragraph('______________________________________________________________')
						p.alignment = WD_ALIGN_PARAGRAPH.CENTER


						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')

						p = document.add_paragraph()
						empresa = Empresa.objects.filter(pk=4)
						parrafo = p.add_run(empresa.nombre)
						parrafo.bold = True
						parrafo.font.size = Pt(16)
						p.alignment = WD_ALIGN_PARAGRAPH.CENTER

						p = document.add_paragraph()
						parrafo = p.add_run('NIT: ')
						parrafo.bold = True
						parrafo.font.size = Pt(16)

						parrafo = p.add_run('XXXXXXXXXXXXXX')
						parrafo.bold = True
						parrafo.font.size = Pt(16)
						parrafo.font.highlight_color = WD_COLOR_INDEX.YELLOW
						p.alignment = WD_ALIGN_PARAGRAPH.CENTER

						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')

						d = date.today()
						fecha_hoy = format_date(d, format='full' , locale='es')
						fecha_hoy_long = format_date(d, format='long' , locale='es')

						p = document.add_paragraph()
						parrafo = p.add_run(fecha_hoy)
						parrafo.bold = True
						parrafo.font.size = Pt(12)
						p.alignment = WD_ALIGN_PARAGRAPH.CENTER

						document.add_page_break()
						
						table = document.add_table(rows=18, cols=2 , style= styles['TableGrid'].name)
						
						hdr_cells = table.rows[0].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Fecha Informe')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(fecha_hoy_long)

						hdr_cells = table.rows[1].cells	

						parrafo = hdr_cells[0].add_paragraph()
						parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
						run = parrafo.add_run(unicode('CONTRATO DE INTERVENTORÍA', 'utf-8'))
						run.bold = True
						a, b = hdr_cells[:2]
						a.merge(b)
						hdr_cells[0].add_paragraph()
						
						hdr_cells = table.rows[2].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('No.')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(interventoria_numero_contrato)

						hdr_cells = table.rows[3].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Contratista')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(interventoria_contratista)

						hdr_cells = table.rows[4].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Nit. Contratista')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(interventoria_nit_contratista)

						hdr_cells = table.rows[5].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Tipo de Contrato')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(interventoria_tipo_contrato)

						hdr_cells = table.rows[6].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run(unicode('Objeto del Contrato Interventoría', 'utf-8' ))

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(interventoria_objecto)

						hdr_cells = table.rows[7].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Valor del Contrato')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run('$ '+str(interventoria_valor_contrato))

						hdr_cells = table.rows[8].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Fecha de Inicio')

						parrafo = hdr_cells[1].add_paragraph()

						run = parrafo.add_run(interventoria_fecha_inicio)

						hdr_cells = table.rows[9].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run(unicode('Duración del Contrato' , 'utf-8'))

						hdr_cells = table.rows[10].cells	
						parrafo = hdr_cells[0].add_paragraph()
						parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
						run = parrafo.add_run('CONTRATO DE OBRA')
						run.bold = True
						run = parrafo.add_run('')
						a, b = hdr_cells[:2]
						a.merge(b)
						hdr_cells[0].add_paragraph()

						hdr_cells = table.rows[11].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('No.')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(str(obra_numero_contrato))

						hdr_cells = table.rows[12].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Contratista')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(obra_contratista)

						hdr_cells = table.rows[13].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Nit. Contratista')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(obra_nit_contratista)

						hdr_cells = table.rows[14].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Objeto del Contrato de Obra')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(obra_objecto)

						hdr_cells = table.rows[15].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Valor del Contrato')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run('$ '+str(obra_valor_contrato))			

						hdr_cells = table.rows[16].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run('Fecha de Inicio')

						parrafo = hdr_cells[1].add_paragraph()
						run = parrafo.add_run(obra_fecha_inicio)	

						hdr_cells = table.rows[17].cells	
						parrafo = hdr_cells[0].add_paragraph()
						run = parrafo.add_run(unicode('Fecha de Terminación', 'utf-8'))
						
						parrafo = hdr_cells[1].add_paragraph()
						
						run = parrafo.add_run(obra_fecha_fin)	
						
						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph(unicode('Información Financiera del Contrato de Obra: ' , 'utf-8'), style=styles['ListNumber'].name)
						
						document.add_paragraph( 'Valor Inicial del Contrato:					$ '+str(obra_valor_contrato)
												, style=styles['ListBullet2'].name)
						document.add_paragraph( 'Valor Actual del Contrato:					$ '+str(obra_valor_actual)
												, style=styles['ListBullet2'].name)
						document.add_paragraph( 'Valor Total Girado:						$ '+str(obra_total_girado)
												, style=styles['ListBullet2'].name)
						document.add_paragraph('')
						

						document.add_paragraph('Personal del Contratista de obra: ', style=styles['ListNumber'].name)

						empleados_contratista_obra = Empleado.objects.filter(estado_id__in = [5,7] , empresa_id = contrato_obra.contratista.id)

						empleados_contratista_obra_total = Empleado.objects.filter(estado_id__in = [5,7] , empresa_id = contrato_obra.contratista.id).values('cargo__nombre').annotate(total=Count('cargo_id')).order_by('cargo__nombre')
						

						if empleados_contratista_obra:
							document.add_paragraph('')
							document.add_paragraph('')

							personal = document.add_table(rows=1, cols=4)
							personal.style = 'TableGrid'
							personal_cells = personal.rows[0].cells

							parrafo = personal_cells[0].add_paragraph()
							run = parrafo.add_run('ITEM')
							run.bold = True
							run.font.size = Pt(9)
							parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

							parrafo = personal_cells[1].add_paragraph()
							run = parrafo.add_run('NOMBRE COMPLETO')
							run.bold = True
							run.font.size = Pt(9)
							parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

							parrafo = personal_cells[2].add_paragraph()
							run = parrafo.add_run('CEDULA')
							run.bold = True
							run.font.size = Pt(9)
							parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

							parrafo = personal_cells[3].add_paragraph()
							run = parrafo.add_run('CARGO')
							run.bold = True
							run.font.size = Pt(9)
							parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

							i = 1

							for item in empleados_contratista_obra:

								row_cells = personal.add_row().cells

								parrafo = row_cells[0].add_paragraph()
								run = parrafo.add_run(str(i))
								run.font.size = Pt(9)
								parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

								parrafo = row_cells[1].add_paragraph()

								nombre_completo = item.persona.nombres+' '+item.persona.apellidos
								run = parrafo.add_run(nombre_completo)
								run.bold = True
								run.font.size = Pt(9)	
								
								parrafo = row_cells[2].add_paragraph()
								run = parrafo.add_run(str(item.persona.cedula))
								run.font.size = Pt(9)

								cargo = ''
								if item.cargo:
									cargo = item.cargo.nombre
								parrafo = row_cells[3].add_paragraph()
								run = parrafo.add_run(cargo)
								run.font.size = Pt(9)
								i+=1
						else:
							p = document.add_paragraph()
							run = p.add_run('No se encontro personal del contratista de obra.')
							run.font.highlight_color = WD_COLOR_INDEX.YELLOW
							paragraph_format = p.paragraph_format							
							paragraph_format.left_indent = Inches(0.5)

						if empleados_contratista_obra_total:
							document.add_paragraph('')
							personal_num = document.add_table(rows=1, cols=2)
							personal_num.style = 'TableGrid'
							personal_num_cells = personal_num.rows[0].cells
							parrafo = personal_num_cells[0].add_paragraph()
							run = parrafo.add_run('CARGO')
							run.bold = True
							run.font.size = Pt(9)
							parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

							parrafo = personal_num_cells[1].add_paragraph()
							run = parrafo.add_run('CANTIDAD')
							run.bold = True
							run.font.size = Pt(9)
							parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

							for item in empleados_contratista_obra_total:
								row_cells = personal_num.add_row().cells

								parrafo = row_cells[0].add_paragraph()
								run = parrafo.add_run(item["cargo__nombre"])
								run.font.size = Pt(9)
								parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

								parrafo = row_cells[1].add_paragraph()
								run = parrafo.add_run(str(item["total"]))
								run.bold = True
								run.font.size = Pt(9)	
								parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER


						document.add_paragraph('')
						document.add_paragraph(unicode('Estado y avance fisico del proyecto: ' , 'utf-8'), style=styles['ListNumber'].name)
						list_proyectos = []
						if interventoria_proyectos:
							sw_foto_avance = 1
							for item in interventoria_proyectos:

								fotos_proyecto = CFotosProyecto.objects.filter(proyecto__id = item.proyecto.id, fecha__year = ano  , fecha__month = mes)

								if fotos_proyecto:
									sw_foto_avance = 0

									list_proyectos.append(item.proyecto.id)
									p = document.add_paragraph('	Proyecto : '+item.proyecto.nombre , style=styles['List'].name)
									run = p.add_run('')
									paragraph_format = p.paragraph_format							
									paragraph_format.left_indent = Inches(0.5)

									p = document.add_paragraph('')
									run = p.add_run('DESCRIBA EL AVANCE DEL PROYECTO.')
									run.font.highlight_color = WD_COLOR_INDEX.YELLOW
									paragraph_format = p.paragraph_format							
									paragraph_format.left_indent = Inches(1)

									total_row = fotos_proyecto.count()
									filas = total_row%2
									total_row = float(total_row)/2
									# print math.ceil(total_row)
									# print int(total_row/2)
									total_row = math.ceil(total_row)
									table_avance = document.add_table( rows = int(total_row),cols=2)		
									table_avance.style = 'TableGrid'

									j = 0
									for i in range(int(total_row)):

										row_cells = table_avance.rows[i].cells

										foto = fotos_proyecto[j]
										newpath = r'static/papelera/'
										filename = str(foto.ruta)
										extension = filename[filename.rfind('.'):]
										nombre = str(foto.id)+extension

										functions.descargarArchivoS3(str(filename), str(newpath)+"/" , nombre )	
										row_cells[0].add_paragraph()
										parrafo = row_cells[0].add_paragraph()
										run = parrafo.add_run()
										

										if ((int(i)+1)!=total_row) or (filas == 0):
											run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(2.8) , height = Inches(2) )
											row_cells[0].add_paragraph()
											j = j+1

											foto = fotos_proyecto[j]
											newpath = r'static/papelera/'
											filename = str(foto.ruta)
											extension = filename[filename.rfind('.'):]
											nombre = str(foto.id)+extension

											functions.descargarArchivoS3(str(filename), str(newpath)+"/" , nombre )	
											row_cells[1].add_paragraph('')
											parrafo = row_cells[1].add_paragraph()
											run = parrafo.add_run()
											run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(2.8) , height = Inches(2))
											row_cells[1].add_paragraph()

										else:

											run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(3) , height = Inches(2))
											parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
											row_cells[0].add_paragraph()
											a, b = row_cells[:2]
											a.merge(b)
										document.add_paragraph('')
										document.add_paragraph('')

										j = j+1

							if sw_foto_avance==1:
								p = document.add_paragraph()
								run = p.add_run('	NO SE ENCONTRARON FOTOS ASOCIADAS.')
								run.font.highlight_color = WD_COLOR_INDEX.YELLOW

						else:
							p = document.add_paragraph()
							run = p.add_run('	NO SE ENCONTRARON PROYECTOS ASOCIADOS AL CONTRATO DE INTERVENTORIA.')
							run.font.highlight_color = WD_COLOR_INDEX.YELLOW

						document.add_paragraph(unicode('Actividades ejecutadas por el Contratista: ' , 'utf-8'), style=styles['ListNumber'].name)
						
						
						for item in interventoria_proyectos:
							list_proyectos.append(item.proyecto.id)
							document.add_paragraph('  	Proyecto : '+item.proyecto.nombre , style=styles['List'].name)
							categorias = ACategoria.objects.filter(contrato_id = item.proyecto.mcontrato_id)
							
							fotos_subcategoria_valida = DFotosSubcategoria.objects.filter(subcategoria__categoria__proyecto_id = item.proyecto.id , ano = ano  , mes = mes).count()
							
							if fotos_subcategoria_valida>0:	
								for c in categorias:	
									subcategorias = BSubcategoria.objects.filter(categoria_id = c.id, proyecto_id = item.proyecto.id)
									
									if subcategorias:	
																				
										p = document.add_paragraph(c.categoria , style=styles['ListParagraph'].name)
										paragraph_format = p.paragraph_format							
										paragraph_format.left_indent = Inches(1)
										for item_subcategoria in subcategorias:

											fotos_subcategoria = DFotosSubcategoria.objects.filter(subcategoria_id = item_subcategoria.id , ano = ano  , mes = mes)

											if fotos_subcategoria:
												p = document.add_paragraph(item_subcategoria.titulo	, style=styles['ListBullet3'].name)
												paragraph_format = p.paragraph_format							
												paragraph_format.left_indent = Inches(1.5)

												total_row = fotos_subcategoria.count()
												filas = total_row%2
												total_row = float(total_row)/2
												# print math.ceil(total_row)
												# print int(total_row/2)
												total_row = math.ceil(total_row)
												table_subcategoria = document.add_table( rows = int(total_row),cols=2)		
												table_subcategoria.style = 'TableGrid'

												j = 0
												for i in range(int(total_row)):

													row_cells = table_subcategoria.rows[i].cells

													foto = fotos_subcategoria[j]
													newpath = r'static/papelera/'
													filename = str(foto.ruta)
													extension = filename[filename.rfind('.'):]
													nombre = str(foto.id)+extension

													functions.descargarArchivoS3(str(filename), str(newpath)+"/" , nombre )	
													row_cells[0].add_paragraph('')
													parrafo = row_cells[0].add_paragraph()
													run = parrafo.add_run()

													if ((int(i)+1)!=total_row) or (filas == 0):
														run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(2.8) , height = Inches(2))
														row_cells[0].add_paragraph()
														j = j+1
														# print "diferente"							
														# row_cells[1].text = "columna dos"

														foto = fotos_proyecto[j]
														newpath = r'static/papelera/'
														filename = str(foto.ruta)
														extension = filename[filename.rfind('.'):]
														nombre = str(foto.id)+extension

														functions.descargarArchivoS3(str(filename), str(newpath)+"/" , nombre )	
														row_cells[1].add_paragraph('')
														parrafo = row_cells[1].add_paragraph()
														run = parrafo.add_run()
														run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(2.8) , height = Inches(2))
														row_cells[1].add_paragraph()

													else:

														run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(3) , height = Inches(2))
														parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
														row_cells[0].add_paragraph()
														a, b = row_cells[:2]
														a.merge(b)

													j = j+1

												document.add_paragraph('')
										
							else:
								p = document.add_paragraph('')
								run = p.add_run('No se encontraron actividades ejecutadas por el contratista en el proyecto.')
								run.font.highlight_color = WD_COLOR_INDEX.YELLOW
								paragraph_format = p.paragraph_format							
								paragraph_format.left_indent = Inches(1)	
												
						document.add_paragraph('')
						document.add_paragraph(unicode('Información Financiera del Contrato de Interventoría:  ' , 'utf-8'), style=styles['ListNumber'].name)

						document.add_paragraph( 'Valor Inicial del Contrato:					$ '+str(interventoria_valor_contrato)
												, style=styles['ListBullet2'].name)
						document.add_paragraph( 'Valor Actual del Contrato:					$ '+str(interventoria_valor_actual)
												, style=styles['ListBullet2'].name)
						document.add_paragraph( 'Valor Total Girado:						$ '+str(interventoria_total_girado)
												, style=styles['ListBullet2'].name)

						document.add_paragraph('')
						document.add_paragraph(unicode('Personal Interventoría: ' , 'utf-8'), style=styles['ListNumber'].name)


						personal_interventoria = Proyecto.funcionario.through.objects.filter(proyecto_id__in = list_proyectos 
																							, funcionario__empresa_id = contrato_interventoria.contratista.id
																							, funcionario__activo = 1
													).values('funcionario__persona__apellidos',
													'funcionario__persona__nombres', 'funcionario__cargo__nombre' ).distinct()

						if personal_interventoria:
							for item in personal_interventoria:
								
								document.add_paragraph( item['funcionario__persona__nombres']+' '+item['funcionario__persona__apellidos']+' - '+item['funcionario__cargo__nombre'] , style=styles['ListNumber2'].name)
						else:
							document.add_paragraph(unicode('No se encontró personal de la interventoria asociado a un proyecto' , 'utf-8'), style=styles['List'].name)
						document.add_paragraph('')
						document.add_paragraph(unicode('Actividades ejecutadas por la Interventoría: ' , 'utf-8'), style=styles['ListNumber'].name)

						for item in interventoria_proyectos:			

							actividades_interventoria = Proyecto_actividad.objects.filter(proyecto_id = item.proyecto.id)

							if actividades_interventoria:
								document.add_paragraph( 'Proyecto : '+item.proyecto.nombre
												, style=styles['ListParagraph'].name)
								for actividad in actividades_interventoria:
									document.add_paragraph( actividad.descripcion , style=styles['ListBullet3'].name)
						document.add_paragraph('')			
						document.add_paragraph(unicode('Conclusiones y recomendaciones: ' , 'utf-8'), style=styles['ListNumber'].name)

						p = document.add_paragraph()
						run = p.add_run('XXXXXXXXXX')
						run.font.highlight_color = WD_COLOR_INDEX.YELLOW

						document.add_paragraph('')
						document.add_paragraph('')
						document.add_paragraph('')

						p = document.add_paragraph()
						run = p.add_run('Cordialmente,')

						document.add_paragraph('')

						nombre_completo = funcionario_firma.persona.nombres+' '+funcionario_firma.persona.apellidos
						p = document.add_paragraph()
						run = p.add_run(nombre_completo)
						run.bold = True
						run.add_break()
						cargo = funcionario_firma.cargo.nombre
						run = p.add_run(cargo)

						document.add_page_break()
						
				else:
					# if contrato_obra:
					p = document.add_paragraph()
					run = p.add_run('No se encontraron contratos de obras asociados a los proyectos del contrato de interventoria.')
					# else:
					# 	p = document.add_paragraph()
					# 	run = p.add_run('Se encontraron mas de un contrato de obra asociados a los proyectos del contrato de interventoria.')
			else:
				p = document.add_paragraph()
				run = p.add_run('No se encontraron proyectos asociados al contrato de interventoria.')
			

			
			nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\prueba.docx'
			document.save(nombreArchivo)


			chunk_size = 108192

			nombreDescarga = 'informeInterventoria.docx'
			response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
				content_type=mimetypes.guess_type(nombreArchivo)[0])
			response['Content-Length'] = os.path.getsize(nombreArchivo)
			response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
			
			return response
		except Exception as e:
			print(e)
			functions.toLog(e,'informe.interventoria')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['POST'])
def GenerarinformeMME_validar (request):
	try:

		contrato= Contrato.objects.get(id=int(request.GET['contrato']))
		limite_ano=int(request.GET['ano'])
		limite_mes=int(request.GET['mes'])
		if limite_mes==12:
			limite_ano=limite_ano+1
			limite_mes=1
		else:
			limite_mes=limite_mes+1

		limite_fecha=str(limite_ano)+'-'+str(limite_mes)+'-'+'1'

		limite_fecha=datetime.strptime(str(limite_fecha),'%Y-%m-%d').date()

		limite_fecha= limite_fecha - timedelta(days=1)

		limite_fecha_contrato	=datetime.strptime(str(contrato.fecha_inicio()),'%Y-%m-%d').date()

		fecha_actual = datetime.now().date()

		#import pdb; pdb.set_trace()

		if str(fecha_actual.year) == str(limite_fecha.year) and str(fecha_actual.month) == str(limite_fecha.month):
			limite_fecha=str(limite_fecha.year)+'-'+str(limite_fecha.month)+'-'+str(fecha_actual.day)
			limite_fecha=datetime.strptime(str(limite_fecha),'%Y-%m-%d').date()

		if str(limite_fecha_contrato)>str(limite_fecha):
			return JsonResponse({'message':'Este contrato empieza luego de la fecha elegida','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)
		elif str(fecha_actual)<str(limite_fecha):
			return JsonResponse({'message':'La fecha elegida todavia aun no pasa','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

		return JsonResponse({'message':'El informe se está generando, por favor espere un momento','success':'ok','data':''},status=status.HTTP_201_CREATED)

	except Exception as e:
			print(e)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def GenerarinformeMME (request):	
	try:
		contrato= Contrato.objects.get(id=int(request.GET['contrato']))
		limite_ano=int(request.GET['ano'])
		limite_mes=int(request.GET['mes'])
		if limite_mes==12:
			limite_ano=limite_ano+1
			limite_mes=1
		else:
			limite_mes=limite_mes+1

		limite_fecha=str(limite_ano)+'-'+str(limite_mes)+'-'+'1'
		limite_fecha=datetime.strptime(str(limite_fecha),'%Y-%m-%d').date()
		limite_fecha= limite_fecha - timedelta(days=1)

		fecha_actual = datetime.now().date()

		if str(fecha_actual.year) == str(limite_fecha.year) and str(fecha_actual.month) == str(limite_fecha.month):
			limite_fecha=str(limite_fecha.year)+'-'+str(limite_fecha.month)+'-'+str(fecha_actual.day)
			limite_fecha=datetime.strptime(str(limite_fecha),'%Y-%m-%d').date()


		response = HttpResponse(content_type='application/vnd.ms-excel;charset=utf-8')
		response['Content-Disposition'] = 'attachment; filename="'+contrato.nombre+' ['+str(limite_fecha.month)+'/'+str(limite_fecha.year)+'].xls"'

		
		proyecto = Proyecto.objects.get(pk=int(request.GET['proyecto_id']))
		# if proyecto:
		# 	proyecto = Proyecto.objects.filter(pk=int(request.GET['proyecto_id']))
		# 	# proyecto = proyecto.first()
		registro_planilla = Planilla.objects.get(id=2)
		soporte = registro_planilla.archivo
		doc = openpyxl.load_workbook(soporte,read_only=False, keep_vba=True, guess_types=True, data_only=False)

		## 1-INFOR-CONTRATO
		nombrehoja=doc.get_sheet_names()[0]
		hoja_0 = doc.get_sheet_by_name(nombrehoja)
		hoja_0.title=nombrehoja

		nombrehoja=doc.get_sheet_names()[1]
		hoja_1 = doc.get_sheet_by_name(nombrehoja)
		hoja_1.title=nombrehoja
		fondo = ""
		b7_escrito=""
		
		if proyecto.tipo_proyecto.id==2:
			fondo = 'FAER'
			hoja_1['A5'].value='ACTA CAFAER'
			hoja_1['B1'].value=contrato.nombre.replace('FAER ','')
			hoja_1['B6'].value='Construir infraestructura eléctrica para ampliar y prestar el servicio de energía eléctrica, en condiciones de calidad y confiabilidad, en las zonas rurales del Sistema Interconectado Nacional - SIN, ubicadas en el Mercado de Comercialización del OPERADOR DE RED, mediante la ejecución del proyecto ELECTRIFICACIÓN RURAL '+proyecto.nombre +' MUNICIPIO DE '+proyecto.municipio.nombre+', DEPARTAMENTO DE '+proyecto.municipio.departamento.nombre+' el cual será construido por el OPERADOR DE RED bajo su responsabilidad, en los términos del presente Contrato, para la entrega al MINISTERIO y el cual es financiado con el Fondo de Apoyo Financiero para la Energización de las Zonas Rurales Interconectadas - FAER'
			b7_escrito='En desarrollo del objeto de este Contrato, para la ejecución del proyecto ELECTRIFICACIÓN RURAL'+proyecto.nombre +' el OPERADOR DE RED deberá:'
			b7_escrito=b7_escrito+'\n1. Construir por su propia cuenta y riesgo la infraestructura objeto del proyecto que luego será entregada al MINISTERIO de conformidad con lo establecido en el presente Contrato.'
			b7_escrito=b7_escrito+'\n2. Aceptar el presupuesto del proyecto.'
			b7_escrito=b7_escrito+'\n3. Administrar los recursos asignados al proyecto de conformidad con lo establecido en el presente contrato.'
			b7_escrito=b7_escrito+'\n4. Contratar y liquidar el encargo Fiduciario para el manejo de los recursos asignados, de conformidad con las especificaciones incluidas en el presente contrato.'
			b7_escrito=b7_escrito+'\n5. Contratar la interventoría, los suministros que sean del caso y la construcción de obras, según corresponda, y liquidar los mismos, de conformidad con las especificaciones incluidas en el presente contrato.'
			b7_escrito=b7_escrito+'\n6. Realizar la dirección y asistencia técnica necesaria para construir y poner en operación el proyecto.'
			b7_escrito=b7_escrito+'\n7. Realizar la socialización previa y periódica del proyecto con la(s) comunidad(es) beneficiada(s). Reportar al MINSITERIO si las comunidades presentan objeciones a la realización del proyecto.'
			b7_escrito=b7_escrito+'\n8. Verificar la localización y condiciones técnicas favorables (replanteo) de las diferentes obras de infraestructura del proyecto en el campo y reportar las diferencias o anomalías al MINISTERIO, previo al inicio de la construcción.'
			b7_escrito=b7_escrito+'\n9. Verificar la necesidad de servidumbres y/o permisos de paso, cuando corresponda; y realizar los requerimientos necesarios a las entidades territoriales para el logro de la consecución y legalización de las mismas, previo de la construcción del proyecto.'
			b7_escrito=b7_escrito+'\n10. Verificar el número y ubicación de los usuarios (replanteo) y reportar al MINISTERIO si existen diferencias, se reportarán al MINISTERIO, previo al inicio de la construcción.'
			b7_escrito=b7_escrito+'\n11. Verificar, previo al inicio de las obras objeto de ejecución, que éstas no se encuentran ubicadas en zonas de alto riesgo.'
			b7_escrito=b7_escrito+'\n12. Verificar la necesidad de licencias y permisos ambientales. Si es el caso, gestionarlas y cumplir con todos los planes exigidos en ellos.'
			b7_escrito=b7_escrito+'\n13. Mantener informado al MINISTERIO sobre los avances e inconvenientes de la construcción del proyecto de tal manera que se puedan tomar decisiones oportunas.'
			b7_escrito=b7_escrito+'\n14. Realizar la entrega a satisfacción del MINISTERIO de la infraestructura energizada y con el respectivo inventario detallado de los activos construidos.'
			b7_escrito=b7_escrito+'\n15. Recibir en aporte para uso y goce, los activos construidos.'
			b7_escrito=b7_escrito+'\n16. Realizar las actividades de Administración, Operación y Mantenimiento, en adelante AOM, y efectuar la reposición de los activos aportados en caso de que sea necesario.'
			hoja_1['B7'].value=b7_escrito

		if proyecto.tipo_proyecto.id==1:
			fondo = 'PRONE'
			hoja_1['A5'].value='ACTA CAPRONE'
			hoja_1['B1'].value=contrato.nombre.replace('PRONE ','')			
			hoja_1['B6'].value='Mejorar la calidad y confiabilidad del servicio de energía eléctrica en los barrios subnormales del Mercado de Comercialización del OPERADOR DE RED ubicados en los municipios del Sistema Interconectado Nacional – SIN conforme los reglamentos técnicos vigentes, mediante la ejecución del proyecto NORMALIZACION DE REDES ELECTRICAS BARRIO '+proyecto.nombre+', el cual será ejecutado por el OPERADOR DE RED bajo su responsabilidad en los términos del presente Contrato, para la entrega al MINISTERIO y el cual es financiado con fondos del Programa de Normalización de Redes Eléctricas - PRONE.'
			b7_escrito='En desarrollo del objeto de este Contrato. El OPERADOR DE RED deberá administrar los recursos del fondo PRONE asignados para llevar a cabo El Proyecto, ejecutando las actividades previas a la contratación de obras, las actividades de contratación de obras, las actividades de suministro, construcción y liquidación de obras, las actividades de recibo y terminación del contrato del administrador.'
			b7_escrito=b7_escrito+'\nAsí mismo, deberá energizar y recibir en aporte, en los términos del articulo 87.9 de la Ley 142 de 1994, la infraestructura construida, realizar las actividades inherentes a la Administración, Operación y Mantenimiento (en adelante AOM), efectuar la reposición de los activos de distribución aportados en caso de ser necesario y ejecutar las actividades referentes a la devolución de la Infraestructura para el Proyecto descrito en el Apéndice Técnico.'
			
			hoja_1['B7'].value=b7_escrito

		hoja_1['B2'].value=fondo

		contratante = Empresa.objects.get(id=contrato.contratante_id)
		hoja_1['B3'].value= contratante.nombre

		contratista = Empresa.objects.get(id=contrato.contratista_id)
		# hoja_1['B4'].value= contratista.nombre

		empresa_contratista=''
		contrato_suministros=contrato.contratista.nombre
		if contrato_suministros.count('AIR-E'):			
			empresa_contratista='AIR-E S.A.S. E.S.P.'			

		elif contrato_suministros.count('MAR'):			
			empresa_contratista='AFINIA S.A.S. ESP'

		hoja_1['B4'].value=empresa_contratista

		ActaAsign=ActaAsignacionRecursosContrato.objects.filter(contrato_id=contrato).exists()
		if ActaAsign:
			ActaAsign=ActaAsignacionRecursosContrato.objects.filter(contrato_id=contrato)
			ActaAsign=ActaAsign.last().actaAsignacion.nombre
			ActaAsign=ActaAsign=int(filter(str.isdigit, str(ActaAsign)))
			hoja_1['B5'].value=ActaAsign
		else:
			hoja_1['B5'].value='No se ha suscrito el contrato aun Acta de Asignacion de Recursos'

		#import pdb; pdb.set_trace()

		if contrato.fecha_firma and contrato.fecha_firma<limite_fecha:
			hoja_1['B8'].value=contrato.fecha_firma
			hoja_1['B8'].number_format='dd/mm/yyyy'
		else:
			hoja_1['B8'].value='No se ha suscrito la fecha de la firma'

		if contrato.fecha_acta_inicio and contrato.fecha_firma<limite_fecha:
			hoja_1['B9'].value=contrato.fecha_acta_inicio
			hoja_1['B9'].number_format='dd/mm/yyyy'
		else:
			hoja_1['B9'].value='No se ha suscrito la fecha del acta de inicio'
	

		cdp=Contrato_CDP.objects.filter(contrato_id=int(contrato.id),fecha__lte=limite_fecha).exists()
		k=0
		if cdp:
			cdp=Contrato_CDP.objects.filter(contrato_id=int(contrato.id),fecha__lte=limite_fecha)
			aux=''			
			for val in cdp:
				k=k+1
				ano=str(val.fecha.year)
				aux=aux+str(val.numero)+' de '+ano
				if k>1:
					aux=aux+','
			hoja_1['B20'].value=aux
			
		else:
			hoja_1['B20'].value='No se ha suscrito el CDP'
		hoja_1['B21'].value='El contrato a la fecha cuenta con ('+str(k)+') CDP'

		count_suspension = 0
		count_reinicio = 0
		count_prorroga = 0
		count_liquidacion = 0
		suspension_sin_fecha_fin = False
			

		vigencias_contrato = VigenciaContrato.objects.filter(contrato_id=contrato.id,fecha_inicio__lte=limite_fecha)
		for vigencia in vigencias_contrato:
			if vigencia.tipo_id==18:
				if vigencia.nombre=='Acta de suspension No.1' or vigencia.nombre=='Acta de suspension No. 1':
					count_suspension=count_suspension+1
					hoja_1['C13'].value=vigencia.fecha_inicio
					hoja_1['C14'].value=str(vigencia.fecha_inicio) + ' - ' + str(vigencia.fecha_fin)
					
					motivo=VigenciaContrato_motivo.objects.filter(vigencia_contrato_id=vigencia.id).exists()
					if motivo:
						motivo=VigenciaContrato_motivo.objects.get(vigencia_contrato_id=vigencia.id)
						hoja_1['C12'].value=motivo.motivo
					else:
						hoja_1['C12'].value='El contrato a la fecha no se ha visto suspendido por las partes'


					if vigencia.fecha_fin == '':
						suspension_sin_fecha_fin=True
					else:
						suspension_sin_fecha_fin=False
						hoja_1['C15'].value=vigencia.fecha_fin
			

			if vigencia.tipo_id==102: 
				if vigencia.nombre == 'Prórroga No.1 del Acta de suspension No.1' or vigencia.nombre=='Prórroga No.1 del Acta de suspension No. 1':
					count_prorroga=count_prorroga+1
					hoja_1['C18'].value=vigencia.fecha_fin

					motivo=VigenciaContrato_motivo.objects.filter(vigencia_contrato_id=vigencia.id).exists()
					if motivo:
						motivo=VigenciaContrato_motivo.objects.get(vigencia_contrato_id=vigencia.id)
						hoja_1['C16'].value=motivo.motivo
					else:
						hoja_1['C16'].value='El contrato no ha surtido prorrogas'


			if vigencia.tipo_id==21: 
				if vigencia.nombre == 'Liquidacion' or vigencia.nombre=='Liquidación' or vigencia.nombre=='LiquidaciÃ³n':
					count_liquidacion=count_liquidacion+1
					if vigencia.fecha_inicio:
						hoja_1['B41'].value=vigencia.fecha_inicio
					else:
						hoja_1['B41'].value='Falta registrar la fecha de liquidación pero la vigencia ya fue cargada'


		if suspension_sin_fecha_fin:
			for vigencia in vigencias_contrato:
				if vigencia.tipo_id==19: 
					if vigencia.nombre=='Acta de reinicio No.1' or vigencia.nombre=='Acta de reinicio No. 1':
						count_reinicio=count_reinicio+1
						hoja_1['C15'].value=vigencia.fecha_inicio
			if count_reinicio==0:
				hoja_1['C15'].value='El contrato no ha surtido reinicios'


		if count_suspension==0:
			hoja_1['C12'].value='El contrato a la fecha no se ha visto suspendido por las partes'
			hoja_1['C13'].value='El contrato a la fecha no se ha visto suspendido por las partes'
			hoja_1['C14'].value='El contrato a la fecha no se ha visto suspendido por las partes'
			hoja_1['C15'].value='El contrato no ha surtido reinicios debido a que no ha sido suspendido'

		
			

		if count_prorroga==0:
			hoja_1['C16'].value='El contrato no ha surtido prorrogas'
			hoja_1['C17'].value='El contrato no ha surtido prorrogas'
			hoja_1['C18'].value='El contrato no ha surtido prorrogas'

		if count_liquidacion==0:			
			hoja_1['B41'].value='El contrato no cuenta con acta de liquidación suscrita'

		hoja_1['B22'].value=empresa_contratista+' no tiene registro presupuestal conforme al estatuto orgánico del presupuesto'
		hoja_1['B23'].value=empresa_contratista+' no tiene registro presupuestal conforme al estatuto orgánico del presupuesto'

		hoja_1['B24'].value=vigencias_contrato[0].valor


		p_vigencia=Contrato_Vigencia_anual.objects.filter(contrato_id=contrato.id).exists()

		if p_vigencia:
			p_vigencia=Contrato_Vigencia_anual.objects.filter(contrato_id=contrato.id)
			iteraccion=24
			l=0
			for p_v in p_vigencia:
				iteraccion=iteraccion+1
				if l==0:
					hoja_1['A'+str(iteraccion)].value='VIGENCIA ORDINARIA '+ str(p_v.ano)+' ('+str(int(p_v.porcentaje*100))+'%)'
				else:
					hoja_1['A'+str(iteraccion)].value='VIGENCIA '+ str(p_v.ano)+' ('+str(int(p_v.porcentaje*100))+'%)'

				#hoja_1['B'+str(iteraccion)].value=hoja_1['B24'].value*float(p_v.porcentaje)
				hoja_1['B'+str(iteraccion)].value='=+B24*'+str(p_v.porcentaje)
				l=l+1
		
		financiacion=Contrato_Financiacion.objects.filter(contrato_id=contrato.id,es_cofinanciacion=True,fecha_suscripcion__lte=limite_fecha).exists()
		if financiacion:
			financiacion=Contrato_Financiacion.objects.filter(contrato_id=contrato.id,es_cofinanciacion=True,fecha_suscripcion__lte=limite_fecha)
			hoja_1['C28'].value=financiacion.last().empresa.nombre
			hoja_1['B29'].value=str(financiacion.last().valor)
			confinanciacion_condiciones=Contrato_financiacion_condicion.objects.filter(financiacion_id=financiacion.last().id,fecha_suscripcion__lte=limite_fecha).exists()
			if confinanciacion_condiciones:
				confinanciacion_condiciones=Contrato_financiacion_condicion.objects.filter(financiacion_id=financiacion.last().id,fecha_suscripcion__lte=limite_fecha)
				aux = ''
				q=0
				for val in confinanciacion_condiciones:
					q=q+1
					if q==1:
						aux=aux+'Condiciones  de confinanciación del contrato '+str(contrato.numero)
						aux=aux+'\n'+str(q)+'. '+val.condicion
					else:
						aux=aux+'\n'+str(q)+'. '+val.condicion
				hoja_1['B30'].value=aux
				#import pdb; pdb.set_trace()
				length = len(hoja_1['B30'].value)
				length = length/5
				hoja_1.row_dimensions[30].height = length
			else:
				hoja_1['B30'].value='N/A'
		else:
			hoja_1['C28'].value='N/A'
			hoja_1['B29'].value='N/A'
			hoja_1['B30'].value='N/A'

		hoja_1['B33'].value='=+B24+B32'

		#import pdb; pdb.set_trace()
		adicion_presupuestal=VigenciaContrato.objects.filter(contrato_id=contrato.id,tipo_id=17,fecha_inicio__lte=limite_fecha).exists()
		if adicion_presupuestal:
			adicion_presupuestal=VigenciaContrato.objects.filter(contrato_id=contrato.id,tipo_id=17,fecha_inicio__lte=limite_fecha).aggregate(suma_adicion=Sum('valor'),count=Count('id'))
			hoja_1['B32'].value=adicion_presupuestal['suma_adicion']

			adicion_presupuestal=VigenciaContrato.objects.filter(contrato_id=contrato.id,tipo_id=17,fecha_inicio__lte=limite_fecha)
			count=0
			for adicion in 	adicion_presupuestal:
				if adicion.valor:
					count=count+1
			hoja_1['B31'].value='El contrato ha surtido ('+str(count)+') adiciones presupuestal'
		else:
			hoja_1['B31'].value='El contrato no ha suscrito adiciones presupuestal'
			hoja_1['B32'].value=0


		val_administracion=Contrato_Administracion.objects.filter(contrato_id=contrato.id)
		if val_administracion:
			#hoja_1['B34'].value=hoja_1['B33'].value*0.5
			hoja_1['B34'].value='=+B33*0.05'
		else:
			hoja_1['B34'].value='N/A'

		banco = FinancieroCuenta.objects.filter(contrato_id=contrato.id).exists()
		if banco:
			banco = FinancieroCuenta.objects.get(contrato_id=contrato.id)
			hoja_1['B35'].value=banco.fiduciaria
			hoja_1['B38'].value=banco.numero
		else:
			hoja_1['B35'].value='N/A'
			hoja_1['B38'].value='N/A'


		fiduciaria= Contrato.objects.filter(mcontrato_id=contrato.id, tipo_contrato_id=15, nombre__icontains='ENCARGO FIDUCIARIO').exists()
		if fiduciaria:
			contrato_fiduciario= Contrato.objects.get(mcontrato_id=contrato.id, tipo_contrato_id=15, nombre__icontains='ENCARGO FIDUCIARIO')
			hoja_1['B36'].value=str(contrato_fiduciario.contratista)
			hoja_1['B37'].value=str(contrato_fiduciario.nombre)
		else:
			hoja_1['B36'].value='N/A'
			hoja_1['B37'].value='N/A'

		hoja_1['B39'].value='=+B37'
		# import pdb; pdb.set_trace()
		
		supervisor=Contrato_supervisor.objects.filter(empresa_id=198,cargo=1,contrato_id=contrato.id)
		if supervisor:
			supervisor = supervisor.last()
			# supervisor=Funcionario.objects.get(empresa_id=198,cargo_id=84,activo=1)
			hoja_1['B43'].value=supervisor.funcionario.persona.nombres+' '+supervisor.funcionario.persona.apellidos
		
		else:
			hoja_1['B43'].value='No se ha suscrito ninguna persona para este puesto'
			supervisor = False

		supervisor_apoyo=Contrato_supervisor.objects.filter(empresa_id=198,cargo=2,contrato_id=contrato.id)
		if supervisor_apoyo:
			supervisor_apoyo = supervisor_apoyo.last()
			# supervisor_apoyo=Funcionario.objects.get(empresa_id=198,cargo=85,activo=1)
			hoja_1['C43'].value=supervisor_apoyo.funcionario.persona.nombres+' '+supervisor_apoyo.funcionario.persona.apellidos
		
		else:
			if supervisor:
				hoja_1['C43'].value=hoja_1['B43'].value			
			else:
				hoja_1['C43'].value='No se ha suscrito ninguna persona para este puesto'
			

		
		vigencia_poliza = VigenciaPoliza.objects.filter(poliza__contrato_id=contrato.id,poliza__tipo_id=45,fecha_inicio__lte=limite_fecha).exists()
		if vigencia_poliza:
			vigencia_poliza = VigenciaPoliza.objects.filter(poliza__contrato_id=contrato.id,poliza__tipo_id=45,fecha_inicio__lte=limite_fecha)			
			
			hoja_1['B40'].value= str(vigencia_poliza[0].aseguradora.nombre + ' - ' +vigencia_poliza[0].numero)

			aprobacion=VigenciaPoliza_AprobacionMME.objects.filter(vigencia_poliza_id=vigencia_poliza[0].id,fecha__lte=limite_fecha).exists()
			if aprobacion:
				aprobacion=VigenciaPoliza_AprobacionMME.objects.get(vigencia_poliza_id=vigencia_poliza[0].id,fecha__lte=limite_fecha)
				hoja_1['F70'].value=aprobacion.fecha
				hoja_1['F70'].number_format='dd/mm/yyyy'
			else:
				hoja_1['F70'].value='No se ha suscrito la fecha de aprobacion del MME'
			
		else:
			hoja_1['B40'].value='El contrato no cuenta con poliza de cumplimiento suscrita'
			hoja_1['C70'].value='El contrato no cuenta con poliza de cumplimiento suscrita'
			hoja_1['C71'].value='El contrato no cuenta con poliza de cumplimiento suscrita'

			hoja_1['D70'].value='El contrato no cuenta con la poliza suscrita'
			hoja_1['E70'].value='El contrato no cuenta con la poliza suscrita'

			hoja_1['F70'].value='No se ha suscrito la fecha de aprobacion del MME'


		

		hoja_1['A70'].value='=+B40'

		if vigencia_poliza.count()==1:

			# hoja_1['C70']='=B24*0,2'
			hoja_1['C70'].value=vigencia_poliza[0].valor

		elif vigencia_poliza.count()>1:
			aux_vigencia_poliza=VigenciaPoliza.objects.filter(poliza__contrato_id=contrato.id,poliza__tipo_id=45, reemplaza=True,fecha_inicio__lte=limite_fecha)

			if aux_vigencia_poliza.count()==0:
				hoja_1['C70'].value=vigencia_poliza[0].valor

			elif aux_vigencia_poliza.count()>0:
				item = aux_vigencia_poliza.last()
				hoja_1['C70'].value=item.valor

		poliza_cumplimiento=VigenciaPoliza.objects.filter(poliza__contrato_id=contrato.id,poliza__tipo_id=45,fecha_inicio__lte=limite_fecha).aggregate(Min('fecha_inicio'))
		hoja_1['D70'].value=poliza_cumplimiento['fecha_inicio__min'];

		qset = (Q(poliza__contrato_id=contrato.id) & Q(poliza__tipo_id=45) & Q(fecha_final__isnull=False) & Q(fecha_inicio__lte=limite_fecha))
		item = VigenciaPoliza.objects.filter(qset).last()
		hoja_1['E70'].value=item.fecha_final





		vigencia_poliza2 = VigenciaPoliza.objects.filter(poliza__contrato_id=contrato.id,poliza__tipo_id=38,fecha_inicio__lte=limite_fecha).exists()
		if vigencia_poliza2:
			vigencia_poliza2 = VigenciaPoliza.objects.filter(poliza__contrato_id=contrato.id,poliza__tipo_id=38,fecha_inicio__lte=limite_fecha)

			if vigencia_poliza2.count()==1:				
				hoja_1['C71'].value=vigencia_poliza2[0].valor

			elif vigencia_poliza2.count()>1:
				aux_vigencia_poliza2=VigenciaPoliza.objects.filter(poliza__contrato_id=contrato.id,poliza__tipo_id=38, reemplaza=True,fecha_inicio__lte=limite_fecha)

				if aux_vigencia_poliza2.count() == 0:
					hoja_1['C71'].value=vigencia_poliza2[0].valor

				elif aux_vigencia_poliza2.count() > 0:	
					item = aux_vigencia_poliza2.last()
					hoja_1['C71'].value=item.valor

			poliza2_cumplimiento=VigenciaPoliza.objects.filter(poliza__contrato_id=contrato.id,poliza__tipo_id=38,fecha_inicio__lte=limite_fecha).aggregate(Min('fecha_inicio'))
			hoja_1['D71'].value=poliza2_cumplimiento['fecha_inicio__min'];

			qset = (Q(poliza__contrato_id=contrato.id) & Q(poliza__tipo_id=38) & Q(fecha_final__isnull=False) & Q(fecha_inicio__lte=limite_fecha))
			item = VigenciaPoliza.objects.filter(qset).last()
			hoja_1['E71'].value=item.fecha_final

		else:
			
			hoja_1['C70'].value=float(hoja_1['B24'].value*0.2)
			hoja_1['C71'].value=hoja_1['B25'].value

			
			hoja_1['D71'].value=hoja_1['D70'].value.date()
			hoja_1['E71'].value=hoja_1['E70'].value.date()

		hoja_1['D70'].number_format='dd/mm/yyyy'
		hoja_1['E70'].number_format='dd/mm/yyyy'
		hoja_1['D71'].number_format='dd/mm/yyyy'
		hoja_1['E71'].number_format='dd/mm/yyyy'

		#hoja_1.insert_rows(7)
		
		
		##2,-BALANCES FINANCIEROS

		nombrehoja=doc.get_sheet_names()[2]
		hoja_2 = doc.get_sheet_by_name(nombrehoja)

		

		valor_ingreso_nov=FinancieroCuentaMovimiento.objects.filter(cuenta__contrato_id=contrato.id, tipo_id=31).exists()
		if valor_ingreso_nov:
			valor_ingreso_nov=FinancieroCuentaMovimiento.objects.filter(cuenta__contrato_id=contrato.id, tipo_id=31)

			#import pdb; pdb.set_trace()

			fecha_inicio_ano=valor_ingreso_nov.first().fecha.year
			fecha_inicio_mes=valor_ingreso_nov.first().fecha.month
			

			fecha_inicio=str(fecha_inicio_ano)+'-'+str(fecha_inicio_mes)+'-01'
			fecha_inicio=datetime.strptime(str(fecha_inicio),'%Y-%m-%d').date()

			#fecha_inicio_dia=fecha_inicio+timedelta(1 * 30)
			
			if int(fecha_inicio.month+1)<13:
				fecha_fin=str(fecha_inicio.year)+'-'+str(fecha_inicio.month+1)+'-'+'01'
				fecha_fin=datetime.strptime(str(fecha_fin),'%Y-%m-%d').date()
				fecha_fin=fecha_fin-timedelta(days=1)

			else:
				fecha_fin=str(fecha_inicio.year+1)+'-01-01'
				fecha_fin=datetime.strptime(str(fecha_fin),'%Y-%m-%d').date()
				fecha_fin=fecha_fin-timedelta(days=1)	


			if str(fecha_fin)>str(limite_fecha):
				fecha_fin=limite_fecha

			i=54
			j=96

			condicion=True

			#import pdb; pdb.set_trace()

			while condicion:

				i=i+1
				j=j+1

				movimientos=FinancieroCuentaMovimiento.objects.filter(
					cuenta__contrato_id=contrato.id,fecha__gte=fecha_inicio,
					fecha__lte=fecha_fin).exists()

				celda_mes='A'+str(i)
				celda_mes_2='B'+str(j)

				if movimientos:
					recaudo_1=0.0
					rendimiento_1=0.0
					pagos_1=0.0
					p_rendimiento_1=0.0
					retefuente_1=0.0
					c_p_proveedores_1=0.0
					gmf_1=0.0
					iva_1=0.0

					
					valor_ingreso_rend=FinancieroCuentaMovimiento.objects.filter(cuenta__contrato_id=contrato.id,
						tipo_id=32,
						fecha__gte=fecha_inicio,
						fecha__lte=fecha_fin).aggregate(suma_rendimiento=Sum('valor'))

					
					if valor_ingreso_rend['suma_rendimiento']:
						rendimiento_1=rendimiento_1+int(valor_ingreso_rend['suma_rendimiento'])
					
					
					valor_egreso=FinancieroCuentaMovimiento.objects.filter(cuenta__contrato_id=contrato.id,
						tipo_id=29,
						fecha__gte=fecha_inicio,
						fecha__lte=fecha_fin).order_by('fecha')


					for val in valor_egreso:
						if val.descripcion.count('RETEFUENTE')>0 or val.descripcion.count('Retefuente')>0 or val.descripcion.count('Retefuente remuneracion especial'):
							retefuente_1=retefuente_1+float(val.valor)

						elif val.descripcion.count('IMPUESTO')>0 or val.descripcion.count('Gravamen')>0 or  val.descripcion.count('GRAVAMEN')>0 or val.descripcion.count('GMF')>0:
							gmf_1=gmf_1+float(val.valor)	
						
						elif val.descripcion.count('IVA')>0 or val.descripcion.count('iva')>0 or val.descripcion.count('I.V.A')>0 or val.descripcion.count('i.v.a')>0:
							iva_1=iva_1+float(val.valor)

						elif val.descripcion.count('DEB TRANSF OFICINA 000000000000')>0 or val.descripcion.count('DEVOLUCION RENDIMIENTOS')>0 or val.descripcion.count('Cargo por Traslado Sebra Enviado')>0 or val.descripcion.count('ajuste Vr Ajuste')>0 or val.descripcion.count('ajuste Vr Reversion')>0 or val.descripcion.count('REINTEGRO')>0 or val.descripcion.count('DEVOLUCIÓN RENDIMIENTOS')>0:
							p_rendimiento_1=p_rendimiento_1+float(val.valor)

						elif val.descripcion.count('PROVEEDOR')>0:
							c_p_proveedores_1=c_p_proveedores_1+float(val.valor)

						elif val.descripcion.count('DEVOLUCION AL MME')>0:
							pagos_1=pagos_1+float(val.valor)

					valor_ingreso=FinancieroCuentaMovimiento.objects.filter(cuenta__contrato_id=contrato.id,
						tipo_id=31,
						fecha__gte=fecha_inicio,
						fecha__lte=fecha_fin).order_by('fecha')

					for val in valor_ingreso:
						if val.descripcion.count('APORTE MME')>0 or val.descripcion.count('COMP')>0 or val.descripcion.count('Ministerio De Minas')>0 or val.descripcion.count('RECAUDO')>0:					
							recaudo_1=recaudo_1+float(val.valor)
						elif val.descripcion.count('IVA')>0 or val.descripcion.count('iva')>0 or val.descripcion.count('I.V.A')>0 or val.descripcion.count('i.v.a')>0:
							iva_1=iva_1-float(val.valor)
						elif val.descripcion.count('DEB TRANSF OFICINA 000000000000')>0 or val.descripcion.count('DEVOLUCION RENDIMIENTOS')>0 or val.descripcion.count('Cargo por Traslado Sebra Enviado')>0 or val.descripcion.count('ajuste Vr Ajuste')>0 or val.descripcion.count('ajuste Vr Reversion')>0 or val.descripcion.count('REINTEGRO')>0 or val.descripcion.count('DEVOLUCIÓN RENDIMIENTOS')>0:
							p_rendimiento_1=p_rendimiento_1-float(val.valor)
						elif val.descripcion.count('PROVEEDOR')>0:
							c_p_proveedores_1=c_p_proveedores_1-float(val.valor)

					hoja_2[str(celda_mes)].value=fecha_inicio
					hoja_2[str(celda_mes)].number_format = 'mmm-yy' 

					hoja_2['C'+str(i)].value=recaudo_1
					hoja_2['D'+str(i)].value=rendimiento_1
					hoja_2['F'+str(i)].value=pagos_1
					hoja_2['G'+str(i)].value=p_rendimiento_1
					hoja_2['H'+str(i)].value=retefuente_1
					hoja_2['I'+str(i)].value=gmf_1
					hoja_2['J'+str(i)].value=iva_1
					hoja_2['K'+str(i)].value=c_p_proveedores_1
					


					if rendimiento_1:
						hoja_2[str(celda_mes_2)].value=fecha_inicio
						hoja_2[str(celda_mes_2)].number_format = 'mmm-yy' 
						hoja_2['D'+str(j)].value=rendimiento_1
						hoja_2['A'+str(j)].fill=PatternFill(start_color="FF0D0D", end_color="FF0D0D", fill_type = "solid")
						hoja_2['C'+str(j)].fill=PatternFill(start_color="FF0D0D", end_color="FF0D0D", fill_type = "solid")


					if int(fecha_inicio.month+1)<13:
						fecha_inicio=str(fecha_inicio.year)+'-'+str(fecha_inicio.month+1)+'-'+'01'
						fecha_inicio=datetime.strptime(str(fecha_inicio),'%Y-%m-%d').date()

					else:
						fecha_inicio=str(fecha_inicio.year+1)+'-01-01'
						fecha_inicio=datetime.strptime(str(fecha_inicio),'%Y-%m-%d').date()

							
					if int(fecha_inicio.month+1)<13:
						fecha_fin=str(fecha_inicio.year)+'-'+str(fecha_inicio.month+1)+'-'+'01'
						fecha_fin=datetime.strptime(str(fecha_fin),'%Y-%m-%d').date()
						fecha_fin=fecha_fin-timedelta(days=1)

					else:
						fecha_fin=str(fecha_inicio.year+1)+'-01-01'
						fecha_fin=datetime.strptime(str(fecha_fin),'%Y-%m-%d').date()
						fecha_fin=fecha_fin-timedelta(days=1)

					if str(fecha_fin)>str(limite_fecha):
						fecha_fin=limite_fecha

					condicion=True

				else:
					
					movimientos=FinancieroCuentaMovimiento.objects.filter(cuenta__contrato_id=contrato.id).order_by('fecha')

					if movimientos.last().fecha<fecha_inicio:
						condicion=False
					else:					
						condicion=True

						if int(fecha_inicio.month+1)<13:
							fecha_inicio=str(fecha_inicio.year)+'-'+str(fecha_inicio.month+1)+'-'+'01'
							fecha_inicio=datetime.strptime(str(fecha_inicio),'%Y-%m-%d').date()

						else:
							fecha_inicio=str(fecha_inicio.year+1)+'-01-01'
							fecha_inicio=datetime.strptime(str(fecha_inicio),'%Y-%m-%d').date()

								
						if int(fecha_inicio.month+1)<13:
							fecha_fin=str(fecha_inicio.year)+'-'+str(fecha_inicio.month+1)+'-'+'01'
							fecha_fin=datetime.strptime(str(fecha_fin),'%Y-%m-%d').date()
							fecha_fin=fecha_fin-timedelta(days=1)

						else:
							fecha_fin=str(fecha_inicio.year+1)+'-01-01'
							fecha_fin=datetime.strptime(str(fecha_fin),'%Y-%m-%d').date()
							fecha_fin=fecha_fin-timedelta(days=1)

						if str(fecha_fin)>str(limite_fecha):
							fecha_fin=limite_fecha
		i=9
		#import pdb; pdb.set_trace()
		if p_vigencia:
			#p_vigencia=Contrato_Vigencia_anual.objects.filter(contrato_id=contrato.id)
			for vig in p_vigencia:			
				contrato_desembolso=Contrato_Desembolso.objects.filter(vigencia_anual_id=vig.id).exists()
				if contrato_desembolso:
					contrato_desembolso=Contrato_Desembolso.objects.filter(vigencia_anual_id=vig.id)
					
					for desembolso in contrato_desembolso:
						i=i+1
						hoja_2['B'+str(i)].value=desembolso.requisito
						hoja_2['C'+str(i)].value=desembolso.vigencia_anual.ano
						hoja_2['D'+str(i)].value=desembolso.valor_requerido

						contrato_desembolso_desembolsados=Contrato_Desembolso_desembolsados.objects.filter(desembolso_id=desembolso.id,fecha_suscripcion__lte=limite_fecha).exists()
						if contrato_desembolso_desembolsados:
							contrato_desembolso_desembolsados=Contrato_Desembolso_desembolsados.objects.filter(desembolso_id=desembolso.id,fecha_suscripcion__lte=limite_fecha).aggregate(suma_valor_desembolsado=Sum('valor_desembolsado'))
							hoja_2['E'+str(i)].value=contrato_desembolso_desembolsados['suma_valor_desembolsado']
							contrato_desembolso_desembolsados=Contrato_Desembolso_desembolsados.objects.filter(desembolso_id=desembolso.id,fecha_suscripcion__lte=limite_fecha)
							hoja_2['G'+str(i)].value=contrato_desembolso_desembolsados.last().fecha_suscripcion
							hoja_2['G'+str(i)].number_format='dd/mm/yyyy'
							hoja_2['H'+str(i)].fill=PatternFill(start_color="FF0D0D", end_color="FF0D0D", fill_type = "solid")
							hoja_2['I'+str(i)].fill=PatternFill(start_color="FF0D0D", end_color="FF0D0D", fill_type = "solid")
			i=i+1
			while i<23:			
				hoja_2.row_dimensions[i].hidden = True
				i=i+1
		else:
			
			i=i+1
			while i<23:
				hoja_2['A'+str(i)].value='N/A'
				hoja_2['B'+str(i)].value='N/A'
				hoja_2['C'+str(i)].value='N/A'
				hoja_2['F'+str(i)].value=0
				if i>12:
					hoja_2.row_dimensions[i].hidden = True
				i=i+1
		j=34
		contrato_remuneracion=Contrato_Remuneracion.objects.filter(contrato_id=contrato.id).exists()
		if contrato_remuneracion:
			contrato_remuneracion=Contrato_Remuneracion.objects.filter(contrato_id=contrato.id)
			for remun in contrato_remuneracion:
				j=j+1
				hoja_2['B'+str(j)].value=remun.requisito
				hoja_2['C'+str(j)].value=remun.valor_requerido

				contrato_Remuneracion_pagos=Contrato_Remuneracion_pagos.objects.filter(remuneracion_id=remun.id).exists()
				if contrato_Remuneracion_pagos:
					contrato_Remuneracion_pagos=Contrato_Remuneracion_pagos.objects.filter(remuneracion_id=remun.id).aggregate(suma_valor_pagado=Sum('valor_pagado'))
					hoja_2['D'+str(j)].value=contrato_Remuneracion_pagos['suma_valor_pagado']
					contrato_Remuneracion_pagos=Contrato_Remuneracion_pagos.objects.filter(remuneracion_id=remun.id)
					hoja_2['F'+str(j)].value=contrato_Remuneracion_pagos.last().fecha_suscripcion
					hoja_2['F'+str(j)].number_format='dd/mm/yyyy'
					hoja_2['G'+str(j)].fill=PatternFill(start_color="FF0D0D", end_color="FF0D0D", fill_type = "solid")
					hoja_2['H'+str(j)].fill=PatternFill(start_color="FF0D0D", end_color="FF0D0D", fill_type = "solid")
			j=j+1
			while j<48:			
				hoja_2.row_dimensions[j].hidden = True
				j=j+1
		else:
			hoja_2['B27'].value=0
			hoja_2['B28'].value=0
			hoja_2['B29'].value=0

			j=j+1
			while j<48:
				hoja_2['A'+str(j)].value='N/A'
				hoja_2['B'+str(j)].value='N/A'
				hoja_2['E'+str(j)].value=0
				if j>36:
					hoja_2.row_dimensions[j].hidden = True
				j=j+1

		

		##FASCP-01A-PROYEC
		nombrehoja=doc.get_sheet_names()[3]
		hoja_3 = doc.get_sheet_by_name(nombrehoja)		

		hoja_3['B3'].value=limite_fecha
		hoja_3['B3'].number_format='mmm-yy'

		hoja_3['L10'].value=empresa_contratista+' no emplea contrato de suministros'

		ActaAsign=ActaAsignacionRecursosContrato.objects.filter(contrato_id=contrato.id).exists()
		if ActaAsign:
			ActaAsign=ActaAsignacionRecursosContrato.objects.filter(contrato_id=contrato.id)
			hoja_3['O9'].value=ActaAsign.last().actaAsignacion.nombre
		
		else:
			hoja_3['O9'].value=hoja_1['B5'].value

		codigo_proyecto=Proyecto_proyecto_codigo.objects.filter(proyecto_id=proyecto.id).exists()
		if codigo_proyecto:
			codigo_proyecto=Proyecto_proyecto_codigo.objects.get(proyecto_id=proyecto.id)
			hoja_3['B9'].value=codigo_proyecto.codigo
		else:
			hoja_3['B9'].value='N/A'
		

		reportado_usuario_persona_id=0

		if contrato.contratista.nombre.count('AIR-E'):
			reportado_usuario_persona_id=5073
		elif contrato.contratista.nombre.count('MAR'):			
			reportado_usuario_persona_id=1595

		reportado_usuario_persona=Funcionario.objects.filter(persona_id=request.user.usuario.persona.id,activo=1).exists()
		if reportado_usuario_persona:
			reportado_usuario_persona=Funcionario.objects.get(persona_id=request.user.usuario.persona.id,activo=1)
			hoja_3['B10'].value=reportado_usuario_persona.persona.nombres+' '+reportado_usuario_persona.persona.apellidos

		# hoja_3['F9'].value=str(datetime.now().strftime('%d/%m/%Y'))
		hoja_3['F9'].value=fecha_actual
		hoja_3['F9'].number_format='dd/mm/yyyy'
		hoja_3['H10'].value=proyecto.municipio.nombre
		hoja_3['H11'].value=proyecto.municipio.departamento.nombre



		interventoria=Contrato.objects.filter(mcontrato_id=contrato.id,tipo_contrato_id=9).exists()
		if interventoria:
			interventoria=Contrato.objects.filter(mcontrato_id=contrato.id,tipo_contrato_id=9)
			if str(interventoria.first().fecha_inicio())<str(limite_fecha):
				interventoria=interventoria.first()
			else:
				interventoria=False
		else:
			interventoria = proyecto.contrato.filter(tipo_contrato_id = 9).exists()
			if interventoria:
				# import pdb; pdb.set_trace()
				interventoria = proyecto.contrato.filter(tipo_contrato_id = 9)
				if str(interventoria.first().fecha_inicio())<str(limite_fecha):
					interventoria=interventoria.first()
				else:
					interventoria=False


		if interventoria:
			hoja_3['L11'].value=interventoria.numero
		else:
			hoja_3['L11'].value='No hay contrato de interventoría suscrito'

		
		if not financiacion:
			hoja_3['F12'].value=''

		contrato_obra = proyecto.contrato.filter(tipo_contrato_id = 8)
		# contrato_obra=Contrato.objects.filter(mcontrato_id=contrato.id,tipo_contrato_id=8).exists()
		if contrato_obra:
			contrato_obra = proyecto.contrato.filter(tipo_contrato_id = 8)
			# contrato_obra=Contrato.objects.filter(mcontrato_id=contrato.id,tipo_contrato_id=8)
			if str(contrato_obra.first().fecha_inicio())<str(limite_fecha):
				contrato_obra=True
			else:
				contrato_obra=False
		if contrato_obra:
			contrato_obra = proyecto.contrato.filter(tipo_contrato_id = 8)
			# contrato_obra=Contrato.objects.filter(mcontrato_id=contrato.id,tipo_contrato_id=8)
			contrato_obra=contrato_obra.first()
			hoja_3['L9'].value=contrato_obra.numero

			aux_fecha_inicio=datetime.strptime(str(contrato_obra.fecha_inicio()),'%Y-%m-%d').date()
		
			ahora=datetime.now()
			ahora=ahora.date()
			print(aux_fecha_inicio)
			if ahora<contrato_obra.fecha_fin():
				diferencia_fechas=ahora-aux_fecha_inicio
				diferencia_fechas=int(diferencia_fechas.days)
				diferencia_fechas=int(diferencia_fechas/30)
				hoja_3['N14'].value=diferencia_fechas
			else:
				diferencia_fechas=contrato_obra.fecha_fin()-aux_fecha_inicio
				diferencia_fechas=int(diferencia_fechas.days)
				diferencia_fechas=int(diferencia_fechas/30)
				hoja_3['N14'].value=diferencia_fechas

		else:
			hoja_3['L9'].value='No hay contrato de obra suscrito'

		valor_pagado_giros = 0 
		valor_pagado_factura = 0

		contratos_asociados = proyecto.contrato.values('id')

		

		# import pdb; pdb.set_trace()
		tranasaciones_pagadas = DetalleGiro.objects.filter(fecha_pago__lte=limite_fecha,encabezado__contrato__in=contratos_asociados,estado_id=enumEstados.pagado).aggregate(Sum('valor_girar'))
		valor_pagado_giros = tranasaciones_pagadas['valor_girar__sum'] if tranasaciones_pagadas['valor_girar__sum'] else 0

		facturas_pagas = FacturaProyecto.objects.filter(factura__fecha_pago__lte=limite_fecha,proyecto_id=proyecto.id,factura__pagada=True).aggregate(Sum('factura__valor_factura'))
		valor_pagado_factura = facturas_pagas['factura__valor_factura__sum'] if facturas_pagas['factura__valor_factura__sum'] else 0

		hoja_3['K13'].value= valor_pagado_giros + valor_pagado_factura
		

		servidumbres=Servidumbre_predio.objects.filter(expediente__proyecto_id=proyecto.id).exists()

		if servidumbres:
			hoja_3['M15'].value=Servidumbre_predio.objects.filter(expediente__proyecto_id=proyecto.id).count()
			hoja_3['K15'].value='Si'
			#------
			count_servidumbres_pendientes=0
			#------
			obj=Servidumbre_predio.objects.filter(expediente__proyecto_id=proyecto.id)

			for predio in obj:
				total_documentos = float(Servidumbre_documento.objects.filter(
				grupo_documento__id=predio.grupo_documento.id).count())
				auxiliar =0
				count = 0

				Servidumbredocumento = Servidumbre_documento.objects.filter(grupo_documento__id=predio.grupo_documento.id)

				for ser_doc in Servidumbredocumento:
						
					auxiliar = float (Servidumbre_predio_documento.objects.filter(
					predio__id=predio.id, documento__id =ser_doc.id ).count())

					if (auxiliar>0):
						count = count +1;					
		
				
				if total_documentos > 0:
					aux=round((count / total_documentos)*100,2)
					if aux<100:
						count_servidumbres_pendientes=count_servidumbres_pendientes+1

		
			hoja_3['P15'].value=count_servidumbres_pendientes
			#------

		else:
			hoja_3['K15'].value='A la fecha no se ha identificado'
			hoja_3['P15'].value='A la fecha no se ha identificado'
		#-- Diseños completos
		Rellenar_actividades(hoja_3,'B15',proyecto,1)		

		#-- Planos completos 
		Rellenar_actividades(hoja_3,'D15',proyecto,2)

		#-- Especificaciones técnicas
		Rellenar_actividades(hoja_3,'F15',proyecto,3)

		#-- Elaboración de pliegos de condiciones
		Rellenar_actividades(hoja_3,'B16',proyecto,4)

		#-- Fecha publicación de pliegos para contratos de obra
		Rellenar_actividades(hoja_3,'D16',proyecto,5)

		#-- Fecha publicación pliegos contratos de interventoría
		Rellenar_actividades(hoja_3,'F16',proyecto,6)

		#-- Fecha aprobación ofertas para contratos de obra
		Rellenar_actividades(hoja_3,'K16',proyecto,7)

		#-- Fecha aprobación ofertas para contratos de interventoría
		Rellenar_actividades(hoja_3,'N16',proyecto,8)

		#-- Revisión y aprobación de diseños, especificaciones y presupuesto
		Rellenar_actividades(hoja_3,'B17',proyecto,9)

		#-- Fecha de revisión y aprobación de diseños, especificaciones y presupuesto
		Rellenar_actividades(hoja_3,'D17',proyecto,10)

		#-- Requiere Permisos de Paso
		Rellenar_actividades(hoja_3,'F17',proyecto,11)

		#-- Numero de permisos suscritos (anexar)
		Rellenar_actividades(hoja_3,'I17',proyecto,12)
		#-- Numero de permisos pendientes
		Rellenar_actividades(hoja_3,'K17',proyecto,13)

		#-- Numero de Actas de socializacion suscritas (Anexar)
		Rellenar_actividades(hoja_3,'M17',proyecto,14)

		#-- Requiere Licencias ambientales
		Rellenar_actividades(hoja_3,'P17',proyecto,15)

		#-- Numero de Licencias suscritas (Anexar)
		Rellenar_actividades(hoja_3,'B18',proyecto,16)

		#-- Numero de Licencias pendientes
		Rellenar_actividades(hoja_3,'D18',proyecto,17)

		#-- Requiere Plan de aprovechamiento Forestal
		Rellenar_actividades(hoja_3,'F18',proyecto,18)

		#-- Numero de Planes de aprovechamiento Forestal suscritos(anexar)
		Rellenar_actividades(hoja_3,'I18',proyecto,19)

		#-- Numero de Planes de aprovechamiento Forestal pendientes
		Rellenar_actividades(hoja_3,'K18',proyecto,20)
		#-- Requiere Permisos ambientales
		Rellenar_actividades(hoja_3,'M18',proyecto,21)

		#-- Numero de Permisos suscritos (Anexar)
		Rellenar_actividades(hoja_3,'P18',proyecto,22)

		#-- Numero de Permisos pendientes
		Rellenar_actividades(hoja_3,'B19',proyecto,23)

		#-- Requiere certificado RETIE(anexar)
		Rellenar_actividades(hoja_3,'D19',proyecto,24)

		#-- Requiere certificado de conformidad de producto RETIE  (Anexar)
		Rellenar_actividades(hoja_3,'F19',proyecto,25)

		#-- Fecha de expedicion Certificado conformidad de RETIE
		Rellenar_actividades(hoja_3,'H19',proyecto,26)

		#-- Fecha de expedicion Certificado de RETIE
		Rellenar_actividades(hoja_3,'K19',proyecto,27)

		#-- Numero de certificaciones de Zonas de Alto Riesgo suscritas(anexar)
		Rellenar_actividades(hoja_3,'P19',proyecto,28)

		#-- Numero de certificaciones de Zonas de Alto Riesgo pendientes
		Rellenar_actividades(hoja_3,'B20',proyecto,29)

		#-- Requiere compra de predios
		Rellenar_actividades(hoja_3,'D20',proyecto,30)

		#-- Numero de predios comprados(anexar)
		Rellenar_actividades(hoja_3,'F20',proyecto,31)

		#-- Numero de predios pendientes
		Rellenar_actividades(hoja_3,'H20',proyecto,32)

		#-- Se requieren consultas previas
		Rellenar_actividades(hoja_3,'K20',proyecto,33)

		#-- Numero de consultas previas realizadas
		Rellenar_actividades(hoja_3,'N20',proyecto,34)

		#-- Fecha de  Acta de Energizacion (Anexar)
		Rellenar_actividades(hoja_3,'P20',proyecto,35)	
		
		#-- Verificación de la localización de usuarios
		Rellenar_actividades(hoja_3,'B21',proyecto,36)	
			
		#-- Numero de Actas de verificación de la localización suscritas (Anexar)
		Rellenar_actividades(hoja_3,'D21',proyecto,37)	
		
		#-- Numero de Actas de verificación de la localización pendientes (Anexar)
		Rellenar_actividades(hoja_3,'F21',proyecto,38)			

		#-- Existe ente responsable de la financiación de redes internas
		Rellenar_actividades(hoja_3,'H21',proyecto,39)	

		#-- Nombre del ente responsable de la financiación
		Rellenar_actividades(hoja_3,'L21',proyecto,40)

		#-- Latitud y longitud de la cabecera municipal en grados		
		Rellenar_actividades(hoja_3,'P21',proyecto,41)

		telefono_administrador=0
		if contrato.contratista.nombre.count('AIR-E'):
			telefono_administrador='035 35-0444'
		elif contrato.contratista.nombre.count('MAR'):
			telefono_administrador='3611000'

		hoja_3['L22'].value=telefono_administrador
		##FASCP-01B-CONTRA-OBRA
		nombrehoja=doc.get_sheet_names()[4]
		hoja_4 = doc.get_sheet_by_name(nombrehoja)	

		if fondo=='PRONE':
			hoja_4['F14'].value='NORMALIZACION DE REDES ELECTRICAS BARRIO '+proyecto.nombre
		elif fondo=='FAER':
			hoja_4['F14'].value='ELECTRIFICACIÓN RURAL '+proyecto.nombre
			
		
		if contrato_obra:
			# contrato_obra = proyecto.contrato.filter(tipo_contrato_id = 8)
			# contrato_obra=Contrato.objects.filter(mcontrato_id=contrato.id,tipo_contrato_id=8)
			# contrato_obra=contrato_obra.first()
			fecha_inicio=datetime.strptime(str(contrato_obra.fecha_inicio()),'%Y-%m-%d').date()
			fecha_final=contrato_obra.fecha_fin()
			

			contrato_obra_meses=fecha_final-fecha_inicio
			contrato_obra_meses=int(contrato_obra_meses.days)
			contrato_obra_meses=round(contrato_obra_meses/30)
			hoja_4['A10'].value=contrato_obra.contratista.nombre
			hoja_4['C15'].value=contrato_obra_meses
			hoja_3['K14'].value=contrato_obra_meses
			hoja_4['P16'].value=fecha_final
			hoja_4['P16'].number_format='dd/mm/yyyy'

			vigencia_contrato_obra = VigenciaContrato.objects.filter(tipo_id__in=[18,19,102],contrato__mcontrato_id=contrato.id,contrato__tipo_contrato_id=8,fecha_inicio__lte=limite_fecha)
			cantidad_prorroga_meses = 0
			cantidad_suspension_meses = 0

			if len(vigencia_contrato_obra)>0:
				for vigencia_aux_obra in vigencia_contrato_obra:

					if vigencia_aux_obra.tipo.id==102:
						aux_vigencia_obra=vigencia_aux_obra.fecha_fin-vigencia_aux_obra.fecha_inicio
						cantidad_prorroga_meses+=round(int(aux_vigencia_obra.days)/30)

					elif vigencia_aux_obra.tipo.id==18:
						vigencia_aux_obra_fecha_fin = vigencia_contrato_obra.filter(pk__gte=vigencia_aux_obra.id,tipo_id=19).first().values('fecha_fin')
						if vigencia_aux_obra_fecha_fin:
							aux_vigencia_obra=vigencia_aux_obra_fecha_fin-vigencia_aux_obra.fecha_inicio
							cantidad_suspension_meses+=round(int(aux_vigencia_obra.days)/30)
						else:
							aux_vigencia_obra=limite_fecha-vigencia_aux_obra.fecha_inicio
							cantidad_suspension_meses+=round(int(aux_vigencia_obra.days)/30)

			if cantidad_prorroga_meses>0:
				hoja_4['L16'].value=cantidad_prorroga_meses
			else:
				hoja_4['L16'].value='No hay prorrogas'

			if cantidad_suspension_meses>0:
				hoja_4['H16'].value=cantidad_suspension_meses
			else:
				hoja_4['H16'].value='No hay suspensiones'


			aux_fecha=contrato_obra.fecha_firma
			if aux_fecha  and aux_fecha<limite_fecha:
				hoja_4['H15'].value=aux_fecha
				hoja_4['H15'].number_format='dd/mm/yyyy'
			else:
				hoja_4['H15'].value=contrato_obra.fecha_inicio()
				hoja_4['H15'].number_format='dd/mm/yyyy'
				
				

			aux_fecha=contrato_obra.fecha_acta_inicio
			if aux_fecha  and aux_fecha<limite_fecha:
				hoja_4['L15'].value=aux_fecha
				hoja_4['L15'].number_format='dd/mm/yyyy'
			else:
				hoja_4['L15'].value='No se ha suscrito la fecha de inicio del contrato de obra'

			anticipo_contrato_obra=DetalleGiro.objects.filter(encabezado__contrato_id=contrato_obra.id,estado__codigo=3).exists()
			if anticipo_contrato_obra:
				anticipo_contrato_obra=DetalleGiro.objects.filter(encabezado__contrato_id=contrato_obra.id,estado__codigo=3).aggregate(suma_valor=Sum('valor_girar'))
				hoja_4['J23'].value=anticipo_contrato_obra['suma_valor']
			

			vigencia_contrato_obra=VigenciaContrato.objects.filter(contrato__mcontrato_id=contrato.id,contrato__tipo_contrato_id=8,tipo_id=16,fecha_inicio__lte=limite_fecha).exists()
			if vigencia_contrato_obra:
				vigencia_contrato_obra=VigenciaContrato.objects.filter(contrato__mcontrato_id=contrato.id,contrato__tipo_contrato_id=8,tipo_id=16,fecha_inicio__lte=limite_fecha)
				hoja_4['C17'].value=vigencia_contrato_obra[0].valor

				vigencia_finalizacion=vigencia_contrato_obra.last()
				if vigencia_finalizacion.tipo.id==21:
					hoja_4['P15'].value=vigencia_finalizacion.fecha_inicio
					hoja_4['P15'].number_format='dd/mm/yyyy'
				else:
					hoja_4['P15'].value='El contrato no ha finalizado'

			else:				
				hoja_4['P15'].value='No hay contrato de obra suscrito'

			adicion_presupuestal_interventoria=VigenciaContrato.objects.filter(contrato_id=contrato_obra.id,tipo_id=17,fecha_inicio__lte=limite_fecha).exists()
			if adicion_presupuestal_interventoria:
				adicion_presupuestal_interventoria=VigenciaContrato.objects.filter(contrato_id=contrato_obra.id,tipo_id=17,fecha_inicio__lte=limite_fecha).aggregate(suma_valor=Sum('valor'))
				hoja_4['H17'].value=adicion_presupuestal_interventoria['suma_valor']


			vigencia_poliza_interven = VigenciaPoliza.objects.filter(poliza__contrato_id=contrato_obra.id,poliza__tipo_id=45,fecha_inicio__lte=limite_fecha).exists()
			if vigencia_poliza_interven:
				vigencia_poliza_interven = VigenciaPoliza.objects.filter(poliza__contrato_id=contrato_obra.id,poliza__tipo_id=45,fecha_inicio__lte=limite_fecha)
				hoja_4['D18'].value=vigencia_poliza_interven[0].numero
				hoja_4['A19'].value=str(vigencia_poliza_interven[0].aseguradora)
			else:
				hoja_4['D18'].value='No se ha suscrito contrato de obra'
				hoja_4['A19'].value='No se ha suscrito contrato de obra'


			vigencia_interventoria_cumplimiento=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=45).exists()
			if vigencia_interventoria_cumplimiento:
				vigencia_interventoria_cumplimiento=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=45)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_cumplimiento:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D20'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C20'].value=fecha_aux
				hoja_4['C20'].number_format='dd/mm/yyyy'

			vigencia_interventoria_pagos=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id__in=[110,119]).exists()
			if vigencia_interventoria_pagos:
				vigencia_interventoria_pagos=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id__in=[110,119])
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_pagos:
					#If str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D21'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C21'].value=fecha_aux
				hoja_4['C21'].number_format='dd/mm/yyyy'

			vigencia_interventoria_buenmanejo=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo__nombre__icontains='anticipo',tipo__app='poliza').exists()
			if vigencia_interventoria_buenmanejo:
				vigencia_interventoria_buenmanejo=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo__nombre__icontains='anticipo',tipo__app='poliza')
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_buenmanejo:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D22'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C22'].value=fecha_aux
				hoja_4['C22'].number_format='dd/mm/yyyy'

			vigencia_interventoria_calidad_equipos=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=44).exists()
			if vigencia_interventoria_calidad_equipos:
				vigencia_interventoria_calidad_equipos=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=44)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_calidad_equipos:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D23'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C23'].value=fecha_aux
				hoja_4['C23'].number_format='dd/mm/yyyy'

			vigencia_interventoria_responsabilidad=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=49).exists()
			if vigencia_interventoria_responsabilidad:
				vigencia_interventoria_responsabilidad=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=49)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_responsabilidad:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D24'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C24'].value=fecha_aux
				hoja_4['C24'].number_format='dd/mm/yyyy'

			vigencia_interventoria_estabilidad=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=46).exists()
			if vigencia_interventoria_estabilidad:
				vigencia_interventoria_estabilidad=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=46)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_estabilidad:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D25'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C25'].value=fecha_aux
				hoja_4['C25'].number_format='dd/mm/yyyy'
			
			
			vigencia_interventoria_responsabilidad_derivada=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=48).exists()
			if vigencia_interventoria_responsabilidad_derivada:
				vigencia_interventoria_responsabilidad_derivada=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=48)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_responsabilidad_derivada:
					#f str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D26'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C26'].value=fecha_aux
				hoja_4['C26'].number_format='dd/mm/yyyy'
				
			else:
				hoja_4['A26'].value=' - '
				hoja_4.row_dimensions[26].hidden= True

			vigencia_interventoria_seguro=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=50).exists()
			if vigencia_interventoria_seguro:
				vigencia_interventoria_seguro=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=50)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_seguro:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D27'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C27'].value=fecha_aux
				hoja_4['C27'].number_format='dd/mm/yyyy'
				
			else:
				hoja_4['A27'].value=' - '
				hoja_4.row_dimensions[27].hidden= True


			vigencia_interventoria_suspension=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=101).exists()
			if vigencia_interventoria_suspension:
				vigencia_interventoria_suspension=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=101)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_suspension:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D28'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C28'].value=fecha_aux
				hoja_4['C28'].number_format='dd/mm/yyyy'
				
			else:
				hoja_4['A28'].value=' - '
				hoja_4.row_dimensions[28].hidden= True


			vigencia_interventoria_calidad=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=43).exists()
			if vigencia_interventoria_calidad:
				vigencia_interventoria_calidad=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=43)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_calidad:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_4['D29'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C29'].value=fecha_aux
				hoja_4['C29'].number_format='dd/mm/yyyy'
				
			else:
				hoja_4['A29'].value=' - '
				hoja_4.row_dimensions[29].hidden= True

			vigencia_interventoria_prestacion_social=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=47).exists()
			if vigencia_interventoria_prestacion_social:
				vigencia_interventoria_prestacion_social=Poliza.objects.filter(contrato_id=contrato_obra.id,tipo_id=47)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_prestacion_social:
					if str(z.fecha_inicio())<str(limite_fecha):
						sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
						if str(fecha_aux)<str(z.fecha_final()):
							fecha_aux=z.fecha_final()
				hoja_4['D30'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_4['C30'].value=fecha_aux
				hoja_4['C30'].number_format='dd/mm/yyyy'
				
			else:
				hoja_4['A30'].value=' - '
				hoja_4.row_dimensions[30].hidden= True

			

		else:
			hoja_4['A10'].value='No hay contrato de obra suscrito'
			hoja_4['A12'].value='No hay contrato de obra suscrito'
			hoja_4['C15'].value='No hay contrato de obra suscrito'
			hoja_4['H15'].value='No hay contrato de obra suscrito'
			hoja_4['H16'].value='No hay contrato de obra suscrito'
			hoja_4['L15'].value='No hay contrato de obra suscrito'
			hoja_4['L16'].value='No hay contrato de obra suscrito'
			hoja_4['P16'].value='No hay contrato de obra suscrito'
			hoja_4['P15'].value='No hay contrato de obra suscrito'
			hoja_4['D18'].value='No hay contrato de obra suscrito'
			hoja_4['A19'].value='No hay contrato de obra suscrito'
		

		# limite_fecha
		qset_actividades =  (Q(presupuesto__cronograma__proyecto__id=proyecto.id))
		actividades_avance_obra = FDetallePresupuesto.objects.filter(qset_actividades).values(
			'actividad__id',
			'actividad__padre',
			'actividad__peso').annotate(total=Sum('cantidad'))

		row_act = 40
		inversion_programada = 0
		total_avance_programado = 0

		if actividades_avance_obra:
			qset_aejecutar = (Q(presupuesto__cronograma__proyecto__id=proyecto.id))
			detalles__aejecutar = FDetallePresupuesto.objects.filter(qset_actividades)



			qset_ejecutar = (Q(detallePresupuesto__presupuesto__cronograma__proyecto__id=proyecto.id)) & (Q(reporteTrabajo__fechaReporte__lte=limite_fecha))  & (Q(cantidad__gt=0))
			detalles__ejecutar = DetalleReporteTrabajo.objects.filter(qset_ejecutar)

			qset_programar = (Q(detallePresupuesto__presupuesto__cronograma__proyecto__id=proyecto.id)) & (Q(periodoProgramacion__fechaDesde__lte=limite_fecha)) & (Q(cantidad__gt=0)) 
			periodo_programar= DetallePeriodoProgramacion.objects.filter(qset_programar)


			
			for act in actividades_avance_obra:
				
				ejecutar = detalles__ejecutar.filter(detallePresupuesto__actividad__id=act['actividad__id']).values('detallePresupuesto__codigoUC','detallePresupuesto__descripcionUC','detallePresupuesto__valorGlobal').annotate(cantidad__sum=Sum('cantidad'))
				programar = periodo_programar.filter(detallePresupuesto__actividad__id=act['actividad__id']).values('detallePresupuesto__codigoUC','detallePresupuesto__descripcionUC','detallePresupuesto__valorGlobal').annotate(cantidad__sum=Sum('cantidad'))
				aejecutar = detalles__aejecutar.filter(actividad__id=act['actividad__id'])

				subtotal_programar = 0
				subtotal_aejecutar = 0
				subtotal_ejecutar = 0

				ejecutar_cantidad = 0
				programar_cantidad = 0

				for ejec in ejecutar:
					subtotal_ejecutar+= float(ejec['detallePresupuesto__valorGlobal']) * float(ejec['cantidad__sum']) if ejec['cantidad__sum'] else 0
					ejecutar_cantidad+= ejec['cantidad__sum']

				for aejec in aejecutar:
					subtotal_aejecutar+= float(aejec.valorGlobal) * float(aejec.cantidad)
				
				for progr in programar:
					subtotal_programar+= float(progr['detallePresupuesto__valorGlobal']) * float(progr['cantidad__sum']) if ejec['cantidad__sum'] else 0
					programar_cantidad+=  progr['cantidad__sum']

				


				porcentaje_aux = float(programar_cantidad)*float(act['actividad__peso']/100)/float(act['total']) if programar_cantidad and act['total'] else 0

				total_avance_programado+=porcentaje_aux
				inversion_programada+=subtotal_programar


				name_aux = CEsquemaCapitulosActividadesG.objects.get(pk=act['actividad__padre']).nombre

				hoja_4['A{0}'.format(row_act)].value=name_aux
				hoja_4['E{0}'.format(row_act)].value='Glb' if name_aux=='Retiro de materiales' or name_aux=='Pruebas y Puesta en Operación' or name_aux=='Liquidación' or name_aux=='Pruebas y puesta en operacion' or name_aux=='Liquidacion' else 'Und'
				hoja_4['F{0}'.format(row_act)].value='{0}%'.format(act['actividad__peso'])
				hoja_4['H{0}'.format(row_act)].value=act['total'] if act['total'] else 0
				hoja_4['L{0}'.format(row_act)].value=ejecutar_cantidad

				hoja_4['J{0}'.format(row_act)].value=subtotal_aejecutar
				hoja_4['M{0}'.format(row_act)].value=subtotal_ejecutar

				row_act+=1

			hoja_4['A67'].value=inversion_programada
			hoja_3['B14'].value=round(total_avance_programado,3)
		else:
			Utilizar_modulo_avanObraGrafico2(hoja_4,hoja_3,limite_fecha,proyecto,row_act)

		if interventoria:
			# interventoria=Contrato.objects.filter(mcontrato_id=contrato.id,tipo_contrato_id=9)
			# interventoria=interventoria.first()
			# import pdb; pdb.set_trace()
			# print(interventoria.id)
			ing_normalizacion=Funcionario.objects.filter(empresa_id=interventoria.contratista.id,cargo_id=26,activo=1).exists()
			if ing_normalizacion:
				ing_normalizacion=Funcionario.objects.filter(empresa_id=interventoria.contratista.id,cargo_id=26,activo=1)
				ing_normalizacion=ing_normalizacion.first()
				hoja_4['J92'].value=ing_normalizacion.persona.nombres+' '+ing_normalizacion.persona.apellidos

				aux_telefono=ing_normalizacion.persona.telefono
				caracter=0
				count_caracter=0
				for t in aux_telefono:
					if t=='-':
						caracter=count_caracter
					count_caracter=count_caracter+1

				telefono_fijo=ing_normalizacion.persona.telefono[int(caracter+1):int(caracter+8)]
				hoja_4['J94'].value='5-'+telefono_fijo
				hoja_4['J95'].value=ing_normalizacion.persona.telefono[0:11]
				hoja_4['J96'].value=ing_normalizacion.persona.correo			
			else:
				hoja_4['J92'].value='Actualmente nadie se encuentra registrado con este puesto en el sistema'				
				hoja_4['J94'].value='Actualmente nadie se encuentra registrado con este puesto en el sistema'
				hoja_4['J95'].value='Actualmente nadie se encuentra registrado con este puesto en el sistema'
				hoja_4['J96'].value='Actualmente nadie se encuentra registrado con este puesto en el sistema'


			if interventoria.contratista.direccion and interventoria.contratista.direccion!='No se encuentra registrada':
				hoja_4['J93'].value=interventoria.contratista.direccion
			else:
				hoja_4['J93'].value= interventoria.contratista.nombre
			
		else:
			hoja_4['J92'].value='No hay contrato de interventoría suscrito'			
			hoja_4['J93'].value='No hay contrato de interventoría suscrito'
			hoja_4['J94'].value='No hay contrato de interventoría suscrito'
			hoja_4['J95'].value='No hay contrato de interventoría suscrito'
			hoja_4['J96'].value='No hay contrato de interventoría suscrito'

		if contrato.contratista:
			
			cargo_funcionario=0			

			if contrato.contratista.nombre.count('AIR-E'):
				cargo_funcionario=23
			elif contrato.contratista.nombre.count('MAR'):
				cargo_funcionario=26

			normalizaicon_redes=Funcionario.objects.filter(empresa_id=contrato.contratista.id,cargo_id=cargo_funcionario,activo=1).exists()
			if normalizaicon_redes:
				normalizaicon_redes=Funcionario.objects.filter(empresa_id=contrato.contratista.id,cargo_id=cargo_funcionario,activo=1)
				normalizaicon_redes=normalizaicon_redes.first()
				hoja_4['C92'].value=normalizaicon_redes.persona.nombres+' '+normalizaicon_redes.persona.apellidos

				aux_telefono=normalizaicon_redes.persona.telefono
				caracter=0
				count_caracter=0
				for t in aux_telefono:
					if t=='-':
						caracter=count_caracter
					count_caracter=count_caracter+1

				telefono_fijo=normalizaicon_redes.persona.telefono[int(caracter+1):int(caracter+8)]
				hoja_4['C94'].value='5-'+telefono_fijo
				hoja_4['C95'].value=normalizaicon_redes.persona.telefono[0:11]
				hoja_4['C96'].value=normalizaicon_redes.persona.correo			
			else:
				hoja_4['C92'].value='Actualmente nadie se encuentra registrado con este puesto en el sistema'
				hoja_4['C95'].value='Actualmente nadie se encuentra registrado con este puesto en el sistema'
				hoja_4['C96'].value='Actualmente nadie se encuentra registrado con este puesto en el sistema'

			if contrato.contratista.direccion and contrato.contratista.direccion!='No se encuentra registrada':
				hoja_4['C93'].value=contrato.contratista.direccion
			else:
				if contrato.contratista.nombre.count('AIR-E'):
					hoja_4['C93'].value='Carrera 55 # 72 - 109 B/quilla'
				elif contrato.contratista.nombre.count('MAR'):			
					hoja_4['C93'].value='CRA 13B 26 - 78 EDIFICIO CHAMBACÚ - PISO 3 / CARTAGENA / BOLIVAR'
		else:
			hoja_4['C92'].value='No hay contrato de interventoría suscrito'			
			hoja_4['C93'].value='No hay contrato de interventoría suscrito'
			hoja_4['C94'].value='No hay contrato de interventoría suscrito'
			hoja_4['C95'].value='No hay contrato de interventoría suscrito'
			hoja_4['C96'].value='No hay contrato de interventoría suscrito'


			
		#
		##FASCP-01C-CONTRA-INTERVEN
		nombrehoja=doc.get_sheet_names()[5]
		hoja_5 = doc.get_sheet_by_name(nombrehoja)	

		if interventoria:

			# interventoria=Contrato.objects.filter(mcontrato_id=contrato.id,tipo_contrato_id=9)
			# interventoria=interventoria.first()
			fecha_inicio=datetime.strptime(str(interventoria.fecha_inicio()),'%Y-%m-%d').date()
			fecha_final=interventoria.fecha_fin()

			interventoria_meses=fecha_final-fecha_inicio
			interventoria_meses=int(interventoria_meses.days)
			interventoria_meses=interventoria_meses/30
			hoja_5['C14'].value=interventoria_meses
			hoja_5['P15'].value=fecha_final
			hoja_5['P15'].number_format='dd/mm/yyyy'

			vigencia_contrato_interventoria = VigenciaContrato.objects.filter(tipo_id__in=[18,19,102],contrato__id=interventoria.id,fecha_inicio__lte=limite_fecha)
			cantidad_prorroga_meses = 0
			cantidad_suspension_meses = 0

			if len(vigencia_contrato_interventoria)>0:
				for vigencia_aux_interventoria in vigencia_contrato_interventoria:

					if vigencia_aux_interventoria.tipo.id==102:
						aux_vigencia_obra=vigencia_aux_interventoria.fecha_fin-vigencia_aux_interventoria.fecha_inicio
						cantidad_prorroga_meses+=round(int(aux_vigencia_obra.days)/30)

					elif vigencia_aux_interventoria.tipo.id==18:
						vigencia_aux_obra_fecha_fin = vigencia_contrato_interventoria.filter(pk__gte=vigencia_aux_interventoria.id,tipo_id=19).first().values('fecha_fin')
						if vigencia_aux_obra_fecha_fin:
							aux_vigencia_obra=vigencia_aux_obra_fecha_fin-vigencia_aux_interventoria.fecha_inicio
							cantidad_suspension_meses+=round(int(aux_vigencia_obra.days)/30)
						else:
							aux_vigencia_obra=limite_fecha-vigencia_aux_interventoria.fecha_inicio
							cantidad_suspension_meses+=round(int(aux_vigencia_obra.days)/30)

			if cantidad_prorroga_meses>0:
				hoja_5['L15'].value=cantidad_prorroga_meses
			else:
				hoja_5['L15'].value='No hay prorrogas'

			if cantidad_suspension_meses>0:
				hoja_5['H15'].value=cantidad_suspension_meses
			else:
				hoja_5['H15'].value='No hay suspensiones'


			vigencia_poliza_interven = VigenciaPoliza.objects.filter(poliza__contrato_id=interventoria.id,poliza__tipo_id=45,fecha_inicio__lte=limite_fecha).exists()
			if vigencia_poliza_interven:
				vigencia_poliza_interven = VigenciaPoliza.objects.filter(poliza__contrato_id=interventoria.id,poliza__tipo_id=45,fecha_inicio__lte=limite_fecha)
				hoja_5['D17'].value=vigencia_poliza_interven[0].numero
				hoja_5['A18'].value=str(vigencia_poliza_interven[0].aseguradora)
			else:
				hoja_5['D17'].value='No se ha suscrito contrato de interventoria'
				hoja_5['A18'].value='No se ha suscrito contrato de interventoria'

			adicion_presupuestal_interventoria=VigenciaContrato.objects.filter(contrato_id=interventoria.id,tipo_id=17,fecha_inicio__lte=limite_fecha).exists()
			if adicion_presupuestal_interventoria:
				adicion_presupuestal_interventoria=VigenciaContrato.objects.filter(contrato_id=interventoria.id,tipo_id=17,fecha_inicio__lte=limite_fecha).aggregate(suma_valor=Sum('valor'))
				hoja_5['F16'].value=adicion_presupuestal_interventoria['suma_valor']
			

			
			# factura_interventoria=Factura.objects.filter(contrato_id=interventoria.id,pagada=1).exists()
			# if factura_interventoria:
			# 	factura_interventoria=Factura.objects.filter(contrato_id=interventoria.id,pagada=1).aggregate(suma_valor=Sum('valor_factura'))
			# 	hoja_5['J20'].value=factura_interventoria['suma_valor']

			vigencia_interventoria_cumplimiento=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=45).exists()
			if vigencia_interventoria_cumplimiento:
				vigencia_interventoria_cumplimiento=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=45)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_cumplimiento:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D19'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C19'].value=fecha_aux
				hoja_5['C19'].number_format='dd/mm/yyyy'

			vigencia_interventoria_calidad=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=43).exists()
			if vigencia_interventoria_calidad:
				vigencia_interventoria_calidad=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=43)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_calidad:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D20'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C20'].value=fecha_aux
				hoja_5['C20'].number_format='dd/mm/yyyy'

			vigencia_interventoria_pagos=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id__in=[110,119]).exists()
			if vigencia_interventoria_pagos:
				vigencia_interventoria_pagos=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id__in=[110,119])
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_pagos:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D22'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C22'].value=fecha_aux
				hoja_5['C22'].number_format='dd/mm/yyyy'

			vigencia_interventoria_responsabilidad=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=49).exists()
			if vigencia_interventoria_responsabilidad:
				vigencia_interventoria_responsabilidad=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=49)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_responsabilidad:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D23'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C23'].value=fecha_aux
				hoja_5['C23'].number_format='dd/mm/yyyy'
				
			vigencia_interventoria_buenmanejo=Poliza.objects.filter(contrato_id=interventoria.id,tipo__nombre__icontains='anticipo',tipo__app='poliza').exists()
			if vigencia_interventoria_buenmanejo:
				vigencia_interventoria_buenmanejo=Poliza.objects.filter(contrato_id=interventoria.id,tipo__nombre__icontains='anticipo',tipo__app='poliza')
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_buenmanejo:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D21'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C21'].value=fecha_aux
				hoja_5['C21'].number_format='dd/mm/yyyy'





				#import pdb; pdb.set_trace()
			menor_fila=1000
			vigencia_interventoria_estabilidad=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=46).exists()
			if vigencia_interventoria_estabilidad:
				vigencia_interventoria_estabilidad=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=46)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_estabilidad:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D24'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C24'].value=fecha_aux				
				hoja_5['C24'].number_format='dd/mm/yyyy'
				if menor_fila>24:
					menor_fila=24
			else:
				hoja_5['A24'].value=' - '
				hoja_5.row_dimensions[24].hidden= True
				hoja_5['J24'].value=' '


			vigencia_interventoria_responsabilidad_derivada=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=48).exists()
			if vigencia_interventoria_responsabilidad_derivada:
				vigencia_interventoria_responsabilidad_derivada=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=48)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_responsabilidad_derivada:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D25'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C25'].value=fecha_aux
				hoja_5['C25'].number_format='dd/mm/yyyy'
				if menor_fila>25:
					menor_fila=25
			else:
				hoja_5['A25'].value=' - '
				hoja_5.row_dimensions[25].hidden= True
				hoja_5['J25'].value=' '



			vigencia_interventoria_seguro=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=50).exists()
			if vigencia_interventoria_seguro:
				vigencia_interventoria_seguro=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=50)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_seguro:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D26'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C26'].value=fecha_aux
				hoja_5['C26'].number_format='dd/mm/yyyy'
				if menor_fila>26:
					menor_fila=26
			else:
				hoja_5['A26'].value=' - '
				hoja_5.row_dimensions[26].hidden= True
				hoja_5['J26'].value=' '


			vigencia_interventoria_suspension=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=101).exists()
			if vigencia_interventoria_suspension:
				vigencia_interventoria_suspension=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=101)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_suspension:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D27'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C27'].value=fecha_aux
				hoja_5['C27'].number_format='dd/mm/yyyy'
				if menor_fila>27:
					menor_fila=27
			else:
				hoja_5['A27'].value=' - '
				hoja_5.row_dimensions[27].hidden= True
				hoja_5['J27'].value=' '

			vigencia_interventoria_calidad_equipos=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=44).exists()
			if vigencia_interventoria_calidad_equipos:
				vigencia_interventoria_calidad_equipos=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=44)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_calidad_equipos:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D28'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C28'].value=fecha_aux
				hoja_5['C28'].number_format='dd/mm/yyyy'
				if menor_fila>28:
					menor_fila=28
			else:
				hoja_5['A28'].value=' - '
				hoja_5.row_dimensions[28].hidden= True
				hoja_5['J28'].value=' '


			vigencia_interventoria_prestacion_social=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=47).exists()
			if vigencia_interventoria_prestacion_social:
				vigencia_interventoria_prestacion_social=Poliza.objects.filter(contrato_id=interventoria.id,tipo_id=47)
				sumatoria_vigencia_interventoria=0
				fecha_aux='0'
				for z in vigencia_interventoria_prestacion_social:
					#if str(z.fecha_inicio())<str(limite_fecha):
					sumatoria_vigencia_interventoria=sumatoria_vigencia_interventoria+z.valor()
					if str(fecha_aux)<str(z.fecha_final()):
						fecha_aux=z.fecha_final()
				hoja_5['D29'].value=sumatoria_vigencia_interventoria
				if fecha_aux=='0':
					fecha_aux='-'
				hoja_5['C29'].value=fecha_aux
				hoja_5['C29'].number_format='dd/mm/yyyy'
				if menor_fila>29:
					menor_fila=29
			else:
				hoja_5['A29'].value=' - '
				hoja_5.row_dimensions[29].hidden= True
				hoja_5['J29'].value=' '

			hoja_5['F'+str(menor_fila)].value='Saldo por amortizar anticipo'


			if fondo=='PRONE':
				hoja_5['A11'].value=''
			elif fondo=='FAER':
				hoja_5['A11'].value='EL CONTRATISTA se obliga para con EL CONTRATANTE a prestar los servicios de INTERVENTORÍA ADMINISTRATIVA, TÉCNICA, JURÍDICA, AMBIENTAL, SOCIAL, FINANCIERA Y CONTABLE DEL PROYECTO FINANCIADO CON RECURSOS DEL FONDO DE APOYO FINANCIERO PARA LA ENERGIZACIÓN DE LAS ZONAS RURALES INTERCONECTADAS FAER EN '+empresa_contratista+' DE LAS ZONAS RURALES INTERCONECTADAS FAER EN '+empresa_contratista+' especificado en el CONTRATO '+contrato.nombre+' PROYECTO DE ELECTRIFICACIÓN RURAL '+proyecto.nombre
				# hoja_5['J22'].value='La forma de pago establecida en el contrato indica que se pagará al contratista mediante facturas'
				# hoja_5['J23'].value='La forma de pago establecida en el contrato indica que se pagará al contratista mediante facturas'
				celda_aux=24
				while celda_aux<30:
					if menor_fila==celda_aux:
						hoja_5['J'+str(celda_aux)].value='La forma de pago establecida en el contrato indica que se pagará al contratista mediante facturas'
					else:
						hoja_5['J'+str(celda_aux)].value=' '
					celda_aux=celda_aux+1

				hoja_5.row_dimensions[22].height = 40
				hoja_5.row_dimensions[23].height = 40
				hoja_5.row_dimensions[menor_fila].height = 40

			vigencia_interventoria=VigenciaContrato.objects.filter(contrato__id=interventoria.id,  tipo_id=16,fecha_inicio__lte=limite_fecha).exists()

			if vigencia_interventoria:
				vigencia_interventoria=VigenciaContrato.objects.filter(contrato__id=interventoria.id,  tipo_id=16,fecha_inicio__lte=limite_fecha)
				aux_fecha=interventoria.fecha_firma
				if aux_fecha and aux_fecha<limite_fecha:
					hoja_5['H14'].value=aux_fecha
				else:
					hoja_5['H14'].value=interventoria.fecha_inicio()

				aux_fecha=interventoria.fecha_acta_inicio
				if aux_fecha  and aux_fecha<limite_fecha:
					hoja_5['L14'].value=aux_fecha
				else:
					hoja_5['L14'].value='No se ha suscrito la fecha de inicio del contrato de interventoria'

				hoja_5['H14'].number_format='dd/mm/yyyy'
				hoja_5['L14'].number_format='dd/mm/yyyy'

				hoja_5['C16'].value=vigencia_interventoria[0].valor

				hoja_5['F9'].value=vigencia_interventoria[0].contrato.contratista.nombre

				vigencia_finalizacion=vigencia_interventoria.last()
				if vigencia_finalizacion.tipo.id==21:
					hoja_5['P14'].value=vigencia_finalizacion.fecha_inicio
					hoja_5['P14'].number_format='dd/mm/yyyy'
				else:
					hoja_5['P14'].value='El contrato no ha finalizado'


				if contrato_obra:			
					vigencia_contrato_obra=VigenciaContrato.objects.filter(contrato_id=contrato_obra.id, contrato__mcontrato_id=contrato.id,contrato__tipo_contrato_id=8, tipo_id=16,fecha_inicio__lte=limite_fecha)
					fecha_inicio_c_obra=vigencia_contrato_obra[0].fecha_inicio
					fecha_inicio_inter=vigencia_interventoria[0].fecha_inicio

					if str(fecha_inicio_inter)>str(fecha_inicio_c_obra):
						hoja_5['P16'].value='SI'
					else:
						hoja_5['P16'].value='NO'
				else:
					hoja_5['P16'].value='NO'
				

			else:
				hoja_5['C16'].value='No hay contrato de interventoría suscrito'
				hoja_5['P16'].value='No hay contrato de interventoría suscrito'		

			vigencia_interventoria_liquidacion=VigenciaContrato.objects.filter(contrato__id=interventoria.id, tipo_id=21,fecha_inicio__lte=limite_fecha).exists()

			# if not vigencia_interventoria_liquidacion:
			# 	hoja_5['J21'].value='Monto será validado ante liquidación del contrato'
			# 	hoja_5.row_dimensions[21].height = 30
				

		else:
			mensaje_no_interventoria='No hay contrato de interventoría suscrito'
			hoja_5['F9'].value=mensaje_no_interventoria
			hoja_5['A11'].value=mensaje_no_interventoria

			hoja_5['D17'].value=mensaje_no_interventoria

			hoja_5['A50'].value=mensaje_no_interventoria

			hoja_5['P47'].value=mensaje_no_interventoria
			hoja_5['P48'].value=mensaje_no_interventoria

			hoja_5['B26'].value=mensaje_no_interventoria
			hoja_5['J26'].value=mensaje_no_interventoria

			hoja_5['C14'].value=mensaje_no_interventoria
			hoja_5['H14'].value=mensaje_no_interventoria
			hoja_5['L14'].value=mensaje_no_interventoria
			hoja_5['P14'].value=mensaje_no_interventoria
			hoja_5['P15'].value=mensaje_no_interventoria
			hoja_5['L15'].value=mensaje_no_interventoria
			hoja_5['H15'].value=mensaje_no_interventoria
			hoja_5['C16'].value=mensaje_no_interventoria
			hoja_5['F16'].value=mensaje_no_interventoria
			hoja_5['P16'].value=mensaje_no_interventoria
			
			hoja_5['c19'].value='-'
			hoja_5['C20'].value='-'
			hoja_5['C21'].value='-'
			hoja_5['C22'].value='-'
			hoja_5['C23'].value='-'
			hoja_5['C24'].value='-'

			hoja_5['D19'].value=mensaje_no_interventoria
			hoja_5['D20'].value=mensaje_no_interventoria
			hoja_5['D21'].value=mensaje_no_interventoria
			hoja_5['D22'].value=mensaje_no_interventoria
			hoja_5['D23'].value=mensaje_no_interventoria
			hoja_5['D24'].value=mensaje_no_interventoria
			hoja_5['D22'].value=mensaje_no_interventoria

			# hoja_5['J19'].value=mensaje_no_interventoria
			# hoja_5['J20'].value=mensaje_no_interventoria
			# hoja_5['J21'].value=mensaje_no_interventoria
			# hoja_5['J22'].value=mensaje_no_interventoria
			# hoja_5['J23'].value=mensaje_no_interventoria
			hoja_5['J24'].value=mensaje_no_interventoria
			# hoja_5['J30'].value=mensaje_no_interventoria
			hoja_5['J16'].value=mensaje_no_interventoria
			hoja_5['F26'].value=mensaje_no_interventoria
			hoja_5['O7'].value=mensaje_no_interventoria
			hoja_5['K7'].value=mensaje_no_interventoria

			hoja_5['A24'].value='Otra? Cual _____________________'
			hoja_5['A25'].value='-'
			hoja_5['A26'].value='-'
			hoja_5['A27'].value='-'
			hoja_5['A28'].value='-'
			hoja_5['A29'].value='-'

			hoja_5.row_dimensions[25].hidden=True
			hoja_5.row_dimensions[26].hidden=True
			hoja_5.row_dimensions[27].hidden=True
			hoja_5.row_dimensions[28].hidden=True
			hoja_5.row_dimensions[29].hidden=True


		nombrehoja=doc.get_sheet_names()[6]
		hoja_6 = doc.get_sheet_by_name(nombrehoja)

		nombrehoja=doc.get_sheet_names()[7]
		hoja_7 = doc.get_sheet_by_name(nombrehoja)

		nombrehoja=doc.get_sheet_names()[8]
		hoja_8 = doc.get_sheet_by_name(nombrehoja)

		nombrehoja=doc.get_sheet_names()[9]
		hoja_9 = doc.get_sheet_by_name(nombrehoja)



		capitulos_cronograma = CcActividadContrato.objects.filter(contrato_id=contrato.id).values('actividad__capitulo__id','actividad__capitulo__nombre','actividad__capitulo__orden').order_by('actividad__capitulo__orden').distinct()
		detalles_cronograma_contrato= CcActividadContrato.objects.filter(contrato_id=contrato.id)
		soportes_cronograma = CcActividadContratoSoporte.objects.filter(actividadcontrato__contrato_id=contrato.id)
		medium_border=Border(left=Side(style='medium'),
			right=Side(style='medium'),
			top=Side(style='medium'),
			bottom=Side(style='medium'))

		row = 10
		for capit in capitulos_cronograma:
			detalles_capitulo_actual = detalles_cronograma_contrato.filter(actividad__capitulo__id=capit['actividad__capitulo__id']).values('id','actividad__descripcion','actividad__orden','inicioprogramado','finprogramado','inicioejecutado','finejecutado','observaciones').order_by('actividad__orden')
			
			hoja_9['A{0}'.format(row)].value='Etapa '+str(capit['actividad__capitulo__orden'])+' - '+capit['actividad__capitulo__nombre']
			hoja_9.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
			
			set_border_cellxcell(hoja_9, 'A{0}:F{0}'.format(row),medium_border)
			set_border_color(hoja_9, 'G{0}:L{0}'.format(row),'bottom','00000000','medium')
			# hoja_9['G{0}'.format(row)].value=''
			row+=1
			for detalle_contrato in detalles_capitulo_actual:
				soporte_act_actual = soportes_cronograma.filter(actividadcontrato__id=detalle_contrato['id'])

				if soporte_act_actual:
					soporte_act_actual = soporte_act_actual.last()
					hoja_9['F{0}'.format(row)].value=soporte_act_actual.nombre

				

				hoja_9['A{0}'.format(row)].value=str(detalle_contrato['actividad__orden'])
				hoja_9['B{0}'.format(row)].value=detalle_contrato['actividad__descripcion']
				hoja_9.cell(column=2, row=row).alignment = Alignment(wrapText=True)
				if detalle_contrato['inicioprogramado']:
					hoja_9['D{0}'.format(row)].value=detalle_contrato['inicioprogramado']
					hoja_9['D{0}'.format(row)].number_format='dd/mm/yyyy'
					hoja_9['C{0}'.format(row)].value='=+E{0}-D{0}'.format(row)
				else:
					hoja_9['D{0}'.format(row)].value='N/A'
					hoja_9['C{0}'.format(row)].value='N/A'

				if detalle_contrato['finprogramado']:
					hoja_9['E{0}'.format(row)].value=detalle_contrato['finprogramado']
					hoja_9['E{0}'.format(row)].number_format='dd/mm/yyyy'
				else:
					hoja_9['E{0}'.format(row)].value='N/A'

				if detalle_contrato['inicioejecutado']:
					hoja_9['H{0}'.format(row)].value=detalle_contrato['inicioejecutado'] if detalle_contrato['inicioejecutado']<=limite_fecha else ''
					hoja_9['H{0}'.format(row)].number_format='dd/mm/yyyy'				
				else:
					hoja_9['H{0}'.format(row)].value='N/A'

				if detalle_contrato['finejecutado']:
					hoja_9['I{0}'.format(row)].value=detalle_contrato['finejecutado'] if detalle_contrato['finejecutado']<=limite_fecha else ''
					hoja_9['I{0}'.format(row)].number_format='dd/mm/yyyy'
				else:
					hoja_9['I{0}'.format(row)].value='N/A'

				hoja_9['L{0}'.format(row)].value=detalle_contrato['observaciones']
				hoja_9.cell(column=12, row=row).alignment = Alignment(wrapText=True)


				

				hoja_9['G{0}'.format(row)].value='=IF(AND(I{0}="",I{0}<=$M$5),IF(H{0}> $B$6,IF(I{0}> $B$6,(IF(E{0}< $B$6,"Presenta Mora",IF(D{0}>= $B$6,"por ejecutar","sin iniciar"))),IF(I{0}="",(IF(E{0}< $B$6,"Presenta Mora",IF(D{0}>= $B$6,"por ejecutar","sin iniciar"))),"Ejecutado")),IF(H{0}="",IF(I{0}> $B$6,(IF(E{0}< $B$6,"Presenta Mora",IF(D{0}>= $B$6,"por ejecutar","sin iniciar"))),IF(I{0}="",(IF(E{0}< $B$6,"Presenta Mora",IF(D{0}>= $B$6,"por ejecutar","sin iniciar"))),"Ejecutado")),"En Ejecución")),"Ejecutado")'.format(row)
				hoja_9.cell(column=7, row=row).alignment = Alignment(wrapText=True)
				hoja_9['J{0}'.format(row)].value='=IFERROR(IF(H{0}>= $B$6,IF(DAYS360(D{0}, $B$6)>0,DAYS360(D{0}, $B$6),"A tiempo"),IF(AND(I{0}="",H{0}=""),IF(DAYS360(D{0}, $B$6)>0,DAYS360(D{0}, $B$6),"A tiempo"),DAYS360(D{0},H{0}))),"N/A")'.format(row)
				hoja_9['K{0}'.format(row)].value='=IFERROR(IF(I{0}="",IF(I{0}="",IF( $B$6>E{0},DAYS360(E{0}, $B$6),IF(E{0}> $B$6,"No ha finalizado","a tiempo")),DAYS360(E{0},I{0})),DAYS360(E{0},I{0})),"N/A")'.format(row)

				row+=1

		# set_border_cellxcell(hoja_9, 'A10:L{0}'.format(row-1),medium_border)
		set_border_color(hoja_9, 'L10:L{0}'.format(row-1),'right','00000000','medium')
		set_border_color(hoja_9, 'G10:G{0}'.format(row-1),'left','00000000','medium')
		set_border_color(hoja_9, 'G10:G{0}'.format(row-1),'right','00000000','medium')
		set_border_color(hoja_9, 'A{0}:L{0}'.format(row-1),'bottom','00000000','medium')

		hoja_1['B11'].value="='5,-CRONOGRAMA'!E{0}".format(row-1) if row>10 else 'N/A'
		# hoja_3['K14'].value="=('5,-CRONOGRAMA'!E{0}-'5,-CRONOGRAMA'!D11)/30".format(row-1)

		nombrehoja=doc.get_sheet_names()[10]
		hoja_10 = doc.get_sheet_by_name(nombrehoja)

		nombrehoja=doc.get_sheet_names()[11]
		hoja_11 = doc.get_sheet_by_name(nombrehoja)

		if not interventoria:
			hoja_8['A2'].value='No se ha suscrito contrato de interventoría'
		#---Fuentes

		
		#Bordes "medium"

		medium_border=Border(left=Side(style='medium'),
			right=Side(style='medium'),
			top=Side(style='medium'),
			bottom=Side(style='medium'))

		medium_top_border=Border(left=Side(style='thin'),
			right=Side(style='thin'),
			top=Side(style='medium'),
			bottom=Side(style='thin'))

		medium_right_border=Border(left=Side(style='thin'),
			right=Side(style='medium'),
			top=Side(style='thin'),
			bottom=Side(style='thin'))

		medium_bot_border=Border(left=Side(style='thin'),
			right=Side(style='thin'),
			top=Side(style='thin'),
			bottom=Side(style='medium'))

		#Bordes "thin"

		thin_border=Border(left=Side(style='thin'),
			right=Side(style='thin'),
			top=Side(style='thin'),
			bottom=Side(style='thin'))

		#Relinear bordes

		#Bordes internos y externos

		set_border_cellxcell(hoja_0, 'D4:E87',thin_border)

		set_border_cellxcell(hoja_1, 'B1:D53',thin_border)
		set_border_cellxcell(hoja_1, 'A68:F71',thin_border)
		set_border_cellxcell(hoja_1, 'A74:F78',thin_border)

		set_border_cellxcell(hoja_2, 'A48:B48',thin_border)
		set_border_cellxcell(hoja_2, 'A52:L53',thin_border)	
		set_border_cellxcell(hoja_2, 'A67:B67',thin_border)	

		set_border_cellxcell(hoja_3, 'A7:P25',thin_border)

		set_border_cellxcell(hoja_4, 'A7:P17',thin_border)
		set_border_cellxcell(hoja_4, 'A18:P31',thin_border)
		set_border_cellxcell(hoja_4, 'A32:P33',thin_border)
		set_border_cellxcell(hoja_4, 'A39:P64',thin_border)
		set_border_cellxcell(hoja_4, 'A61:P75',thin_border)

		set_border_cellxcell(hoja_4, 'A91:P96',thin_border)
		set_border_cellxcell(hoja_4, 'A101:P157',thin_border)

		# set_border_cellxcell(hoja_4, 'B178:G179',thin_border)
		set_border_cellxcell(hoja_4, 'B186:G187',thin_border)
		set_border_cellxcell(hoja_4, 'B209:G210',thin_border)
		set_border_cellxcell(hoja_4, 'I186:O187',thin_border)
		set_border_cellxcell(hoja_4, 'I209:O210',thin_border)
		set_border_cellxcell(hoja_4, 'B244:G245',thin_border)
		set_border_cellxcell(hoja_4, 'B267:G268',thin_border)
		set_border_cellxcell(hoja_4, 'I244:O245',thin_border)
		set_border_cellxcell(hoja_4, 'I267:O268',thin_border)
		set_border_cellxcell(hoja_4, 'B304:G305',thin_border)
		set_border_cellxcell(hoja_4, 'B327:G328',thin_border)
		set_border_cellxcell(hoja_4, 'I304:O305',thin_border)
		set_border_cellxcell(hoja_4, 'I327:O328',thin_border)

		set_border_cellxcell(hoja_5, 'A7:P38',thin_border)
		set_border_cellxcell(hoja_5, 'A47:P51',thin_border)

		set_border_cellxcell(hoja_6, 'A7:P68',thin_border)
		set_border_cellxcell(hoja_6, 'A77:P117',thin_border)

		set_border_cellxcell(hoja_7, 'A2:D2',thin_border)

		set_border_cellxcell(hoja_8, 'A1:D2',thin_border)

		set_border_cellxcell(hoja_9, 'A1:G3',thin_border)
		set_border_cellxcell(hoja_9, 'A7:L9',thin_border)

		set_border_cellxcell(hoja_10, 'H1:J2',thin_border)

		set_border_cellxcell(hoja_11, 'B1:G2',thin_border)

		#Border outside
		set_border_cellxcell(hoja_0, 'B2:E2',medium_border)
		set_border_cellxcell(hoja_0, 'E4:E87',medium_right_border)

		set_border_cellxcell(hoja_1, 'D1:D64',medium_right_border)

		
		set_border_cellxcell(hoja_2, 'A48:B48',medium_bot_border)
		set_border_cellxcell(hoja_2, 'B95:C95',medium_top_border)

		set_border_cellxcell(hoja_3, 'A1:P5',medium_border)

		set_border_cellxcell(hoja_4, 'A2:P5',medium_border)
		set_border_cellxcell(hoja_4, 'A34:P37',medium_border)
		set_border_cellxcell(hoja_4, 'A97:P100',medium_border)
		set_border_cellxcell(hoja_4, 'A158:P162',medium_border)
		set_border_cellxcell(hoja_4, 'A217:P220',medium_border)
		set_border_cellxcell(hoja_4, 'A277:P280',medium_border)

		set_border_cellxcell(hoja_5, 'A2:P5',medium_border)
		set_border_cellxcell(hoja_5, 'A52:O53',medium_border)

		set_border_cellxcell(hoja_6, 'A2:P5',medium_border)
		set_border_cellxcell(hoja_6, 'A27:P30',medium_border)

		# set_border_cellxcell(hoja_9, 'A9:E9',medium_border)
		# set_border_cellxcell(hoja_9, 'A24:F24',medium_border)
		# set_border_cellxcell(hoja_9, 'A29:F29',medium_border)
		# set_border_cellxcell(hoja_9, 'A39:F39',medium_border)

		#Border blue
		set_border_color(hoja_1, 'F65:F79','right','000000FF','thick')
		set_border_color(hoja_2, 'L52:L53','right','000000FF','thick')
		set_border_color(hoja_3, 'P1:P28','right','000000FF','thick')
		set_border_color(hoja_4, 'P1:P272','right','000000FF','thick')
		set_border_color(hoja_5, 'P1:P63','right','000000FF','thick')
		set_border_color(hoja_6, 'P1:P117','right','000000FF','thick')

		set_border_color(hoja_3, 'A28:P28','bottom','000000FF','thick')
		set_border_color(hoja_5, 'A63:P63','bottom','000000FF','thick')
		set_border_color(hoja_6, 'A117:P117','bottom','000000FF','thick')

		set_border_color(hoja_4, 'A150:P150','top','000000FF','thick')
		set_border_color(hoja_4, 'A97:P97','top','000000FF','thick')

		#Border black
		set_border_color(hoja_0, 'D4:E4','top','00000000','medium')
		set_border_color(hoja_0, 'D88:E88','top','00000000','medium')

		set_border_color(hoja_1, 'A65:D65','top','00000000','medium')
		set_border_color(hoja_1, 'A67:F67','top','00000000','medium')
		set_border_color(hoja_1, 'A72:C72','top','00000000','medium')
		set_border_color(hoja_1, 'F72:F72','top','00000000','medium')
		set_border_color(hoja_1, 'A74:F74','top','00000000','medium')
		set_border_color(hoja_1, 'A79:F79','top','00000000','medium')
		set_border_color(hoja_1, 'A55:D55','top','00000000','thin')

		set_border_color(hoja_2, 'A24:D24','top','00000000','medium')
		set_border_color(hoja_2, 'A26:D26','top','00000000','medium')
		set_border_color(hoja_2, 'C95:D95','top','00000000','thin')
		set_border_color(hoja_2, 'A95:D95','top','00000000','medium')
		set_border_color(hoja_2, 'A92:D92','bottom','00000000','medium')
		set_border_color(hoja_2, 'D95:D95','bottom','00000000','thin')
		set_border_color(hoja_2, 'A92:D92','bottom','00000000','thin')
		set_border_color(hoja_2, 'B1:B2','right','00000000','medium')
		set_border_color(hoja_2, 'B26:B27','right','00000000','medium')
		set_border_color(hoja_2, 'D95:D96','right','00000000','medium')

		set_border_color(hoja_3, 'A26:P26','bottom','00000000','thin')

		set_border_color(hoja_4, 'A83:P83','bottom','00000000','medium')
		set_border_color(hoja_4, 'A67:P67','top','00000000','medium')
		set_border_color(hoja_4, 'A68:P68','top','00000000','medium')
		set_border_color(hoja_4, 'A90:P90','top','00000000','medium')
		set_border_color(hoja_4, 'A91:P91','top','00000000','medium')

		set_border_color(hoja_5, 'A45:P45','top','00000000','medium')
		set_border_color(hoja_5, 'A55:P55','top','00000000','medium')
		set_border_color(hoja_5, 'A46:P46','top','00000000','medium')

		set_border_color(hoja_6, 'A75:P75','top','00000000','medium')
		set_border_color(hoja_6, 'A76:P76','top','00000000','medium')

		set_border_color(hoja_7, 'D1:D2','right','00000000','medium')

		set_border_color(hoja_8, 'D1:D3','right','00000000','medium')

		set_border_color(hoja_9, 'E5:E8','right','00000000','medium')
		set_border_color(hoja_9, 'A5:E5','top','00000000','medium')
		set_border_color(hoja_9, 'A6:E6','top','00000000','medium')
		set_border_color(hoja_9, 'A7:L7','top','00000000','medium')
		set_border_color(hoja_9, 'A10:L10','top','00000000','medium')

		set_border_color(hoja_9, 'L7:L9','right','00000000','medium')

		set_border_color(hoja_9, 'B1:B3','right','00000000','medium')
		set_border_color(hoja_9, 'G1:G2','right','00000000','medium')
		set_border_color(hoja_9, 'A3:G3','top','00000000','medium')

		
		doc.save(response)
		return response


	except Exception as e:
		functions.toLog(e,'InformeMME')
		return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def Rellenar_actividades(hoja,celda,proyecto,tipo):
	actividad=Proyecto_Actividad_contrato.objects.filter(proyecto_id=proyecto.id,actividad_id=int(tipo)).exists()
	if actividad:
		actividad=Proyecto_Actividad_contrato.objects.filter(proyecto_id=proyecto.id,actividad_id=int(tipo))
		hoja[str(celda)]=actividad[0].valor
	else:
		hoja[str(celda)]='No se ha suscrito esta información'

def set_border_color(ws, cell_range,border,color,type_border):
	#border.left.style
	
	rows = ws[cell_range]
	for row in rows:
		for c in row:
			#import pdb; pdb.set_trace()
			if border=='top':
				c.border = Border(left=Side(style=c.border.left.style, color=c.border.left.color),
						right=Side(style=c.border.right.style, color=c.border.right.color),
						top=Side(style=type_border, color=color),
						bottom=Side(style=c.border.bottom.style, color=c.border.bottom.color))
			if border=='bottom':
				c.border = Border(left=Side(style=c.border.left.style, color=c.border.left.color),
						right=Side(style=c.border.right.style, color=c.border.right.color),
						top=Side(style=c.border.top.style, color=c.border.top.color),
						bottom=Side(style=type_border, color=color))
			if border=='right':
				c.border = Border(left=Side(style=c.border.left.style, color=c.border.left.color),
						right=Side(style=type_border, color=color),
						top=Side(style=c.border.top.style, color=c.border.top.color),
						bottom=Side(style=c.border.bottom.style, color=c.border.bottom.color))
			if border=='left':
				c.border = Border(left=Side(style=type_border, color=color),
						right=Side(style=c.border.right.style, color=c.border.right.color),
						top=Side(style=c.border.top.style, color=c.border.top.color),
						bottom=Side(style=c.border.bottom.style, color=c.border.bottom.color))

def set_border_cellxcell(ws, cell_range, style_border):
	
	rows = ws[cell_range]
	for row in rows:
		for c in row:
			#import pdb; pdb.set_trace()
			c.border = style_border

def Utilizar_modulo_avanObraGrafico2(hoja_4,hoja_3,limite_fecha,proyecto,row_act):
	try:
		from avanceObraGrafico2.models import FDetallePresupuesto, KDetalleReporteTrabajo, CEsquemaCapitulosActividadesG, JCantidadesNodo

		qset_actividades =  (Q(presupuesto__cronograma__proyecto__id=proyecto.id))
		actividades_avance_obra = FDetallePresupuesto.objects.filter(qset_actividades).values(
			'actividad__id',
			'actividad__padre',
			'actividad__peso').annotate(total=Sum('cantidad'))

		inversion_programada = 0
		total_avance_programado = 0

		qset_aejecutar = (Q(presupuesto__cronograma__proyecto__id=proyecto.id))
		detalles__aejecutar = FDetallePresupuesto.objects.filter(qset_actividades)

		qset_ejecutar = (Q(detallepresupuesto__presupuesto__cronograma__proyecto__id=proyecto.id)) & (Q(reporte_trabajo__fechaTrabajo__lte=limite_fecha)) & (Q(cantidadEjecutada__gt=0))
		detalles__ejecutar = KDetalleReporteTrabajo.objects.filter(qset_ejecutar)

		qset_programar = (Q(detallepresupuesto__presupuesto__cronograma__proyecto__id=proyecto.id)) & (Q(cantidad__gt=0)) 
		periodo_programar= JCantidadesNodo.objects.filter(qset_programar)

		for act in actividades_avance_obra:
			ejecutar = detalles__ejecutar.filter(detallepresupuesto__actividad__id=act['actividad__id']).values('detallepresupuesto__codigoUC','detallepresupuesto__descripcionUC','detallepresupuesto__valorGlobal').annotate(cantidad__sum=Sum('cantidadEjecutada'))
			programar = periodo_programar.filter(detallepresupuesto__actividad__id=act['actividad__id']).values('detallepresupuesto__codigoUC','detallepresupuesto__descripcionUC','detallepresupuesto__valorGlobal').annotate(cantidad__sum=Sum('cantidad'))
			aejecutar = detalles__aejecutar.filter(actividad__id=act['actividad__id'])


			subtotal_programar = 0
			subtotal_aejecutar = 0
			subtotal_ejecutar = 0

			ejecutar_cantidad = 0
			programar_cantidad = 0


			for ejec in ejecutar:
				subtotal_ejecutar+= float(ejec['detallepresupuesto__valorGlobal']) * float(ejec['cantidad__sum']) if ejec['cantidad__sum'] else 0
				ejecutar_cantidad+= ejec['cantidad__sum']

			for aejec in aejecutar:
				subtotal_aejecutar+= float(aejec.valorGlobal) * float(aejec.cantidad)

			for progr in programar:
				subtotal_programar+= float(progr['detallepresupuesto__valorGlobal']) * float(progr['cantidad__sum']) if progr['cantidad__sum'] else 0
				programar_cantidad+=  progr['cantidad__sum']


			porcentaje_aux = float(programar_cantidad)*float(act['actividad__peso']/100)/float(act['total']) if programar_cantidad and act['total'] else 0
			# print(total_avance_programado)
			total_avance_programado+=porcentaje_aux
			inversion_programada+=subtotal_programar

			name_aux = CEsquemaCapitulosActividadesG.objects.get(pk=act['actividad__padre']).nombre

			hoja_4['A{0}'.format(row_act)].value=name_aux
			hoja_4['E{0}'.format(row_act)].value='Glb' if name_aux=='Retiro de materiales' or name_aux=='Pruebas y Puesta en Operación' or name_aux=='Liquidación' or name_aux=='Pruebas y puesta en operacion' or name_aux=='Liquidacion' else 'Und'
			hoja_4['F{0}'.format(row_act)].value='{0}%'.format(act['actividad__peso'])
			hoja_4['H{0}'.format(row_act)].value=act['total'] if act['total'] else 0
			hoja_4['L{0}'.format(row_act)].value=ejecutar_cantidad

			hoja_4['J{0}'.format(row_act)].value=subtotal_aejecutar
			hoja_4['M{0}'.format(row_act)].value=subtotal_ejecutar

			row_act+=1

		hoja_4['A67'].value=inversion_programada
		hoja_3['B14'].value=round(total_avance_programado,3)
	except Exception as e:
		functions.toLog(e,'InformeMME, Utilizar_modulo_avanObraGrafico2')

def informeFotosProyecto(request):
	if request.method == 'GET':
		try:
			document = Document()
			tipo_contrato = tipoC()

			styles = document.styles
			font = styles['Normal'].font
			font.name = 'Arial'

			contrato = request.GET['contrato'] if 'contrato' in request.GET else None;
			foto = request.GET['foto'] if 'foto' in request.GET else None;#es el periodo de la foto
			fechaDesde = request.GET['fechaDesde'] if 'fechaDesde' in request.GET else None;
			fechaHasta = request.GET['fechaHasta'] if 'fechaHasta' in request.GET else None;

			qset=(Q(asociado_reporte= 1))#CONSULTA TODAS LAS FOTOS ASOCIADAS A UN REPORTE
			qset = qset & ( Q(proyecto__contrato__tipo_contrato__id = tipo_contrato.contratoProyecto) )
			if contrato and contrato is not None:
				qset = qset & ( Q(proyecto__mcontrato_id = contrato ) )
				# FILTRA POR EL MACRO CONTRATO QUE TENGA EL PROYECTO
			if foto and foto is not None:
				qset = qset & ( Q(tipo_id = foto ) )
				# FILTRA POR EL TIPO DE FOTOS DEL PROYECTO
			if (fechaDesde and (fechaHasta is not None) ):
				qset = qset & ( Q(fecha__gte = fechaDesde ) )
				# INICIA EL FILTRO POR LA FECHA DE LA FOTO
			if ( (fechaDesde is not None) and fechaHasta ):
				qset = qset & ( Q(fecha__lte =fechaHasta)  )
				# FINALIZA EL FILTRO POR LA FECHA DE LA FOTO
			if (fechaDesde and fechaHasta):	
				qset = qset & ( Q(fecha__range = (fechaDesde , fechaHasta) ) )
				# BUSCA ENTRE LAS FECHAS SELECCIONADA 

			convenio = Contrato.objects.get(pk = contrato)
			proyectos =CFotosProyecto.objects.filter(qset).values('proyecto__id' ,'proyecto__nombre','proyecto__municipio__departamento__nombre' ,'proyecto__municipio__nombre').order_by('proyecto__nombre').distinct()

			# print proyectos.query
			if proyectos:

				document.add_paragraph('')
				document.add_paragraph('')
				document.add_paragraph('')
				document.add_paragraph('')
				p=document.add_paragraph('')
				parrafo = p.add_run('REGISTRO FOTOGRAFICO')
				parrafo.bold = True
				parrafo.font.size = Pt(20)
				p.alignment = WD_ALIGN_PARAGRAPH.CENTER

				document.add_paragraph('')
				document.add_paragraph('')
				document.add_paragraph('')
				document.add_paragraph('')

				p=document.add_paragraph('')
				parrafo = p.add_run('CONVENIO')
				parrafo.bold = True
				parrafo.font.size = Pt(20)
				parrafo.add_break()	
				parrafo = p.add_run(convenio.nombre)
				parrafo.bold = True
				parrafo.font.size = Pt(20)
				p.alignment = WD_ALIGN_PARAGRAPH.CENTER

				document.add_page_break()



				for proyecto in proyectos:

					qset1 = (Q(proyecto__id = proyecto['proyecto__id'] ) )  & ( Q(asociado_reporte = True ) )

					if foto and foto is not None:
						qset1 = qset1 & ( Q(tipo_id = foto ) )

					if (fechaDesde and (fechaHasta is not None) ):
						qset1 = qset1 & ( Q(fecha__gte = fechaDesde ) )
						# INICIA EL FILTRO POR LA FECHA DE LA FOTO

					if ( (fechaDesde is not None) and fechaHasta ):
						qset1 = qset1 & ( Q(fecha__lte =fechaHasta)  )
						# FINALIZA EL FILTRO POR LA FECHA DE LA FOTO

					if (fechaDesde and fechaHasta):	
						qset1 = qset1 & ( Q(fecha__range = (fechaDesde , fechaHasta) ) )
					fotos_proyecto =CFotosProyecto.objects.filter(qset1)
					if fotos_proyecto:					
						contratistas = Proyecto.objects.filter(pk=proyecto['proyecto__id'],contrato__tipo_contrato__id = tipo_contrato.contratoProyecto).values('contrato__contratista__nombre').order_by('contrato__contratista__nombre').distinct()
						p=document.add_paragraph('')
						if contratistas.count()>1:
							parrafo = p.add_run('Contratistas: ')
						else:
							parrafo = p.add_run('Contratista: ')						
						parrafo.font.size = Pt(14)
						parrafo.bold = True
						count = 1
						for contratista in contratistas:
							if count==contratistas.count():
								parrafo = p.add_run(contratista['contrato__contratista__nombre'])
							else:
								parrafo = p.add_run(contratista['contrato__contratista__nombre']+' - ')
							count+=1
						parrafo.font.size = Pt(14)
						parrafo.add_break()	
						parrafo = p.add_run('Proyecto: ')
						parrafo.bold = True
						parrafo = p.add_run(proyecto['proyecto__nombre'])
						parrafo.font.size = Pt(12)
						parrafo.add_break()	
						parrafo = p.add_run('Departamento: ')
						parrafo.bold = True
						parrafo = p.add_run(proyecto['proyecto__municipio__departamento__nombre'])
						parrafo.font.size = Pt(12)
						parrafo.add_break()	
						parrafo = p.add_run('Municipio: ')
						parrafo.bold = True
						parrafo = p.add_run(proyecto['proyecto__municipio__nombre'])
						parrafo.font.size = Pt(12)
						
						p.alignment = WD_ALIGN_PARAGRAPH.CENTER
						p=document.add_paragraph('')
					
						for foto in fotos_proyecto:

							newpath = r'static/papelera/'
							filename = str(foto.ruta)
							extension = filename[filename.rfind('.'):]
							nombre = str(foto.id)+extension

							functions.descargarArchivoS3(str(filename), str(newpath)+"/" , nombre )	
							parrafo = document.add_paragraph('')
							run = parrafo.add_run()

							run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(2.8) , height = Inches(2))
							parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

						p=document.add_paragraph('')
			else:
				p=document.add_paragraph('No se encontraron fotos para el convenio seleccionado')


			nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\prueba.docx'
			document.save(nombreArchivo)

			chunk_size = 108192

			nombreDescarga = 'informeFotoProyecto.docx'
			response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
				content_type=mimetypes.guess_type(nombreArchivo)[0])
			response['Content-Length'] = os.path.getsize(nombreArchivo)
			response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
			
			return response
		except Exception as e:
			print(e)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@login_required
def index(request):
	return render(request, 'informe/index.html',{'model':'proyecto','app':'informe'})

@login_required
def interventoria(request):
	return render(request, 'informe/informeInterventoria.html',{'model':'proyecto','app':'informe'})

@login_required
def registroSistema(request):
	return render(request, 'informe/registroSistema.html',{'model':'proyecto','app':'informe'})

@login_required
def fotosProyecto(request):
	qsetEstados = Tipo.objects.filter(app='administradorFotos') 
	return render(request, 'informe/fotosProyecto.html',{ 'estados' : qsetEstados , 'model':'multa','app':'informe'})

@login_required
def informeMME(request):
	tipo=tipoC()
	
	qset = (Q(contrato__tipo_contrato=tipo.m_contrato))& (Q(edita='1')) &(Q(empresa=request.user.usuario.empresa.id) & Q(participa=1))
	ListMacro = EmpresaContrato.objects.filter(qset).order_by("contrato__nombre")

	return render(request, 'informe/informeMME.html',{'macro':ListMacro,'app':'informe'})


@login_required
def obtenerColumnas(request):
	if request.method == 'GET':
		cursor = connection.cursor()
		try:
			opcion = request.GET['opcion']
			empresa_id = request.user.usuario.empresa.id
			cursor.callproc('[dbo].[consultar_informe_dinamico]', [opcion, empresa_id, ' and 1=2'])		
			columns = cursor.description
			lista = []
			i = 0
			print(columns)
			for col in columns:
				lista.append(
					{"procesar": True,
					"columna":columns[i][0],
					"condicional": '',
					"texto": '',
					"entre1": '',
					"entre2": '',
					"tipo":columns[i][3]})
				i = i + 1

			return JsonResponse({'message':'', 'success':'ok', 'data': lista})
		except Exception as e:
			functions.toLog(e, 'informe.obtenerColumnas')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''})
	
@login_required
def generarInformeDinamico(request):
	if request.method == 'POST':
		cursor = connection.cursor()
		try:
			#import pdb; pdb.set_trace()
			opcion = request.POST['opcion']
			empresa_id = request.user.usuario.empresa.id
			condicion = request.POST['condicion']
			incluir = json.loads(request.POST['incluir'])			
			ahora=datetime.now()
			nombre=str(ahora.year)+str(ahora.month)+str(ahora.day)+str(ahora.hour)+str(ahora.minute)+str(ahora.second)
			response = HttpResponse(content_type='application/vnd.ms-excel;charset=utf-8')
			response['Content-Disposition'] = 'attachment; filename="{}-{}.xlsx"'.format(opcion, nombre)
			workbook = xlsxwriter.Workbook(response, {'in_memory': True})
			worksheet = workbook.add_worksheet(opcion)
			format1=workbook.add_format({'border':1,'font_size':12,'bold':True, 'bg_color':'#4a89dc','font_color':'white'})
			format2=workbook.add_format({'border':1})
			format_date=workbook.add_format({'border':1,'num_format': 'yyyy-mm-dd'})
			format_money = workbook.add_format({'border':1, 'num_format': '$#,##0'})

			cursor.callproc('[dbo].[consultar_informe_dinamico]', [opcion, empresa_id, condicion])		
			columns = cursor.description 
			length_list = [len(x) for x in columns]
			lista = [{columns[index][0]:column for index, column in enumerate(value)} for value in cursor.fetchall()]
			for i, width in enumerate(length_list):
				worksheet.set_column(i, i, 30)
			
			columnasV = []
			i = 0
			for col in columns:
				columnasV.append({'columna':columns[i][0], 'tipo': columns[i][3]})
				i = i + 1

			col = 0	
			for column in columnasV:
				if column['columna'] in incluir:				
					worksheet.write(0, col, column['columna'], format1)
					col = col + 1

			row=1				
			for item in lista:
				col=0				
				for column in columnasV:
					if column['columna'] in incluir:
						if column['tipo'] == 10:
							worksheet.write(row, col, item[column['columna']], format_date)
						elif column['columna'] == 'Valor':
							worksheet.write(row, col, item[column['columna']], format_money)	
						elif column['columna'] == 'Soporte':
							worksheet.write_url(row, col, item[column['columna']], string='Ver Soporte', cell_format=format2)		
						else:	
							worksheet.write(row, col, item[column['columna']], format2)
						col = col + 1
				row = row + 1	
				
			workbook.close()	
			return response
		except Exception as e:
			functions.toLog(e, 'informe.generarInformeDinamico')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''})

@login_required
def informeDinamico(request):
	# qsetEstados = Tipo.objects.filter(app='administradorFotos') 
	return render(request, 'informe/informeDinamico.html',{ })

@login_required
def crearEstruturaContrato(request):
	try:
		mcontrato_id = request.GET['mcontrato_id'].split(',')
		mContratos = Contrato.objects.filter(id__in=mcontrato_id)
		ahora=datetime.now()
		archivo=str(ahora.year)+str(ahora.month)+str(ahora.day)+str(ahora.hour)+str(ahora.minute)+str(ahora.second)
		ruta = '{0}\\estructuras\\{1}'.format(settings.STATICFILES_DIRS[0], 'macrocontratos-' + archivo)
		os.mkdir(ruta)
		file = open("{0}\\logs.txt".format(ruta), "w")
		for mContrato in mContratos:			
			contratos = Contrato.objects.filter(mcontrato_id=mContrato.id)
			rutaMacroContrato = '{0}\\{1}'.format(ruta, mContrato.numero.strip())		
			if not os.path.exists(rutaMacroContrato):
				os.mkdir(rutaMacroContrato)			
						
			for contrato in contratos:
				rutaContrato = '{0}\\{1}'.format(rutaMacroContrato, contrato.numero.strip())
				vigencias = contrato.vigencia_contrato() #VigenciaContrato.objects.filter(contrato_id=contrato.id)
				if not os.path.exists(rutaContrato):
					os.mkdir(rutaContrato)
				if vigencias:	
					for vigencia in vigencias:
						filename = str(vigencia.soporte).strip()
						extension = filename[filename.rfind('.'):]
						nombre = re.sub('[^A-Za-z0-9 ]+', '', vigencia.nombre.strip())
						resultado = functions.descargarArchivoS3(str(vigencia.soporte), rutaContrato + '\\', "{0}{1}".format(nombre, extension))
						if resultado['success'] == False:
							file.write("El archivo descargo del S3: {0}/{1}{2}".format(rutaContrato, str(vigencia.soporte), os.linesep))	
		
		return JsonResponse({'message':'La estructura se ha creado con exito.','success':'ok','data': None})		
	except Exception as e:
		functions.toLog(e, 'informe.crearEstruturaContrato')
		return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''})

@login_required
def generarInformeTrimestral(request):
	if request.method == 'GET':
		cursor = connection.cursor()
		try:
					
			ahora=datetime.now()
			nombre=str(ahora.year)+str(ahora.month)+str(ahora.day)+str(ahora.hour)+str(ahora.minute)+str(ahora.second)
			response = HttpResponse(content_type='application/vnd.ms-excel;charset=utf-8')
			response['Content-Disposition'] = 'attachment; filename="Informe de nivel de actualizacion de datos de SININ.xlsx"'
			workbook = xlsxwriter.Workbook(response, {'in_memory': True})
			format1=workbook.add_format({'border':1,'font_size':12,'bold':True, 'bg_color':'#4a89dc','font_color':'white'})
			format2=workbook.add_format({'border':1})
			format_date=workbook.add_format({'border':1,'num_format': 'yyyy-mm-dd'})
			format_money = workbook.add_format({'border':1, 'num_format': '$#,##0'})

			sheets = ['Macro contratos', 'Resumen de contratos', 'Contratos', 'Resumen info tecnica diseno', 
						'Resumen info tecnica replanteo', 'Resumen info tecnica ejecucion',
						'Liq de contratos PRONE 20 13-14', 'Liq de contratos PRONE 2012',
						'Liq de contratos FAER', 'Correspondencia enviada']
			ids = [0, 0, 0, 0, 0, 0, 1, 9, 15, 0]
			f=0
			for sh in sheets:				
				worksheet = workbook.add_worksheet(sh)				
				cursor.callproc('[dbo].[obtener_informe_trimestral]', [sh, ids[f]])		
				columns = cursor.description 
				length_list = [len(x) for x in columns]
				lista = [{columns[index][0]:column for index, column in enumerate(value)} for value in cursor.fetchall()]
				
				# if sh != 'Resumen de contratos':
				for i, width in enumerate(length_list):
					worksheet.set_column(i, i, 30)
				
				columnasV = []
				i = 0
				for col in columns:
					columnasV.append({'columna':columns[i][0], 'tipo': columns[i][3]})
					i = i + 1

				col = 0	
				for column in columnasV:					
					worksheet.write(0, col, column['columna'], format1)
					col = col + 1

				row=1				
				for item in lista:
					col=0				
					for column in columnasV:						
						if column['tipo'] == 10:
							worksheet.write(row, col, item[column['columna']], format_date)
						elif column['columna'] == 'Valor':
							worksheet.write(row, col, item[column['columna']], format_money)	
						elif column['columna'] == 'Soporte':
							worksheet.write_url(row, col, item[column['columna']], string='Ver Soporte', cell_format=format2)		
						else:	
							worksheet.write(row, col, item[column['columna']], format2)
						col = col + 1
					row = row + 1
				# else:
				# 	worksheet.add_table('A1:D1', {'data': lista,
				# 						'columns': [{'header': 'Numero'},
				# 						{'header': 'Nombre'},
				# 						{'header': 'Estado'},
				# 						{'header': 'Cantidad'}
				# 						]})
		
				f = f + 1

			workbook.close()	
			return response
		except Exception as e:
			functions.toLog(e, 'informe.generarInformeTrimestral')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''})

@login_required
def crearEstruturaEncabezadoGiros(request):
	try:
		mContratos = Contrato.objects.filter(contratista__id=4, tipo_contrato__id=12)
		ahora=datetime.now()
		archivo=str(ahora.year)+str(ahora.month)+str(ahora.day)+str(ahora.hour)+str(ahora.minute)+str(ahora.second)
		ruta = '{0}\\estructuras\\{1}'.format(settings.STATICFILES_DIRS[0], 'EncabezadosGiros-' + archivo)
		os.mkdir(ruta)
		file = open("{0}\\logs.txt".format(ruta), "w")
		for mContrato in mContratos:			
			contratos = Contrato.objects.filter(mcontrato__id=mContrato.id)
			rutaMacroContrato = '{0}\\{1}'.format(ruta, mContrato.numero.strip())		
			if not os.path.exists(rutaMacroContrato):
				os.mkdir(rutaMacroContrato)			
						
			for contrato in contratos:
				rutaContrato = '{0}\\{1}'.format(rutaMacroContrato, contrato.numero.strip())
				encabezadosGiros = DEncabezadoGiro.objects.filter(contrato_id=contrato.id)
				if not os.path.exists(rutaContrato):
					os.mkdir(rutaContrato)
				if encabezadosGiros:	
					for giro in encabezadosGiros:
						filename = str(giro.soporte).strip()
						extension = filename[filename.rfind('.'):]
						nombre = re.sub('[^A-Za-z0-9 ]+', '', giro.nombre.nombre.strip())
						try:
							shutil.copy("E:\\AWS S3\\source-sinin-prueba\\media\\{}".format(str(giro.soporte)), rutaContrato + '\\' + "{0}{1}".format(nombre, extension))
						except Exception as e:
							file.write("El archivo descargo del S3: {0}/{1}{2}".format(rutaContrato, str(e), os.linesep))
						
						# resultado = functions.descargarArchivoS3(str(giro.soporte), rutaContrato + '\\', "{0}{1}".format(nombre, extension))
						# if resultado['success'] == False:
						# 	file.write("El archivo descargo del S3: {0}/{1}{2}".format(rutaContrato, str(giro.soporte), os.linesep))	
		
		return JsonResponse({'message':'La estructura se ha creado con exito.','success':'ok','data': None})		
	except Exception as e:
		functions.toLog(e, 'informe.crearEstruturaEncabezadoGiros')
		return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''})


