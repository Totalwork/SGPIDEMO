from django.shortcuts import render
#,render_to_response
from django.template import RequestContext
from django.db.models import Q

from django.urls import reverse
from django.core.serializers import serialize
from rest_framework import viewsets, serializers

from contrato.models import EmpresaContrato
from contrato.enumeration import tipoC


from parametrizacion.models import Funcionario
from empresa.models import Empresa
from empresa.views import EmpresaSerializer

from tipo.models import Tipo
from tipo.views import TipoSerializer

from usuario.views import UsuarioSerializer
from usuario.models import Usuario

from contrato.models import Contrato
from contrato.enumeration import tipoC

from proyecto.models import Proyecto

from avance_de_obra.models import BCronograma,Porcentaje,CIntervaloCronograma

from logs.models import Logs,Acciones
from django.db import transaction
from django.contrib.auth.decorators import login_required

from rest_framework.renderers import JSONRenderer
from rest_framework.views import exception_handler
from rest_framework.response import Response
from rest_framework.request import Request
from rest_framework import status
from rest_framework.pagination import PageNumberPagination


from django.http import HttpResponse,JsonResponse
from rest_framework.parsers import MultiPartParser, FormParser
from django.db.models.deletion import ProtectedError
from django.core.paginator import Paginator

from .models import Planilla,Tag
from .enumeration import TipoInforme

import xlsxwriter
from datetime import *
import os
from coasmedas.functions import functions
from docx import Document
from docx.shared import Inches , Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.text import WD_ALIGN_PARAGRAPH , WD_COLOR_INDEX
from django.conf import settings


import mimetypes
from django.http import StreamingHttpResponse
from wsgiref.util import FileWrapper
from django.db import connection
# Create your views here.


#Api rest para planilla
class PlanillaSerializer(serializers.HyperlinkedModelSerializer):

	empresa=EmpresaSerializer(read_only=True)
	empresa_id=serializers.PrimaryKeyRelatedField(write_only=True,queryset=Empresa.objects.all())

	tipo=TipoSerializer(read_only=True)
	tipo_id=serializers.PrimaryKeyRelatedField(write_only=True,queryset=Tipo.objects.filter(app='informe_ministerio_planilla'))

	class Meta:
		model = Planilla
		fields=('id','archivo','empresa','empresa_id','tipo_id','tipo',)

class InformePlanillaViewSet(viewsets.ModelViewSet):
	
	model=Planilla
	model_log=Logs
	model_acciones=Acciones
	nombre_modulo='informe_ministerio.planilla'
	queryset = model.objects.all()
	serializer_class = PlanillaSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)


	def list(self, request, *args, **kwargs):
		try:
			queryset = super (InformePlanillaViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			sin_paginacion= self.request.query_params.get('sin_paginacion',None)

			if dato:
				qset = qset &(
					Q(archivo__icontains=dato)
					)
			
				queryset = self.model.objects.filter(qset)


			if sin_paginacion is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					return self.get_paginated_response({'message':'','success':'ok',
					'data':serializer.data})
	
			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok',
					'data':serializer.data})			
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()
			try:
				serializer = PlanillaSerializer(data=request.DATA,context={'request': request})

				if serializer.is_valid():
					serializer.save(archivo=self.request.FILES.get('archivo'),tipo_id=request.DATA['tipo_id'],empresa_id=request.DATA['empresa_id'])
					logs_model=self.model_log(usuario_id=request.user.usuario.id,accion=self.model_acciones.accion_crear,nombre_modelo=self.nombre_modulo,id_manipulado=serializer.data['id'])
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
			  		'data':''},status=status.HTTP_400_BAD_REQUEST)

	@transaction.atomic
	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			sid = transaction.savepoint()
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = PlanillaSerializer(instance,data=request.DATA,context={'request': request},partial=partial)
				if serializer.is_valid():
					serializer.save(archivo=self.request.FILES.get('archivo'),tipo_id=request.DATA['tipo_id'],empresa_id=request.DATA['empresa_id'])
					logs_model=self.model_log(usuario_id=request.user.usuario.id,accion=self.model_acciones.accion_actualizar,nombre_modelo=self.nombre_modulo,id_manipulado=instance.id)
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

	def destroy(self,request,*args,**kwargs):
		try:
			instance = self.get_object()
			self.perform_destroy(instance)
			return Response({'message':'El registro se ha eliminado correctamente','success':'ok',
				'data':''},status=status.HTTP_201_CREATED)
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''},status=status.HTTP_400_BAD_REQUEST)	

#Fin Api rest para planilla




#Api rest para tag
class TagSerializer(serializers.HyperlinkedModelSerializer):

	planilla=PlanillaSerializer(read_only=True,many=True)

	class Meta:
		model = Tag
		fields=('id','nombre','modelo','campo','planilla',)

class TagViewSet(viewsets.ModelViewSet):
	
	model=Tag
	model_log=Logs
	model_acciones=Acciones
	nombre_modulo='informe_ministerio.tag'
	queryset = model.objects.all()
	serializer_class = TagSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)


	def list(self, request, *args, **kwargs):
		try:
			queryset = super (TagViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			sin_paginacion= self.request.query_params.get('sin_paginacion',None)

			if dato:
				qset = qset &(
					Q(nombre__icontains=dato)
					)
			
				queryset = self.model.objects.filter(qset)


			if sin_paginacion is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					return self.get_paginated_response({'message':'','success':'ok',
					'data':serializer.data})
	
			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok',
					'data':serializer.data})			
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()
			try:
				serializer = TagSerializer(data=request.DATA,context={'request': request})

				if serializer.is_valid():
					serializer.save()
					logs_model=self.model_log(usuario_id=request.user.usuario.id,accion=self.model_acciones.accion_crear,nombre_modelo=self.nombre_modulo,id_manipulado=serializer.data['id'])
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
			  		'data':''},status=status.HTTP_400_BAD_REQUEST)

	@transaction.atomic
	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			sid = transaction.savepoint()
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = TagSerializer(instance,data=request.DATA,context={'request': request},partial=partial)
				if serializer.is_valid():
					serializer.save()
					logs_model=self.model_log(usuario_id=request.user.usuario.id,accion=self.model_acciones.accion_actualizar,nombre_modelo=self.nombre_modulo,id_manipulado=instance.id)
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

	def destroy(self,request,*args,**kwargs):
		try:
			instance = self.get_object()
			self.perform_destroy(instance)
			return Response({'message':'El registro se ha eliminado correctamente','success':'ok',
				'data':''},status=status.HTTP_201_CREATED)
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''},status=status.HTTP_400_BAD_REQUEST)	

#Fin Api rest para tag


@login_required
def excel_verificar_datos(request):
	try:
		response = HttpResponse(content_type='application/vnd.ms-excel;charset=utf-8')
		response['Content-Disposition'] = 'attachment; filename="informe_verificacion.xls"'
		
		workbook = xlsxwriter.Workbook(response, {'in_memory': True})
		worksheet = workbook.add_worksheet('Datos')
		format1=workbook.add_format({'border':2,'font_size':9,'bold':True,'align':'center'})
		format3=workbook.add_format({'border':0,'font_size':9,'bold':True,'align':'center'})
		format2=workbook.add_format({'border':2,'font_size':9})
		format4=workbook.add_format({'border':0,'font_size':9})

		meses=['Enero','Febrero','Marzo','Abril','Mayo','Junio','Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']

		worksheet.set_column('A:E',20,format3)

		disenos=None
		qset=None
		contrato_id=request.GET.get('contrato_id',None)
		desde=request.GET.get('desde',None)
		hasta=request.GET.get('hasta',None)
		ano=request.GET.get('ano',None)

		contrato=Contrato.objects.get(pk=contrato_id)
		worksheet.merge_range('A1:F1', 'Proyectos en los que el avance de obra no refleja cambios, respecto al periodo anterior', format3)
		worksheet.write('A3', 'Contrato', format3)
		worksheet.write('B3', contrato.nombre, format4)
		worksheet.write('A4', 'Periodo analizado', format3)
		worksheet.write('B4', meses[int(desde)-1]+" - "+meses[int(hasta)-1], format4)

		worksheet.write('A5', 'Interventor', format1)
		worksheet.write('B5', 'Departamento', format1)
		worksheet.write('C5', 'Municipio', format1)
		worksheet.write('D5', 'Proyecto', format1)		
		worksheet.write('E5', 'Cronograma', format1)


		proyectos=Proyecto.objects.filter(mcontrato_id=contrato_id)
		fecha_limite=datetime.strptime(str(ano)+"-"+str(hasta)+"-30", '%Y-%m-%d').date()
		fecha_desde=datetime.strptime(str(ano)+"-"+str(desde)+"-1", '%Y-%m-%d').date()
		intervalo = (int(hasta)- int(desde))+1;
		fecha=fecha_limite- timedelta(days=intervalo*365/12)
		fecha_atras=datetime.strptime(str(fecha), '%Y-%m-%d').date()
		

		row=5
		col=0
		tipo=tipoC()
		for item in proyectos:
			cronogramas=BCronograma.objects.filter(proyecto_id=item.id)

			interventor=""
			for contratos in item.contrato.all():
				if contratos.tipo_contrato.id==tipo.interventoria:
					interventor=contratos.nombre

			for item2 in cronogramas:
				if item2.fecha_inicio_cronograma<fecha_atras:
					dias_atras=(item2.fecha_inicio_cronograma - fecha_atras)/item2.periodicidad.numero_dias
					dias_siguientes=(item2.fecha_inicio_cronograma - fecha_limite)/item2.periodicidad.numero_dias
					dias_desde=(item2.fecha_inicio_cronograma - fecha_desde)/item2.periodicidad.numero_dias
					
					intervalos=CIntervaloCronograma.objects.filter(cronograma__id=item2.id,tipo_linea=3)

					if abs(dias_atras.days)<=len(intervalos) and abs(dias_desde.days)<=len(intervalos):
						porcentaje=Porcentaje.objects.filter(intervalo__intervalo__lte=abs(dias_atras.days),intervalo__cronograma_id=item2.id,intervalo__sinAvance=False,tipo_linea=3,intervalo__tipo_linea=3).values('porcentaje').order_by('-porcentaje').first()
						porcentaje2=Porcentaje.objects.filter(intervalo__intervalo__lte=abs(dias_siguientes.days),intervalo__cronograma_id=item2.id,intervalo__sinAvance=False,tipo_linea=3,intervalo__tipo_linea=3).values('porcentaje').order_by('-porcentaje').first()
						
						if porcentaje is not None and porcentaje2 is not None:
							if porcentaje['porcentaje']==porcentaje2['porcentaje'] and porcentaje['porcentaje']<100 and porcentaje2['porcentaje']<100:
								worksheet.write(row, col,interventor,format2)
								worksheet.write(row, col+1,item.municipio.departamento.nombre,format2)
								worksheet.write(row, col+2,item.municipio.nombre,format2)
								worksheet.write(row, col+3,item.nombre,format2)
								worksheet.write(row, col+4,item2.nombre,format2)

								row +=1
		workbook.close()

		return response

	except Exception as e:
		functions.toLog(e,'informe_ministerio')
		return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''})	


@login_required
def generar_informe(request):
	try:
		tipoinforme=TipoInforme()
		t = datetime.now()
		nombreArchivo = str(t.year)+str(t.month)+str(t.day)+str(t.hour)+str(t.minute)+str(t.second)		
		#newpath = r'media/administrador_tarea/'+str(nombreArchivo)+"/"
		#newpath = r'media/administrador_tarea/201771915951/informe_ministerio_eca.docx'

		# if not os.path.exists(newpath):
		# 	os.makedirs(newpath)

		planilla=Planilla.objects.filter(empresa_id=request.user.usuario.empresa.id,tipo_id=tipoinforme.informe_eca)
		#functions.descargarArchivoS3(str(planilla[0].archivo),str(newpath))		
		document =Document(planilla[0].archivo)
		tags=Tag.objects.filter(planilla=planilla[0].id)

		# styles=document.styles
		# table_styles = [s for s in styles ]

		# for style in table_styles:
		# 	print(style.name)

		meses=['Enero','Febrero','Marzo','Abril','Mayo','Junio','Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']
		dias=['31','28','31','30','31','30','31','31','30','31','30','31']

		cursor = connection.cursor()
		for item in tags:
			for p in document.paragraphs:
				if item.nombre in p.text:
					if item.tag_especial==False:
						if request.GET.get(item.nombre_variable) is not None:
							cursor.execute("SELECT "+item.campo+" FROM "+item.modelo+" "+item.inner+" where "+item.modelo+".id="+request.GET.get(item.nombre_variable));
							texto=""
							index=0
							for item2 in cursor.fetchall():
								texto_consulta=str(item2[index].encode('utf-8'))
								texto=texto+" "+texto_consulta.decode('utf-8')
								index=index+1

							if item.mayuscula==False:
								text_after=p.text.replace(item.nombre, texto)
							else:
								text_after=p.text.replace(item.nombre, texto.upper())

							p_list=p.runs
							style_p=p.runs[0].style
							p.clear()
							p.add_run(text_after,style_p)
					else: 
						if item.nombre in '<tabla>' or item.nombre in '<img>':
							texto=str(p.text.encode('utf-8')).split(";")
							encabezado=str(texto[1]).split(",")	
							p.clear()	
							variables=str(texto[3]).split(",")
							mylist=[]
							for item_variable in variables:
								if item_variable!='':
									if item_variable =='desde' or item_variable =='hasta':
										mes=''
										if int(request.GET.get(item_variable))<10:
											mes='0'+str(request.GET.get(item_variable.strip()))
										else:
											mes=request.GET.get(item_variable.strip())
										dia=''
										if item_variable=='desde':
											dia='01'

										if item_variable=='hasta':
											dia=str(dias[int(request.GET.get(item_variable.strip()))-1])
										
										fecha=str(request.GET.get('ano'))+'-'+mes+'-'+dia
										mylist.append(fecha)
									else:
										mylist.append(request.GET.get(item_variable.strip()))

							cursor.execute(texto[2],mylist)
							table = document.add_table(rows=1, cols=len(encabezado),style='mibordepersonal')
							addTableAfterParagraph(table,p)

							personal_cells = table.rows[0].cells
							index=0

							for item3 in encabezado:
								if item3 !='':
									parrafo=personal_cells[index].add_paragraph()
									run = parrafo.add_run(item3.strip().upper())
									run.bold = True
									run.font.size = Pt(9)
									parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
									index=index+1

							total=False
							posicion_variable=None
							nombre_total=''
							if len(texto)>4:
								totales=str(texto[4]).split(",")
								total=True
								nombre_total=totales[0]
								posicion_variable=totales[1]

							total_pagado=0
							index_principal=0

							for item2 in cursor.fetchall():
								index=0
								row_cells = table.add_row().cells
								for x in range(0,len(item2)):
									parrafo = row_cells[index].add_paragraph()	
									if x==0:
										run = parrafo.add_run(str(index_principal+1))
									else:
										if item2[x] is None:
											run = parrafo.add_run('')
										else:	
											run = parrafo.add_run(str(item2[x]).decode('utf-8'))
										
									run.font.size = Pt(9)
									parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
									index=index+1	


								if total==True:
									total_pagado=total_pagado+item2[int(posicion_variable)]

								index_principal=index_principal+1	

							if total==True:
								row_cells = table.add_row().cells
								parrafo = row_cells[0].add_paragraph()	
								run = parrafo.add_run(str(nombre_total).upper())
								run.font.size = Pt(9)
								parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
								a, b = row_cells[:int(posicion_variable)-1]
								a.merge(b)

								parrafo = row_cells[int(posicion_variable)].add_paragraph()	
								run = parrafo.add_run(str(total_pagado))
								run.font.size = Pt(9)
								parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
								a, b = row_cells[int(posicion_variable):len(encabezado)]
								a.merge(b)


						else:
							if request.GET.get(item.nombre_variable) is not None:
								if item.nombre_variable=='desde' or item.nombre_variable=='hasta':
									if item.mayuscula==False:
										text_after=p.text.replace(item.nombre, meses[int(request.GET.get(item.nombre_variable))-1])
									else:
										text_after=p.text.replace(item.nombre, str(meses[int(request.GET.get(item.nombre_variable))-1]).upper())
								else:
									if item.mayuscula==False:
										text_after=p.text.replace(item.nombre, request.GET.get(item.nombre_variable))
									else:
										text_after=p.text.replace(item.nombre, str(request.GET.get(item.nombre_variable)).upper())

								p_list=p.runs
								style_p=p.runs[0].style
								p.clear()
								p.add_run(text_after,style_p)


		# for table in document.tables:
		# 	for row in table.rows:
		# 		for cell in row.cells:
		# 			print cell.text.encode('utf-8')


		#                 text = inline[i].text.replace(item.nombre, 'new text')
		#                 inline[i].text = text

				# text = document.paragraphs[p].text.replace(item.nombre, 'prueba')
				# document.paragraphs[p].text = text


		nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\prueba.docx'
		document.save(nombreArchivo)


		chunk_size = 108192

		nombreDescarga = 'informe_ministerio_eca.docx'
		response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
				content_type=mimetypes.guess_type(nombreArchivo)[0])
		response['Content-Length'] = os.path.getsize(nombreArchivo)
		response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
			
		return response

	except Exception as e:
		functions.toLog(e,'informe_ministerio')
		return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''})	



def addTableAfterParagraph(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


@login_required
def informe_ministerio(request):
	tipo=tipoC()
	qset = (Q(contrato__tipo_contrato=tipo.m_contrato)) &(Q(empresa=request.user.usuario.empresa.id) & Q(participa=1))
	ListMacro = EmpresaContrato.objects.filter(qset)

	ListFuncionario=Funcionario.objects.filter(empresa_id=request.user.usuario.empresa.id)

	return render(request, 'informe_ministerio/informe_ministerio.html',{'model':'informe','app':'informe','macro':ListMacro,'funcionarios':ListFuncionario})
