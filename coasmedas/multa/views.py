# -*- coding: utf-8 -*- 
from sinin4.functions import functions

# import win32com.client as win32

import shutil
import time
import os
from io import StringIO
from django.db.models import Max

from django.shortcuts import render, render_to_response
from django.urls import reverse
from django.db.models import Max

from rest_framework import viewsets, serializers, response
from django.db.models import Q
from rest_framework.validators import UniqueTogetherValidator

from django.db import IntegrityError,transaction
from .models import SolicitudConsecutivo , ConjuntoEvento , Evento , Solicitud , SolicitudEmpresa , SolicitudEvento , SolicitudHistorial , SolicitudSoporte , SolicitudApelacion , SolicitudPronunciamiento
from django.template import RequestContext
from rest_framework.renderers import JSONRenderer
from rest_framework.views import exception_handler
from rest_framework.response import Response
from rest_framework.request import Request
from rest_framework import status
from rest_framework.pagination import PageNumberPagination
import xlsxwriter
import json
from django.http import HttpResponse,JsonResponse
from rest_framework.parsers import MultiPartParser, FormParser
from django.db import connection
from django.contrib.auth.decorators import login_required

from babel.dates import format_date, format_datetime, format_time
from babel.numbers import format_number, format_decimal, format_percent
from numbertoletters import number_to_letters

import openpyxl
from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from openpyxl.styles import Border, Side, PatternFill, Font, GradientFill, Alignment


from contrato.models import Contrato,EmpresaContrato,VigenciaContrato
from contrato.views import ContratoSerializer
from tipo.models import Tipo
from tipo.views import TipoSerializer

from empresa.models import Empresa
from empresa.views import EmpresaSerializer
from usuario.models import Usuario ,Persona

from estado.models import Estado , Estados_posibles
from estado.views import EstadoSerializer
from parametrizacion.models import  Funcionario , Departamento
from parametrizacion.views import FuncionarioSerializer
from correspondencia.models import  CorrespondenciaEnviada ,CorrespondenciaConsecutivo ,CorrespondenciaSoporte ,CorresPfijo ,CorrespondenciaPlantilla, CorrespondenciaRadicado
from correspondencia.views import CorrespondenciaEnviadaSerializer
from correspondencia.functions import buscaEtiqueta , buscaEtiqueta2
from correspondencia_recibida.models import CorrespondenciaRecibida , CorrespondenciaRecibidaAsignada
from empresa.models import Empresa , EmpresaAcceso

from empresa.models import Empresa

from multa.enumeration import EstadoMulta , Notificaciones
from contrato.enumeration import tipoC
from proyecto.models import Proyecto 
from proyecto.views import ProyectoSerializer
from contrato.enumeration import tipoC

from logs.models import Logs,Acciones
from django.db import transaction,connection
from django.db.models.deletion import ProtectedError

from datetime import date
import datetime

from django.conf import settings
from docx import Document
from docx.shared import Inches , Pt , Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import os
import mimetypes
from django.http import StreamingHttpResponse
from wsgiref.util import FileWrapper

from rest_framework.decorators import api_view


from adminMail.models import Mensaje

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml 

from docx.oxml.shared import OxmlElement

from docx.shared import Pt

# encoding=utf8  
import sys  

# reload(sys)  
# sys.setdefaultencoding('utf8')

global tipoRadicado , tipoConsecutivo , tagsHtml
tipoRadicado = 56
tipoConsecutivo = 57 

# extensiones validas para los soportes de la multas
extensiones_permitidas = ['.doc' , '.docx' , '.pdf' ]


@api_view(['GET'])
def select_filter_multa(request):
	if request.method == 'GET':
		try:
			
			tipo_contrato = tipoC()

			empresaActual = request.user.usuario.empresa.id

			qsetEstadosMultas= Estado.objects.filter(app = "multa").values('id', 'nombre')
			contratistas = Contrato.objects.filter(empresacontrato__participa = True , empresacontrato__empresa = request.user.usuario.empresa.id , tipo_contrato_id = tipo_contrato.contratoProyecto).values_list('contratista_id')

			# CONTRATISTAS DEL PROYECTO
			# DATOS REQUERIDO PARA LA CONSULTA DE PROYECTO
			qsetEmpresas_contratistas = Empresa.objects.filter(id__in = contratistas , esContratista = True).values('id', 'nombre')

			# id del contrato no asignado
			qsetMcontratos = Contrato.objects.filter(empresacontrato__participa = True , empresacontrato__empresa = request.user.usuario.empresa.id , tipo_contrato_id = tipo_contrato.m_contrato).values('id', 'nombre').exclude(pk = 1843)

			# empresas solicitantes
			empresas_acceso = EmpresaAcceso.objects.filter(empresa_id = empresaActual).values_list('empresa_ver_id')
			empresas = Empresa.objects.filter(id__in = empresas_acceso , esContratista = True).values('id', 'nombre')
			
			return JsonResponse({'message':'','success':'ok','data':{
										  'estados_solicitudes' : list(qsetEstadosMultas)
										, 'solicitantes' : list(empresas)
										, 'macro_contratos' : list(qsetMcontratos)
										, 'contratistas' : list(qsetEmpresas_contratistas)
									 }})

		except Exception as e:
			print(e)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)



# Create your views here.
class EmpresaLiteSerializer(serializers.HyperlinkedModelSerializer):
	"""docstring for ClassName"""	
	class Meta:
		model = Empresa
		fields=('id' , 'nombre'	)

class ContratoLiteSerializer(serializers.HyperlinkedModelSerializer):
	"""docstring for ClassName"""	
	estado = EstadoSerializer(read_only=True)
	tipo_contrato = TipoSerializer(read_only=True)
	contratista = EmpresaLiteSerializer(read_only=True)
	contratante = EmpresaLiteSerializer(read_only=True)
	class Meta:
		model = Contrato
		fields=('id' , 'nombre'	, 'contratista' , 'contratante' , 'tipo_contrato' , 'numero' , 'estado')

class ContratoSuperLiteSerializer(serializers.HyperlinkedModelSerializer):
	"""docstring for ClassName"""	
	contratista = EmpresaLiteSerializer(read_only=True)

	class Meta:
		model = Contrato
		fields=('id' , 'nombre'	, 'contratista' , 'numero' )

class EmpresaContratoLiteSerializer(serializers.HyperlinkedModelSerializer):

	contrato = ContratoLiteSerializer(read_only=True)

	class Meta:
		model = EmpresaContrato
		fields=('id','contrato')

class CorresPfijoLiteSerializer(serializers.HyperlinkedModelSerializer):
	"""docstring for ClassName"""	
	class Meta:
		model = CorresPfijo
		fields=('id', 'nombre')

class CorrespondenciaEnviadaLiteSerializer(serializers.HyperlinkedModelSerializer):
	
	prefijo = CorresPfijoLiteSerializer(read_only =True)
	
	class Meta:
		model = CorrespondenciaEnviada
		fields=('id' , 'consecutivo', 'prefijo'	)



class SolicitudConsecutivoLiteSaveSerializer(serializers.HyperlinkedModelSerializer):

	empresa_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Empresa.objects.all())
	class Meta:
		model = SolicitudConsecutivo
		fields = ('consecutivo' , 'empresa_id' )

class SolicitudConsecutivoSerializer(serializers.HyperlinkedModelSerializer):
	empresa = EmpresaLiteSerializer(read_only = True)
	empresa_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Empresa.objects.all())
	class Meta:
		model = SolicitudConsecutivo
		fields = ('consecutivo' , 'empresa' ,'empresa_id' )

class SolicitudConsecutivoViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de informacion , consecutivos  de multas asignados por empresa ,
	puede utilizar el parametro (consecutivo , empresa) a traves del cual, se podra buscar por cada uno de estos filtros
	<br>consecutivo = [numero] <br>
	empresa = [numero].
	"""
	model = SolicitudConsecutivo
	queryset = model.objects.all()
	serializer_class = SolicitudConsecutivoSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(SolicitudConsecutivoViewSet, self).get_queryset()
			consecutivo = self.request.query_params.get('consecutivo', None)
			empresa = self.request.query_params.get('empresa', None)

			qset=(~Q(id=0))

			if consecutivo or empresa:
				if consecutivo:
					qset = qset & ( Q(consecutivo = consecutivo) )					
				if empresa:
					qset =  qset & ( Q(empresa = empresa) )

			queryset = self.model.objects.filter(qset)
	
			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})

			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})
		except Exception as e:
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			
			try:
				serializer = SolicitudConsecutivoLiteSaveSerializer(data=request.DATA,context={'request': request})
				# print serializer
				if serializer.is_valid():

					sc = SolicitudConsecutivo.objects.filter(empresa_id = request.DATA['empresa_id'])
					
					if sc.count()>0:
						return Response({'message':'La empresa solo puede tener un unico registro','success':'fail',
							'data':serializer.data},status=status.HTTP_400_BAD_REQUEST)
					else:						
						serializer.save(empresa_id =  request.DATA['empresa_id'] )
						return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
							'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					# print(serializer.errors)
					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "consecutivo" in serializer.errors:
						mensaje = serializer.errors['consecutivo'][0]+" En el campo consecutivo"
					else: 
						mensaje = 'datos requeridos no fueron recibidos'
					return Response({'message': mensaje ,'success':'fail',
			 			'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				# print(e)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error', 'data':''},status=status.HTTP_400_BAD_REQUEST)

	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = SolicitudConsecutivoSerializer(instance,data=request.DATA,context={'request': request},partial=partial)
				if serializer.is_valid():
					self.perform_update(serializer)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

class ConjuntoEventoSerializer(serializers.HyperlinkedModelSerializer):

	class Meta:
		model = ConjuntoEvento
		fields = ('id' , 'nombre' )

class ConjuntoEventoViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de informacion tecnica del proyectos ,
	puede utilizar el parametro (campo , proyecto) a traves del cual, se podra buscar por cada uno de estos filtros
	<br>campo = [numero] <br>
	proyecto = [numero].
	"""
	model = ConjuntoEvento
	queryset = model.objects.all()
	serializer_class = ConjuntoEventoSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(ConjuntoEventoViewSet, self).get_queryset()
			nombre = self.request.query_params.get('dato', None)
			
			if nombre: 
				qset = ( Q(nombre__icontains = nombre) )
				queryset = self.model.objects.filter(qset)
		
			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})

			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			try:
				serializer = ConjuntoEventoSerializer(data=request.DATA,context={'request': request})
				if serializer.is_valid():
					serializer.save()
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					if "non_field_errors" in serializer.errors:
						mensaje = 'Existe un registro con el nombre ingresado.'
					elif "nombre" in serializer.errors:
						mensaje = serializer.errors['nombre'][0]
					else:
						mensaje = 'datos requeridos no fueron recibidos'
					return Response({'message': mensaje ,'success':'fail',
			 			'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:				
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error', 'data':''},status=status.HTTP_400_BAD_REQUEST)

	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = ConjuntoEventoSerializer(instance,data=request.DATA,context={'request': request},partial=partial)
				if serializer.is_valid():
					self.perform_update(serializer)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

class EventoSerializer(serializers.HyperlinkedModelSerializer):

	conjunto = ConjuntoEventoSerializer(read_only = True)
	conjunto_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=ConjuntoEvento.objects.all())

	class Meta:
		model = Evento
		fields = ('id' , 'nombre' , 'valor' , 'conjunto' , 'conjunto_id' )

		validators=[
				serializers.UniqueTogetherValidator(
				queryset=model.objects.all(),
				fields=( 'conjunto_id' , 'nombre' ),
				message=('El nombre del evento  no puede  estar repetido en el conjunto.')
				)
				]

class EventoViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de informacion tecnica del proyectos ,
	puede utilizar el parametro (campo , proyecto) a traves del cual, se podra buscar por cada uno de estos filtros
	<br>campo = [numero] <br>
	proyecto = [numero].
	"""
	model = Evento
	queryset = model.objects.all()
	serializer_class = EventoSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(EventoViewSet, self).get_queryset()
			nombre = self.request.query_params.get('dato', None)
			conjunto = self.request.query_params.get('conjunto', None)

			qset=(~Q(id=0))
			if nombre:
				qset = qset & ( Q(nombre__icontains = nombre) )
			if conjunto:
				qset = qset & ( Q(conjunto_id = conjunto) )

			queryset = self.model.objects.filter(qset)
			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})

			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			try:
				serializer = EventoSerializer(data=request.DATA,context={'request': request})
				if serializer.is_valid():
					serializer.save(conjunto_id =  request.DATA['conjunto_id'])
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:

					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "nombre" in serializer.errors:
						mensaje = serializer.errors["nombre"][0]
					elif "valor" in serializer.errors:
						mensaje = serializer.errors["valor"][0]
					else:
						mensaje = 'datos requeridos no fueron recibidos'

					return Response({'message': mensaje ,'success':'fail', 'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				print(e)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
			  		'data':''},status=status.HTTP_400_BAD_REQUEST)

	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = EventoSerializer(instance,data=request.DATA,context={'request': request},partial=partial)
				if serializer.is_valid():
					serializer.save(conjunto_id =  request.DATA['conjunto_id'])
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:

					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "nombre" in serializer.errors:
						mensaje = serializer.errors["nombre"][0]
					elif "valor" in serializer.errors:
						mensaje = serializer.errors["valor"][0]
					else:
						mensaje = 'datos requeridos no fueron recibidos'

					return Response({'message': mensaje ,'success':'fail', 'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				print(e)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)


class SolicitudSerializer(serializers.HyperlinkedModelSerializer):

	contrato = ContratoLiteSerializer(read_only = True)
	contrato_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Contrato.objects.all())

	correspondenciasolicita = CorrespondenciaEnviadaSerializer(read_only = True)
	correspondenciasolicita_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=CorrespondenciaEnviada.objects.all())

	firmaImposicion = FuncionarioSerializer(read_only = True)

	correspondenciadescargo = CorrespondenciaEnviadaSerializer(read_only = True)

	class Meta:
		model = Solicitud
		fields = (	'id','contrato' , 'contrato_id' , 'correspondenciasolicita' , 'correspondenciasolicita_id'
					, 'consecutivo' , 'diasApelar' , 'firmaImposicion' , 'fechaDiligencia' , 'soporte' 
					, 'valorSolicitado' , 'valorImpuesto' , 'correspondenciadescargo','codigoReferencia'
				)

# solo mostrar datos
class SolicitudLiteSerializer(serializers.HyperlinkedModelSerializer):

	contrato = ContratoLiteSerializer(read_only = True)

	correspondenciasolicita = CorrespondenciaEnviadaLiteSerializer(read_only = True)

	class Meta:
		model = Solicitud
		fields = (	'id','contrato' , 'correspondenciasolicita' , 'consecutivo' , 'valorSolicitado' ,'valorImpuesto' , 'codigoReferencia')


# SOLICITUDES ELABORADAS
class SolicitudSuperLiteSerializer(serializers.HyperlinkedModelSerializer):

	contrato = ContratoSuperLiteSerializer(read_only = True)

	correspondenciasolicita = CorrespondenciaEnviadaLiteSerializer(read_only = True)

	class Meta:
		model = Solicitud
		fields = (	'id', 'consecutivo' , 'contrato' , 'correspondenciasolicita' , 'valorImpuesto', 'codigoReferencia' )

# se usa el serializador por ID
class SolicitudPorIdSerializer(serializers.HyperlinkedModelSerializer):

	contrato = ContratoLiteSerializer(read_only = True)
	estado = serializers.SerializerMethodField()
	fechasolicitud = serializers.SerializerMethodField()

	class Meta:
		model = Solicitud
		fields = (	'id', 'consecutivo' , 'contrato' , 'estado' , 'fechasolicitud' ,'valorSolicitado' , 'valorImpuesto' )


	def get_estado(self, obj):
		try:
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.id).latest('id')
			lista = {'id': datos.estado_id, 'nombre': datos.estado.nombre}
			return lista
		except SolicitudHistorial.DoesNotExist:
			lista = {'id': 0, 'nombre': ''}
			return lista

	def get_fechasolicitud(self, obj):
		try:
			# print obj.solicitud_id
			es=EstadoMulta()
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.id , estado_id = es.solicitada).latest('id')
			
			fecha = datos.fecha
			
			return fecha.strftime("%Y-%m-%d") 
		except Exception as e:
			return ''


class SolicitudViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de informacion tecnica del proyectos ,
	puede utilizar el parametro (campo , proyecto) a traves del cual, se podra buscar por cada uno de estos filtros
	<br>campo = [numero] <br>
	proyecto = [numero].
	"""
	model = Solicitud
	queryset = model.objects.all()
	serializer_class = SolicitudSerializer
	nombre_modulo='multa.solicitud'

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = SolicitudPorIdSerializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(SolicitudViewSet, self).get_queryset()

			correspondenciasolicita_id = self.request.query_params.get('correspondenciasolicita', None)
			consecutivo = self.request.query_params.get('consecutivo', None)
			contrato_id = self.request.query_params.get('contrato_id', None)
			correspondenciadescargo_id = self.request.query_params.get('correspondenciadescargo', None)
			firmaImposicion_id = self.request.query_params.get('firmaImposicion', None)

			qset=(~Q(correspondenciasolicita_id=0))

			if correspondenciasolicita_id:
				qset = qset & ( Q(correspondenciasolicita_id = correspondenciasolicita_id) )
			if consecutivo:
				qset = qset & ( Q(consecutivo__icontains = consecutivo) )
			if contrato_id:
				qset = qset & ( Q(contrato_id = contrato_id) )
			if correspondenciadescargo_id:
				qset = qset & ( Q(correspondenciadescargo_id = correspondenciadescargo_id) )
			if firmaImposicion_id:
				qset = qset & ( Q(firmaImposicion_id = firmaImposicion_id) )
			

			if qset:
				queryset = self.model.objects.filter(qset)
			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})

			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()
			try:

				request.DATA['valorImpuesto'] = request.DATA['valorSolicitado']
				fecha = request.DATA['fechaEnvio']
				# INSERTAR PRIMERO CORRESPONDENCIA PARA PODER INSERTAR UNA SOLICITUD DE MULTA
				usuarioActual = request.DATA['usuarioSolicitante_id']
				d = datetime.datetime.now()
				u = Usuario.objects.get(pk = usuarioActual)
				p = Persona.objects.get(pk = u.persona_id)
				empresaIdActual = u.empresa.id
				empresaConsecutivoDigitado = u.empresa.consecutivoDigitado

				destinatario = int(request.DATA['destinatario_id'])
				usuarioDestino = Usuario.objects.get(pk = destinatario)

				if empresaConsecutivoDigitado:
					consecutivo = request.DATA['consecutivo_carta'] if 'consecutivo_carta' in request.DATA else None;

					# se valida que el consecutivo que esta digitado no exista
					validacion = CorrespondenciaEnviada.objects.filter(consecutivo = consecutivo , prefijo_id = request.DATA['prefijo_id'] , anoEnvio = int(fecha[:4]))
					
					if validacion:
						mensaje='El consecutivo ya se encuentra registrado.'
						return Response({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)	


				else:
				# SE  INCREMENTA EL CONSECUTIVO EN 1
					cC = CorrespondenciaConsecutivo.objects.get(prefijo_id =request.DATA['prefijo_id'], ano =int(fecha[:4]))	
					consecutivo = cC.numero	
					request.DATA['consecutivo_carta'] = consecutivo				


				try:
					usuarioDestinoFuncionario = Funcionario.objects.get(persona_id = usuarioDestino.persona.id
																	 ,empresa_id = usuarioDestino.empresa.id)
					cargo = usuarioDestinoFuncionario.cargo.nombre
				except Funcionario.DoesNotExist:
					cargo = 'No registra'


				contrato_id = request.DATA['contrato_id']
				try:
					contrato = Contrato.objects.get(pk = contrato_id)
					convenio = ''

					if contrato.mcontrato is not None:
						convenio = ' - Convenio '+str(contrato.mcontrato.nombre)

					referencia = contrato.nombre+'  No '+str(contrato.numero)+convenio;

				except Contrato.DoesNotExist:
					pass		

				insert_list_logs = []

				cE= CorrespondenciaEnviada(	consecutivo = consecutivo ,fechaEnvio = fecha,anoEnvio = int(fecha[:4])
									,asunto = unicode('Recomendación de imposición de multa por incumplimiento de obligaciones contractuales.', 'utf-8')	
									,referencia = referencia  
									,grupoSinin = 1
									,persona_destino = usuarioDestino.persona.nombres+' '+usuarioDestino.persona.apellidos
									,cargo_persona = cargo
									,direccion = usuarioDestino.empresa.direccion	
									,telefono = usuarioDestino.persona.telefono
									,empresa_destino = usuarioDestino.empresa.nombre
									,contenido = request.DATA['contenido']
									,contenidoHtml = request.DATA['contenidoHtml']
									,clausula_afectada = request.DATA['clausula_afectada']
									,clausula_afectadaHtml = request.DATA['clausula_afectadaHtml']
									,privado = 0, anulado = 0
									,ciudad_id = request.DATA['ciudad_id']
									,municipioEmpresa_id = request.DATA['ciudad_destinatario_id']
									,prefijo_id = request.DATA['prefijo_id']
									,usuarioSolicitante_id = usuarioActual
									,empresa_id = empresaIdActual
									,firma_id = request.DATA['firmaSolicitud_id']
					)

				cE.save()
				# SE REGISTRA EL LOG DE LA CORRESPONDENCIA ENNVIADA
				insert_list_logs.append(Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='correspondencia.CorrespondenciaEnviada',id_manipulado=cE.id))
				if not empresaConsecutivoDigitado:
					cC.numero = cC.numero+1
					cC.save()
		
				radicado = CorrespondenciaRadicado.objects.get(empresa_id = usuarioDestino.empresa.id, ano = d.year)
				cR = CorrespondenciaRecibida(
					 radicado = radicado.numero
					,fechaRecibida = fecha
					, anoRecibida = int(fecha[:4]) 
					,remitente = p.nombres+' '+p.apellidos
					,asunto = unicode('Recomendación de imposición de multa por incumplimiento de obligaciones contractuales.', 'utf-8')
					,privado = 0 # carta privada
					,correspondenciaEnviada_id = cE.id
					,empresa_id = usuarioDestino.empresa.id
					,usuarioSolicitante_id = destinatario
					)
				cR.save()
				# SE REGISTRA EL LOG DE LA CORRESPONDENCIA RECIBIDA 
				insert_list_logs.append(Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='correspondencia_recibida.CorrespondenciaRecibida',id_manipulado= cR.id ))
				radicado.numero = radicado.numero+1
				radicado.save()

				cRa=CorrespondenciaRecibidaAsignada(
					correspondenciaRecibida_id = cR.id
					,usuario_id = destinatario
					,fechaAsignacion = d
					,estado_id = 33# estado de la correspondencia por revisar
					,respuesta_id = None, copia = 0 # boleano false porque no es copia
					)
				cRa.save()
				# SE REGISTRA EL LOG DE LA CORRESPONDENCIA A QUIEN SE LE ASIGNA 
				insert_list_logs.append(Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='correspondencia_recibida.CorrespondenciaRecibidaAsignada',id_manipulado= cRa.id ))
	
				request.DATA['correspondenciasolicita_id'] = cE.id
				serializer = SolicitudSerializer(data=request.DATA,context={'request': request})
				# print "debug3"
				if serializer.is_valid():

					try:
						
						serializer.save(contrato_id = contrato_id , correspondenciasolicita_id = cE.id )
						# print "debug4"
						empresaPropietaria = SolicitudEmpresa(propietario = True ,empresa_id = empresaIdActual ,solicitud_id = serializer.data["id"])
						empresaPropietaria.save()
						

						if int(usuarioDestino.empresa.id) != int(empresaIdActual):

							empresaDestino = SolicitudEmpresa(propietario = False ,empresa_id = usuarioDestino.empresa.id ,solicitud_id = serializer.data["id"])
							empresaDestino.save()

						es=EstadoMulta()
						# print "debug5"
						solHistorial = SolicitudHistorial(soporte = None
											,comentarios = unicode('Elaboración de la solicitud de multa' , 'utf-8')
											,estado_id = es.elaborada
											,solicitud_id = serializer.data["id"]
											,usuario_id = usuarioActual
							)
						solHistorial.save()
						# print "debug6"
						ListEventosId = request.DATA['listado_eventos_asignados_id'] if 'listado_eventos_asignados_id' in request.DATA else None;
						ListEventosDia = request.DATA['listado_eventos_asignados_dia'] if 'listado_eventos_asignados_dia' in request.DATA else None;
						# print "debug7"

						# print str(request.DATA['listado_eventos_asignados_id'])+" == listado eventos id"
						if ListEventosId:
							ListEventosId = ListEventosId.split(',')
							ListEventosDia = ListEventosDia.split(',')
							if len(ListEventosId)>0:
								# print len(ListEventosId)
								insert_list = []
								for index in range(len(ListEventosId)):
									insert_list.append(SolicitudEvento(evento_id = ListEventosId[index]
																		, numeroimcumplimiento = ListEventosDia[index] 
																		,solicitud_id = serializer.data["id"])
														)

								SolicitudEvento.objects.bulk_create(insert_list)

						if insert_list_logs:
							Logs.objects.bulk_create(insert_list_logs)

						transaction.savepoint_commit(sid)
						return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
							'data':serializer.data},status=status.HTTP_201_CREATED)
					except Exception as e:
						print(e)
						transaction.savepoint_rollback(sid)
						return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
						'data':''},status=status.HTTP_400_BAD_REQUEST)
					
				else:
					print(serializer.errors)
					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "fechaDiligencia" in serializer.errors:
						mensaje = serializer.errors["fechaDiligencia"][0]
					elif "valorSolicitado" in serializer.errors:
						mensaje = serializer.errors["valorSolicitado"][0]
					else:
						mensaje = 'datos requeridos no fueron recibidos'
					transaction.savepoint_rollback(sid)
					return Response({'message': mensaje ,'success':'fail', 'data':''},status=status.HTTP_400_BAD_REQUEST)
			
			except CorrespondenciaConsecutivo.DoesNotExist:
				transaction.savepoint_rollback(sid)						
				mensaje='No existe el (consecutivo ó radicado) para la fecha de envio solicitada.'
				return Response({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)

			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

def generateSolicitud(request):
	if request.method == 'POST':
		sid = transaction.savepoint()
		try:
			es=EstadoMulta()
			lista=request.POST['_content']
			respuesta= json.loads(lista)
			# funcionario que firma la diligencia
			funcionario_id = respuesta['funcionario']
			# fecha de diligencia
			fechaDiligencia = respuesta['fechaDiligencia']
			# id de la carta de solicitud
			solicitud_id = respuesta['solicitud']
			# usuario que hace la peticion o transaccion
			usuario_id = respuesta['usuario']

			empresa = request.user.usuario.empresa.id


			try:
				consecutivo_solicitud = SolicitudConsecutivo.objects.get(empresa_id = empresa);
			except SolicitudConsecutivo.DoesNotExist:
				return JsonResponse({'message':unicode(' Su compañia no ha creado los consecutivos de las multas', 'utf-8'),'status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)
			

			consecutivo_multa = consecutivo_solicitud.consecutivo

			solicitud = Solicitud.objects.get(pk = solicitud_id)
			solicitud.fechaDiligencia = fechaDiligencia
			solicitud.firmaImposicion_id = funcionario_id
			solicitud.consecutivo = consecutivo_multa
			solicitud.save()

			consecutivo_solicitud.consecutivo = (int(consecutivo_multa)+1)
			consecutivo_solicitud.save()

			solicitud_historial = SolicitudHistorial(solicitud_id = solicitud.id
													,estado_id = es.generada
													,usuario_id = usuario_id
													,comentarios = unicode('Generación de multa', 'utf-8')  )
			solicitud_historial.save()

			logs_model = Logs(usuario_id=request.user.usuario.id
									,accion=Acciones.accion_crear
									,nombre_modelo='multa.SolicitudHistorial'
									,id_manipulado=solicitud_historial.id)

			logs_model.save()

			transaction.savepoint_commit(sid)
			return JsonResponse({'message': 'El registro ha sido guardado exitosamente consecutivo No '+str(consecutivo_multa),'success':'ok','data': ''})
		except Exception as e:
			transaction.savepoint_rollback(sid)
			print(e)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		

@transaction.atomic
def registerCodigoOF(request):
	if request.method == 'POST':
		sid = transaction.savepoint()
		try:
			es=EstadoMulta()
			lista=request.POST['_content']
			respuesta= json.loads(lista)
			# codigo OF 
			codigo_of = respuesta['codigoOF']
			# fecha de diligencia
			fechaDiligencia = respuesta['fechaDiligencia']
			# id de la carta de solicitud
			solicitud_id = respuesta['solicitud']
			# usuario que hace la peticion o transaccion
			usuario_id = respuesta['usuario']

			solicitud = Solicitud.objects.get(pk = solicitud_id)
			solicitud.codigoOF = codigo_of
			solicitud.save()

			solicitud_historial = SolicitudHistorial(solicitud_id = solicitud.id
													,estado_id = es.pendiente_contabilizacion
													,usuario_id = usuario_id
													,comentarios = 'Registro del codigo OF')
			solicitud_historial.save()
			transaction.savepoint_commit(sid)
			return JsonResponse({'message': 'El registro ha sido guardado exitosamente','success':'ok','data': ''})
		except Exception as e:
			print(e)
			transaction.savepoint_rollback(sid)
			return JsonResponse({'message': 'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@transaction.atomic
def registerCodigoReferencia(request):
	if request.method == 'POST':
		sid = transaction.savepoint()
		try:
			es=EstadoMulta()
			lista=request.POST['_content']
			respuesta= json.loads(lista)
			# codigo OF 
			codigo = respuesta['codigo']
			# fecha de diligencia
			fechaDiligencia = respuesta['fechaDiligencia']
			# id de la carta de solicitud
			solicitud_id = respuesta['solicitud']
			# usuario que hace la peticion o transaccion
			usuario_id = respuesta['usuario'] if 'usuario' in respuesta else request.user.usuario.id;

			solicitud = Solicitud.objects.get(pk = solicitud_id)
			solicitud.codigoReferencia = codigo
			solicitud.save()

			solicitud_historial = SolicitudHistorial(solicitud_id = solicitud.id
													,estado_id = es.contabilizada
													,usuario_id = usuario_id
													,comentarios = unicode('Registro del codigo de referencia', 'utf-8')  )
			solicitud_historial.save()
			transaction.savepoint_commit(sid)
			return JsonResponse({'message': 'El registro ha sido guardado exitosamente','success':'ok','data': ''})
		except Exception as e:
			print(e)
			transaction.savepoint_rollback(sid)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)				

def updateValorImpuesto(request):
	if request.method == 'POST':
		sid = transaction.savepoint()
		try:

			lista=request.POST['_content']
			respuesta= json.loads(lista)

			usuario = respuesta['usuario'] if 'usuario' in respuesta else request.user.usuario.id;
			solicitud_id = respuesta['solicitud'] if 'solicitud' in respuesta else None;
			valor = respuesta['valor'] if 'valor' in respuesta else None;

			if solicitud_id is None or valor is None:
				# print solicitud_id
				return JsonResponse({'message': 'datos requeridos no fueron recibidos','success':'fail','data': ''},status=status.HTTP_400_BAD_REQUEST)

			solicitud = Solicitud.objects.get(pk = solicitud_id)
			solicitud.valorImpuesto = valor
			solicitud.save()

			transaction.savepoint_commit(sid)
			return JsonResponse({'message': 'El registro ha sido guardado exitosamente','success':'ok','data': ''})

		except Exception as e:
			transaction.savepoint_rollback(sid)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)				


# se usa para consulta paginada
class SolicitudEmpresaLiteSerializer(serializers.HyperlinkedModelSerializer):

	solicitud = SolicitudLiteSerializer(read_only = True)
	totalMultaContrato = serializers.SerializerMethodField()
	estado = serializers.SerializerMethodField()
	fechasolicitud = serializers.SerializerMethodField()
	totalDescargos = serializers.SerializerMethodField()
	fechaelaboracion = serializers.SerializerMethodField()

	class Meta:
		model = SolicitudEmpresa
		fields = ('id' , 'solicitud' , 'propietario' ,'totalMultaContrato' ,'estado' ,'fechasolicitud' ,'totalDescargos' , 'fechaelaboracion')


	def get_totalMultaContrato(self, obj):
		es=EstadoMulta()
		# return Solicitud.objects.filter(contrato_id = obj.solicitud.contrato_id).count()
		total_multa = 0
		idSolicitud = Solicitud.objects.filter(contrato_id = obj.solicitud.contrato_id).values_list("id")

		queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = idSolicitud).values('solicitud_id').annotate(id=Max('id')).values("id")

		if queryHistorial:
			qset = ( Q(pk__in = queryHistorial) )

			qset = qset & ( Q(estado_id__in = [es.contabilizada,es.pendiente_contabilizacion,es.confirmada] ))

			total_multa = SolicitudHistorial.objects.filter(qset).count()	

		return str(total_multa);

	def get_estado(self, obj):
		try:
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id).latest('id')
			lista = {'id': datos.estado_id, 'nombre': datos.estado.nombre}
			return lista
		except SolicitudHistorial.DoesNotExist:
			lista = {'id': 0, 'nombre': ''}
			return lista

	def get_fechaelaboracion(self, obj):
		try:
			es=EstadoMulta()
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id , estado_id = es.elaborada).latest('id')
			return datos.fecha
		except Exception as e:
			return ''

	def get_fechasolicitud(self, obj):
		try:
			# print obj.solicitud_id
			es=EstadoMulta()
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id , estado_id = es.solicitada).latest('id')
			return datos.fecha
		except Exception as e:
			return ''

	def get_totalDescargos(self, obj):
		try:
			return SolicitudApelacion.objects.filter(solicitud_id = obj.solicitud_id).count()
		except Exception as e:
			return 0

# SE USA PARA LAS SOLICITUDES ELABORADAS
class SolicitudEmpresaSuperLiteSerializer(serializers.HyperlinkedModelSerializer):

	solicitud = SolicitudSuperLiteSerializer(read_only = True)	
	totalMultaContrato = serializers.SerializerMethodField()
	fechaelaboracion = serializers.SerializerMethodField()

	class Meta:
		model = SolicitudEmpresa
		fields = ('id' , 'solicitud' , 'propietario' ,'totalMultaContrato' , 'fechaelaboracion')


	def get_totalMultaContrato(self, obj):
		es=EstadoMulta()
		# return Solicitud.objects.filter(contrato_id = obj.solicitud.contrato_id).count()
		total_multa = 0
		idSolicitud = Solicitud.objects.filter(contrato_id = obj.solicitud.contrato_id).values_list("id")

		queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = idSolicitud).values('solicitud_id').annotate(id=Max('id')).values("id")

		if queryHistorial:
			qset = ( Q(pk__in = queryHistorial) )

			qset = qset & ( Q(estado_id__in = [es.contabilizada,es.pendiente_contabilizacion,es.confirmada] ))

			total_multa = SolicitudHistorial.objects.filter(qset).count()	

		return str(total_multa);

	def get_fechaelaboracion(self, obj):
		try:
			es=EstadoMulta()
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id , estado_id = es.elaborada).latest('id')
			return datos.fecha
		except Exception as e:
			return ''

# SE USA PARA LA CONSULTA EN GENERAL DE TODAS LAS MULTAS
class SolicitudEmpresaGeneralLiteSerializer(serializers.HyperlinkedModelSerializer):

	solicitud = SolicitudSuperLiteSerializer(read_only = True)
	totalMultaContrato = serializers.SerializerMethodField()
	estado = serializers.SerializerMethodField()
	fechasolicitud = serializers.SerializerMethodField()
	totalDescargos = serializers.SerializerMethodField()

	class Meta:
		model = SolicitudEmpresa
		fields = ('id' , 'solicitud' , 'propietario' ,'totalMultaContrato' ,'estado' ,'fechasolicitud' ,'totalDescargos' )


	def get_totalMultaContrato(self, obj):
		es=EstadoMulta()
		# return Solicitud.objects.filter(contrato_id = obj.solicitud.contrato_id).count()
		total_multa = 0
		idSolicitud = Solicitud.objects.filter(contrato_id = obj.solicitud.contrato_id).values_list("id")

		queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = idSolicitud).values('solicitud_id').annotate(id=Max('id')).values("id")

		if queryHistorial:
			qset = ( Q(pk__in = queryHistorial) )

			qset = qset & ( Q(estado_id__in = [es.contabilizada,es.pendiente_contabilizacion,es.confirmada] ))

			total_multa = SolicitudHistorial.objects.filter(qset).count()	

		return str(total_multa);

	def get_estado(self, obj):
		try:
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id).latest('id')
			lista = {'id': datos.estado_id, 'nombre': datos.estado.nombre}
			return lista
		except SolicitudHistorial.DoesNotExist:
			lista = {'id': 0, 'nombre': ''}
			return lista

	def get_fechasolicitud(self, obj):
		try:
			# print obj.solicitud_id
			es=EstadoMulta()
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id , estado_id = es.solicitada).latest('id')
			return datos.fecha
		except Exception as e:
			return ''

	def get_totalDescargos(self, obj):
		try:
			return SolicitudApelacion.objects.filter(solicitud_id = obj.solicitud_id).count()
		except Exception as e:
			return 0


class SolicitudEmpresaSerializer(serializers.HyperlinkedModelSerializer):

	empresa = EmpresaLiteSerializer(read_only = True)
	empresa_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Empresa.objects.all())

	solicitud = SolicitudSerializer(read_only = True)
	solicitud_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Solicitud.objects.all())

	totalMultaContrato = serializers.SerializerMethodField()
	estado = serializers.SerializerMethodField()
	fechasolicitud = serializers.SerializerMethodField()
	totalDescargos = serializers.SerializerMethodField()
	fechaelaboracion = serializers.SerializerMethodField()

	class Meta:
		model = SolicitudEmpresa
		fields = ('id' ,'empresa' ,'empresa_id' ,'solicitud' ,'solicitud_id' ,'propietario' ,'totalMultaContrato' ,'estado' ,'fechasolicitud' ,'totalDescargos' , 'fechaelaboracion')

		validators=[
				serializers.UniqueTogetherValidator(
				queryset=model.objects.all(),
				fields=( 'empresa_id' , 'solicitud_id' ),
				message=('El nombre del evento  no puede  estar repetido en el conjunto.')
				)
				]
	def get_totalMultaContrato(self, obj):
		es=EstadoMulta()
		# return Solicitud.objects.filter(contrato_id = obj.solicitud.contrato_id).count()
		total_multa = 0
		idSolicitud = Solicitud.objects.filter(contrato_id = obj.solicitud.contrato_id).values_list("id")

		queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = idSolicitud).values('solicitud_id').annotate(id=Max('id')).values("id")

		if queryHistorial:
			qset = ( Q(pk__in = queryHistorial) )

			qset = qset & ( Q(estado_id__in = [es.contabilizada,es.pendiente_contabilizacion,es.confirmada] ))

			total_multa = SolicitudHistorial.objects.filter(qset).count()	

		return str(total_multa);

	def get_estado(self, obj):
		try:
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id).latest('id')
			lista = {'id': datos.estado_id, 'nombre': datos.estado.nombre}
			return lista
		except SolicitudHistorial.DoesNotExist:
			lista = {'id': 0, 'nombre': ''}
			return lista

	def get_fechaelaboracion(self, obj):
		try:
			es=EstadoMulta()
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id , estado_id = es.elaborada).latest('id')
			return datos.fecha
		except Exception as e:
			return ''

	def get_fechasolicitud(self, obj):
		try:
			# print obj.solicitud_id
			es=EstadoMulta()
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id , estado_id = es.solicitada).latest('id')
			return datos.fecha
		except Exception as e:
			return ''

	def get_totalDescargos(self, obj):
		try:
			return SolicitudApelacion.objects.filter(solicitud_id = obj.solicitud_id).count()
		except Exception as e:
			return 0

class SolicitudEmpresaViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de informacion tecnica del proyectos ,
	puede utilizar el parametro (campo , proyecto) a traves del cual, se podra buscar por cada uno de estos filtros
	<br>campo = [numero] <br>
	proyecto = [numero].
	"""
	model = SolicitudEmpresa
	queryset = model.objects.all()
	serializer_class = SolicitudEmpresaSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(SolicitudEmpresaViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			# VARIABLES SOLICITUD EMPRESA
			# propietario es una variable boolean
			propietario = self.request.query_params.get('propietario', None)
			empresa_id = self.request.query_params.get('empresa', request.user.usuario.empresa.id)
			solicitud_id = self.request.query_params.get('solicitud', None)			
			
			# macro contrato 
			macro_contrato_id = self.request.query_params.get('macro_contrato', None)
			# estado de los parametros del filtro	
			consecutivo = self.request.query_params.get('consecutivo', None)
			fecha_desde = self.request.query_params.get('fecha_desde', None)
			fecha_hasta = self.request.query_params.get('fecha_hasta', None)
			numero_contrato_obra = self.request.query_params.get('numero_contrato_obra', None)
			contratista_id = self.request.query_params.get('contratista', None)			
			solicitante_id = self.request.query_params.get('solicitante', None)
			contrato_id = self.request.query_params.get('contrato_id', None)
			
			# VARIABLES SOLICITUD
			correspondenciasolicita_id = self.request.query_params.get('correspondenciasolicita', None)
			correspondenciadescargo_id = self.request.query_params.get('correspondenciadescargo', None)
			firmaImposicion_id = self.request.query_params.get('firmaImposicion', None)

			# VARIABLES HISTORIAL
			estado_id = self.request.query_params.get('estado', None)

			# tipos de solicitudes
			solicitudes_elaboradas = self.request.query_params.get('solicitudes_elaboradas', None)
			solicitudes_solicitadas = self.request.query_params.get('solicitudes_solicitadas', None)
			solicitudes_consulta = self.request.query_params.get('solicitudes_consulta', None)
			es=EstadoMulta()

			qset=(~Q(id=0))
			# VARIABLES SOLICITUD EMPRESA
			if dato:
				qset = qset & ( Q(solicitud__contrato__numero__icontains = dato) |
							Q(solicitud__consecutivo__icontains = dato) |
							Q(solicitud__contrato__nombre__icontains = dato))	
			if propietario:

				if solicitudes_consulta is None or solicitudes_consulta == "":
					qset = qset & ( Q(propietario = propietario) )
			
			if empresa_id:
				qset = qset & ( Q(empresa_id = empresa_id) )
			if solicitud_id:
				qset = qset & ( Q(solicitud_id = solicitud_id) )

			# VARIABLES SOLICITUD
			if correspondenciasolicita_id:
				qset = qset & ( Q(solicitud__correspondenciasolicita_id = correspondenciasolicita_id) )		

			if correspondenciadescargo_id:
				qset = qset & ( Q(solicitud__correspondenciadescargo_id = correspondenciadescargo_id) )
			if firmaImposicion_id:
				qset = qset & ( Q(solicitud__firmaImposicion_id = firmaImposicion_id) )

			if macro_contrato_id:
				qset = qset & ( Q(solicitud__contrato__mcontrato_id = macro_contrato_id) )
			if consecutivo:
				qset = qset & ( Q(solicitud__consecutivo__icontains = consecutivo) )
			if numero_contrato_obra:
				qset = qset &( Q(solicitud__contrato__numero__icontains = numero_contrato_obra) )
			if contratista_id:
				qset = qset &( Q(solicitud__contrato__contratista__id = contratista_id) )
			if contrato_id:
				qset = qset &( Q(solicitud__contrato__id = contrato_id) )	

			if fecha_desde or fecha_hasta:

				if (fecha_desde and (fecha_hasta is None) ):
					qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__gte= fecha_desde) 
									& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )
				
				if ( (fecha_desde is None) and fecha_hasta ):
					qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__lte=fecha_hasta)  
									& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )

				if (fecha_desde and fecha_hasta):	
					qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__range = (fecha_desde , fecha_hasta))
									& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )

			
			queryset = self.model.objects.filter(qset).distinct()

			idSolicitud = []
			for i in queryset:
				idSolicitud.append(i.solicitud_id)
				# print i.solicitud_id

			queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = idSolicitud).values('solicitud_id').annotate(id=Max('id')).values("id")		

			if queryset:

				qset = ( Q(solicitud__fk_Solicitud_SolicitudHistorial__id__in = queryHistorial) )

				# VARIABLES HISTORIAL
				if estado_id:
					qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = estado_id))	

				if solicitante_id:
					qset = qset & ( Q(empresa_id = solicitante_id) & Q(propietario = True) )
				else:
					qset = qset & ( Q(empresa_id = empresa_id) )	

				if solicitudes_consulta:
					qset = qset & (~Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id__in = [
							es.elaborada,
						]))
					queryset = self.model.objects.filter(qset).distinct()
				else:
					queryset = self.model.objects.filter(qset).distinct()

			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					# serializer = self.get_serializer(page,many=True)
					serializer_context = {
						'request': request
					}

					if solicitudes_elaboradas:
						serializer = SolicitudEmpresaSuperLiteSerializer(page,many=True,context=serializer_context).data						
					elif solicitudes_solicitadas:
						serializer = SolicitudEmpresaLiteSerializer(page,many=True,context=serializer_context).data
					elif solicitudes_consulta:
						serializer = SolicitudEmpresaGeneralLiteSerializer(page,many=True,context=serializer_context).data
					else:
						serializer = SolicitudEmpresaLiteSerializer(page,many=True,context=serializer_context).data

					
					return self.get_paginated_response({'message':'','success':'ok','data': { 
							'solicitudes' : serializer
						}})

			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			try:
				serializer = SolicitudSerializer(data=request.DATA,context={'request': request})
				if serializer.is_valid():
					serializer.save(conjunto_id =  request.DATA['conjunto_id'])
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:

					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "nombre" in serializer.errors:
						mensaje = serializer.errors["nombre"][0]
					elif "valor" in serializer.errors:
						mensaje = serializer.errors["valor"][0]
					else:
						mensaje = 'datos requeridos no fueron recibidos'
					return Response({'message': mensaje ,'success':'fail', 'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				print(e)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
			  		'data':''},status=status.HTTP_400_BAD_REQUEST)

	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = SolicitudSerializer(instance,data=request.DATA,context={'request': request},partial=partial)
				if serializer.is_valid():
					self.perform_update(serializer)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

class SolicitudEventoSerializer(serializers.HyperlinkedModelSerializer):

	evento = EventoSerializer(read_only = True)
	evento_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Evento.objects.all())

	solicitud = SolicitudSerializer(read_only = True)
	solicitud_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Solicitud.objects.all())

	class Meta:
		model = SolicitudEvento
		fields = ('id' , 'evento' , 'evento_id' , 'solicitud' , 'solicitud_id' , 'numeroimcumplimiento' )


class SolicitudEventoViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de informacion tecnica del proyectos ,
	puede utilizar el parametro (campo , proyecto) a traves del cual, se podra buscar por cada uno de estos filtros
	<br>campo = [numero] <br>
	proyecto = [numero].
	"""
	model = SolicitudEvento
	queryset = model.objects.all()
	serializer_class = SolicitudEventoSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(SolicitudEventoViewSet, self).get_queryset()
			nombre = self.request.query_params.get('dato', None)
			conjunto = self.request.query_params.get('conjunto', None)

			qset=None
			if nombre:
				qset = ( Q(nombre__icontains = nombre) )
			if conjunto:
				if qset:
					qset = ( Q(conjunto_id = conjunto) )
				else:
					qset = qset & ( Q(conjunto_id = conjunto) )

			if qset:
				queryset = self.model.objects.filter(qset)
			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})

			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			try:
				serializer = SolicitudEventoSerializer(data=request.DATA,context={'request': request})
				if serializer.is_valid():
					serializer.save(conjunto_id =  request.DATA['conjunto_id'])
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:

					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "nombre" in serializer.errors:
						mensaje = serializer.errors["nombre"][0]
					elif "valor" in serializer.errors:
						mensaje = serializer.errors["valor"][0]
					else:
						mensaje = 'datos requeridos no fueron recibidos'
					return Response({'message': mensaje ,'success':'fail', 'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				print(e)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
			  		'data':''},status=status.HTTP_400_BAD_REQUEST)

	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = SolicitudEventoSerializer(instance,data=request.DATA,context={'request': request},partial=partial)
				if serializer.is_valid():
					self.perform_update(serializer)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

class SolicitudHistorialSerializer(serializers.HyperlinkedModelSerializer):

	estado = EstadoSerializer(read_only = True)
	estado_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Estado.objects.all())

	solicitud = SolicitudSerializer(read_only = True)
	solicitud_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Solicitud.objects.all())

	class Meta:
		model = SolicitudHistorial
		fields = ('id' , 'fecha' , 'solicitud' , 'solicitud_id' , 'estado' , 'estado_id' , 'comentarios', 'soporte')

class SolicitudHistorialViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de informacion tecnica del proyectos ,
	puede utilizar el parametro (campo , proyecto) a traves del cual, se podra buscar por cada uno de estos filtros
	<br>campo = [numero] <br>
	proyecto = [numero].
	"""
	model = SolicitudHistorial
	queryset = model.objects.all()
	serializer_class = SolicitudHistorialSerializer
	nombre_modulo = 'multa.SolicitudHistorialViewSet'

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(SolicitudHistorialViewSet, self).get_queryset()
			fecha = self.request.query_params.get('fecha', None)
			estado_id = self.request.query_params.get('estado', None)
			solicitud_id = self.request.query_params.get('solicitud', None)
			usuario_id = self.request.query_params.get('usuario', None)

			qset=(~Q(id=0))
			if fecha:
				qset = qset & ( Q(fecha = fecha) )
			if estado_id:
				qset = qset & ( Q(estado_id = estado_id) )
			if solicitud_id:
				qset = qset & ( Q(solicitud_id = solicitud_id) )
			if usuario_id:
				qset = qset & ( Q(usuario_id = usuario_id) )

			queryset = self.model.objects.filter(qset)
			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})

			serializer = self.get_serializer(queryset,many=True)
			return Response({'message':'','success':'ok','data':serializer.data})
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			try:
				serializer = SolicitudHistorialSerializer(data=request.DATA,context={'request': request})
				if serializer.is_valid():
					es = EstadoMulta()
					estado_id = request.DATA['estado_id']
					archivo = request.FILES['soporte'] if 'soporte' in request.FILES else '';

					if archivo!='':

						filename, file_extension = os.path.splitext(archivo.name)

						if ( int(es.anulada) == int(estado_id) ):

							if (file_extension.lower()!=".pdf" and file_extension.lower()!=".zip"):
								return Response({'message': 'Solo se permiten archivos (PDF y ZIP).' ,'success':'fail',
										'data':''},status=status.HTTP_400_BAD_REQUEST)

						else:

							if file_extension.lower()!=".pdf":
								return Response({'message': 'Solo se permiten archivos (PDF).' ,'success':'fail',
										'data':''},status=status.HTTP_400_BAD_REQUEST)

						t = datetime.datetime.now()
					
						archivo.name = str(request.user.usuario.empresa.id)+'-'+str(request.user.usuario.id)+'-'+str(t.year)+str(t.month)+str(t.day)+str(t.hour)+str(t.minute)+str(t.second)+str(t.microsecond)+str(file_extension)
						destino = archivo
					else:
						destino = ''
					if ((int(estado_id) == int(es.notificada_contratista) or int(estado_id) == int(es.contabilizada) or int(estado_id) == int(es.solicitada)) and (destino=='')):
						mensaje = 'Debe cargar un soporte'
						return Response({'message': mensaje ,'success':'fail', 'data':''},status=status.HTTP_400_BAD_REQUEST)


					serializer.save(solicitud_id =  request.DATA['solicitud_id'] 
									,estado_id =  estado_id
									,usuario_id = request.user.usuario.id
									,soporte = destino 
									)
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					print(serializer.errors)
					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "fecha" in serializer.errors:
						mensaje = serializer.errors["fecha"][0]
					else:
						mensaje = 'datos requeridos no fueron recibidos'
					return Response({'message': mensaje ,'success':'fail', 'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
			  		'data':''},status=status.HTTP_400_BAD_REQUEST)

	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = SolicitudHistorialSerializer(instance,data=request.DATA,context={'request': request},partial=partial)
				if serializer.is_valid():
					self.perform_update(serializer)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

#Api rest para SolicitudSoporte
class SolicitudSoporteSerializer(serializers.HyperlinkedModelSerializer):

	solicitud = SolicitudSerializer(read_only = True)
	solicitud_id = serializers.PrimaryKeyRelatedField(write_only=True , queryset=Solicitud.objects.all())

	class Meta:
		model = SolicitudSoporte
		fields=('id' ,'nombre' , 'solicitud' , 'solicitud_id' ,'soporte' ,'anulado' )


class SolicitudSoporteLiteSerializer(serializers.HyperlinkedModelSerializer):

	
	class Meta:
		model = SolicitudSoporte
		fields=('id' ,'nombre' , 'soporte' ,'anulado' )


class SolicitudSoporteViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de fondos de proyectos , 
	puede utilizar el parametro (dato) a traves del cual puede consultar todos los fondos.
	"""
	model = SolicitudSoporte
	queryset = model.objects.all()
	serializer_class = SolicitudSoporteSerializer
	nombre_modulo = 'multas.soporte'

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(SolicitudSoporteViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			solicitudId = self.request.query_params.get('solicitud', None)
			# anulado = self.request.query_params.get('anulado', None)

			# qset = ( Q(anulado = 0) )
			qset = ( Q(anulado = 0) )
			if dato:
				qset = qset & ( Q(nombre__icontains=dato) )

			if solicitudId:
				qset = qset &( Q(solicitud_id = solicitudId) )
			
			queryset = self.model.objects.filter(qset)

			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			serializer_context = {
				'request': request,
			}


			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					return self.get_paginated_response({'message':'','success':'ok',
					'data':serializer.data})

			# serializer = self.get_serializer(queryset,many=True)
			serializer = SolicitudSoporteLiteSerializer(queryset,many=True,context=serializer_context)	

			return Response({'message':'','success':'ok','data':serializer.data})				
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()
			try:
				serializer = SolicitudSoporteSerializer(data=request.DATA,context={'request': request})

				archivos = request.FILES.getlist('soporte[]') if 'soporte[]' in request.FILES else None;

				if serializer.is_valid():
					sw = 0

					if archivos is None:

						return Response({'message': 'Solo se permiten archivos (WORD o PDF).' ,'success':'fail',
			 			'data':''},status=status.HTTP_400_BAD_REQUEST)

					for archivo in archivos:

						filename, file_extension = os.path.splitext(archivo.name)						

						t = datetime.datetime.now()
						nombre = filename

						archivo.name = str(request.user.usuario.empresa.id)+'-'+str(request.user.usuario.id)+'-'+str(t.year)+str(t.month)+str(t.day)+str(t.hour)+str(t.minute)+str(t.second)+str(t.microsecond)+str(file_extension)
						destino = archivo

						s= SolicitudSoporte(soporte = destino , nombre = nombre , anulado = 0 ,  solicitud_id = request.POST['solicitud_id'])
						s.save()

						# serializer.save(soporte = destino , nombre = nombre , anulado = 0 ,  solicitud_id = request.POST['solicitud_id'])	
						# print serializer.data['id']
						# SE REGISTRA LA CARGA DEL SOPORTE
						logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='multa.SolicitudSoporte',id_manipulado=s.id)
						logs_model.save()
					
					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok', 'data':serializer.data},status=status.HTTP_201_CREATED)										

				else:

					print(serializer.errors)
					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "nombre" in serializer.errors:
						mensaje = serializer.errors["nombre"][0]+" En el campo nombre"
					elif "anulado" in serializer.errors:
						mensaje = serializer.errors["anulado"][0]+" En el campo anulado"
					else:
						mensaje = 'datos requeridos no fueron recibidos'


					return Response({'message': mensaje ,'success':'fail',
						'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
			  		'data':''},status=status.HTTP_400_BAD_REQUEST)

@transaction.atomic
def destroySolicitudSoporte(request):
	if request.method == 'POST':
		sid = transaction.savepoint()
		try:
			lista=request.POST['_content']
			respuesta= json.loads(lista)
			myList = respuesta['soporte_id']
			SolicitudSoporte.objects.filter(id__in = myList).update(anulado = 1)


			insert_list = []
			for i in myList:
				insert_list.append(Logs(usuario_id=request.user.usuario.id
										,accion=Acciones.accion_borrar
										,nombre_modelo='multa.SolicitudSoporte'
										,id_manipulado=i)
										)
			if insert_list:
				Logs.objects.bulk_create(insert_list)

			transaction.savepoint_commit(sid)
			
			return JsonResponse({'message':'El registro se ha eliminado correctamente','success':'ok','data': ''})
		except Exception as e:
			print(e)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#Fin api rest para CorrespondenciaSoporte

# CREAR SOLICITUD DE LA MULTA
@transaction.atomic
def uploapCartaSolicitudSoporte(request):
	if request.method == 'POST':
		nombre_modulo= 'multa.uploapCartaSolicitudSoporte'
		sid = transaction.savepoint()
		try:
			usuarioActual = request.user.usuario.id
			archivo = request.FILES['soporte'] if 'soporte' in request.FILES else None;

			solicitud = request.POST['solicitud_id'] if 'solicitud_id' in request.POST else 0;
			correspondencia = request.POST['correspondencia_id'] if 'correspondencia_id' in request.POST else 0;	

			solicitud = Solicitud.objects.get(pk = solicitud)

			solicitud_soportes = SolicitudSoporte.objects.filter( solicitud_id = solicitud.id , anulado = False)

			if solicitud_soportes.count()==0:
				return JsonResponse({'message': 'Senor(a) usuario(a) se debe subir primero las pruebas de la solicitud.' ,'success':'fail',
						'data':''},status=status.HTTP_400_BAD_REQUEST)

			if archivo is None:
				return JsonResponse({'message': 'Solo se permiten archivos (PDF).' ,'success':'fail',
						'data':''},status=status.HTTP_400_BAD_REQUEST)
			
			filename, file_extension = os.path.splitext(archivo.name)

			if file_extension.lower()!=".pdf":
				return JsonResponse({'message': 'Solo se permiten archivos (PDF).' ,'success':'fail',
			 			'data':''},status=status.HTTP_400_BAD_REQUEST)


			t = datetime.datetime.now()
			
			# valida si seleccionan para cambiar el nombre del archivo
			if request.POST['nombre'] == '':
				nombre = filename
			else:	
				nombre = request.POST['nombre']

			archivo.name = str(request.user.usuario.empresa.id)+'-'+str(request.user.usuario.id)+'-'+str(t.year)+str(t.month)+str(t.day)+str(t.hour)+str(t.minute)+str(t.second)+str(file_extension)
			destino = archivo

			correspondenciaSoporte= CorrespondenciaSoporte(soporte = destino , nombre = nombre , anulado = 0 ,  correspondencia_id = correspondencia)
			correspondenciaSoporte.save()

			es=EstadoMulta()
			# print correspondenciaSoporte.id
			solHistorial = SolicitudHistorial(soporte = None
								,comentarios = 'Carta de solicitud'
								,estado_id = es.solicitada
								,solicitud_id = solicitud.id
								,usuario_id = usuarioActual
				)
			solHistorial.save()

			insert_list = []
			
			insert_list.append(Logs(usuario_id=request.user.usuario.id
										,accion=Acciones.accion_crear
										,nombre_modelo='correspondencia.CorrespondenciaEnviada.proyecto'
										,id_manipulado=correspondenciaSoporte.id))

			insert_list.append(Logs(usuario_id=request.user.usuario.id
									,accion=Acciones.accion_crear
									,nombre_modelo='multa.SolicitudHistorial'
									,id_manipulado=solHistorial.id))
									
			Logs.objects.bulk_create(insert_list)			

			
			# ENVIAR CORREOS DE LA SOLICITUD DE LA MULTA
			notificacion = Notificaciones()
			# se consultan todos los proyectos asociados al contrato
			proyectos = Proyecto.objects.filter(contrato = solicitud.contrato.id ).values_list('id')
			proyecto_funcionarios = Proyecto.funcionario.through.objects.filter(proyecto_id__in = proyectos , funcionario__notificaciones = notificacion.multas_solicitud_actualizacion_de_estados ).values_list('funcionario_id')


			transaction.savepoint_commit(sid)

			formato_fecha = "%Y-%m-%d"
			hoy = date.today()
			hoy = datetime.datetime.strptime(str(hoy), formato_fecha)

			funcionarios_notificar = Funcionario.objects.filter( pk__in = proyecto_funcionarios )

			for f in funcionarios_notificar:
				
				empresaId = f.empresa_id
				
				queryHistorial = SolicitudHistorial.objects.filter(solicitud_id = solicitud.id ).values('solicitud_id').annotate(id=Max('id')).values("id")
				
				qset = qset ( Q(solicitud__fk_Solicitud_SolicitudHistorial__id__in = queryHistorial) )

				es=EstadoMulta()
				qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id__in = [
						es.solicitada
					]))

				queryset = SolicitudEmpresa.objects.filter(qset)

				if queryset:

					correo_envio = ''
					correo_envio = f.persona.correo+';'
					row=1
					col=0
					contenido='<div style="background-color: #34353F; height:40px;"><h3><p style="Color:#FFFFFF;"><strong>SININ </strong>- <b>S</b>istema <b>In</b>tegral de <b>In</b>formacion</p></h3></div><br/><br/>'
					contenido+='Estimado usuario(a), <br/><br/>nos permitimos comunicarle que las siguientes multas han sido <strong >solicitadas</strong> <br/><br/>'
					contenido+='<br><br><table border=1>'							
					contenido+='<tr bgcolor="#DED9D9">'
					contenido+='<td width="5%" >Consecutivo carta</td>'
					contenido+='<td width="5%" >Fecha solicitud</td>'
					contenido+='<td width="20%" >Solicitante</td>'
					contenido+='<td width="10%" >No Contrato</td>'
					contenido+='<td width="35%" >Contrato</td>'
					contenido+='<td width="20%">Contratista</td>'
					contenido+='<td width="10%" >Valor multa</td>'
					contenido+='</tr>'

					for objecto in queryset:

						meses = ''
						fecha = ''
						solicitante = SolicitudEmpresa.objects.get(solicitud_id = objecto.solicitud.id )
		
						contenido+='<tr style="color:black;">'
						contenido+='<td> '+objecto.solicitud.correspondenciasolicita.consecutivo+'</td>'							
						contenido+='<td> '+hoy+'</td>'
						contenido+='<td> '+solicitante.empresa.nombre+'</td>';
						contenido+='<td> '+objecto.solicitud.contrato.numero+'</td>'
						contenido+='<td> '+objecto.solicitud.contrato.nombre+'</td>'							
						contenido+='<td> '+objecto.solicitud.contrato.contratista.nombre+'</td>'
						contenido+='<td> $ '+objecto.solicitud.valorImpuesto+'</td>'					
						contenido+='</tr>'

						row +=1
					
					contenido+='<br><br>Favor no responder este correo, es de uso informativo unicamente.<br/><br/>'
					contenido+='Gracias,<br/><br/><br/>'
					contenido+='Equipo SININ<br/>'
					contenido+='soporte@sinin.co<br/>'
					contenido+='<a href="http://52.42.43.115:8000/usuario/">SININ</a><br/>'

					mail = Mensaje(
						remitente=settings.REMITENTE,
						destinatario=correo_envio,
						asunto='Recomendación de imposición de multa por incumplimiento de obligaciones contractuales.',
						contenido=contenido,
						appLabel='multas',
						tieneAdjunto=True,
						adjunto='',
						copia=''
					)

					mail.Send()

			# FINALIZA EL ENVIO DE CORREOS 

			return JsonResponse({'message':'El registro ha sido guardado exitosamente','success':'ok','data': ''})
		
		except Proyecto.DoesNotExist:
			transaction.savepoint_rollback(sid)
			mensaje='No existe un proyecto asociado al contrato solicitado.'
			return JsonResponse({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)	

		except Exception as e:
			functions.toLog(e,nombre_modulo)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#Api rest para SolicitudApelacion
class SolicitudApelacionSerializer(serializers.HyperlinkedModelSerializer):

	solicitud = SolicitudSerializer(read_only = True , allow_null = True)
	solicitud_id = serializers.PrimaryKeyRelatedField(write_only=True , queryset=Solicitud.objects.all())
	totalPronunciamientos = serializers.SerializerMethodField()
	fechasolicitud = serializers.SerializerMethodField()

	class Meta:
		model = SolicitudApelacion
		fields=('id' ,'fecha' , 'comentarios' , 'fecha_transacion' , 'soporte' , 'solicitud' ,'solicitud_id' ,'totalPronunciamientos' , 'fechasolicitud')

	def get_totalPronunciamientos(self, obj):
		return SolicitudPronunciamiento.objects.filter(apelacion_id = obj.id).count()

	def get_fechasolicitud(self, obj):
		try:
			# print obj.solicitud_id
			es=EstadoMulta()
			datos = SolicitudHistorial.objects.filter(solicitud_id = obj.solicitud_id , estado_id = es.solicitada).latest('id')
			return datos.fecha
		except Exception as e:
			return ''

class SolicitudApelacionLiteSerializer(serializers.HyperlinkedModelSerializer):

	class Meta:
		model = SolicitudApelacion
		fields=('id' ,'fecha' , 'comentarios' , 'fecha_transacion' , 'soporte' )

class SolicitudApelacionViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de fondos de proyectos , 
	puede utilizar el parametro (dato) a traves del cual puede consultar todos los fondos.
	"""
	model = SolicitudApelacion
	queryset = model.objects.all()
	serializer_class = SolicitudApelacionSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(SolicitudApelacionViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			solicitudId = self.request.query_params.get('solicitud', None)
			empresaId = self.request.query_params.get('empresa', None)

			# empresa que hace la peticion
			empresa_actual = self.request.query_params.get('empresa_actual', request.user.usuario.empresa.id)

			# variables de busqueda
			macro_contrato_id = self.request.query_params.get('macro_contrato', None)
			contratista_id = self.request.query_params.get('contratista', None)
			numero_contrato_obra = self.request.query_params.get('numero_contrato_obra', None)
			solicitante_id = self.request.query_params.get('solicitante', None)
			estado_id = self.request.query_params.get('estado', None)
			fecha_desde = self.request.query_params.get('desde', None)
			fecha_hasta = self.request.query_params.get('hasta', None)
			
			qset=(~Q(id=0))
			if dato:
				qset = qset & ( Q(solicitud__consecutivo__icontains=dato) )

			if solicitudId:
				qset = qset &( Q(solicitud__id = solicitudId) )

			if empresaId:
				qset = qset &( Q(solicitud__fk_Solicitud_SolicitudEmpresa__id = empresaId) )

			# filtro de busqueda
			if macro_contrato_id:
				qset = qset & ( Q(solicitud__contrato__mcontrato__id = macro_contrato_id) )

			if contratista_id:
				qset = qset &( Q(solicitud__contrato__contratista__id = contratista_id) )

			if numero_contrato_obra:
				qset = qset &( Q(solicitud__contrato__numero__icontains = numero_contrato_obra) )

			if solicitante_id:
				qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudEmpresa__empresa__id = solicitante_id) )


			if fecha_desde or fecha_hasta:

				es=EstadoMulta()

				if (fecha_desde and (fecha_hasta is None) ):
					qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__gte= fecha_desde) 
									& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )
				
				if ( (fecha_desde is None) and fecha_hasta ):
					qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__lte=fecha_hasta)  
									& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )

				if (fecha_desde and fecha_hasta):	
					qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__range = (fecha_desde , fecha_hasta))
									& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )


			queryset = self.model.objects.filter(qset).distinct()	

			idSolicitud = []
			for i in queryset:
				idSolicitud.append(i.solicitud_id)
				# print i.solicitud_id

			queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = idSolicitud).values('solicitud_id').annotate(id=Max('id')).values("id")

			qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__id__in = queryHistorial) )

			if estado_id:
				# print "estado_id"
				qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = estado_id))
			
			queryset = self.model.objects.filter(qset)

			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:

					serializer_context = {
						'request': request
					}

					serializer = SolicitudApelacionSerializer(page,many=True,context=serializer_context).data

					qsetEstadosMultas= Estado.objects.filter(app = "multa")
					estadosMultasData = EstadoSerializer(qsetEstadosMultas,many=True).data

					qsetEmpresas= Empresa.objects.all()
					empresasData = EmpresaLiteSerializer(qsetEmpresas,many=True).data

					qsetMcontratos= EmpresaContrato.objects.filter(empresa_id = empresa_actual , contrato__tipo_contrato_id = 12)
					mcontratosData = EmpresaContratoLiteSerializer(qsetMcontratos,many=True).data

					return self.get_paginated_response({'message':'','success':'ok','data': { 
							'descargos' : serializer,
							'estados_solicitudes' : estadosMultasData,
							'solicitantes' : empresasData
							,'macro_contratos' : mcontratosData
						}})

			serializer = SolicitudApelacionLiteSerializer(queryset,many=True)

			return Response({'message':'','success':'ok',
				'data':serializer.data})				
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()
			try:
				serializer = SolicitudApelacionSerializer(data=request.DATA,context={'request': request})

				if serializer.is_valid():

					
					archivo = request.FILES['soporte']
					filename, file_extension = os.path.splitext(archivo.name)
					t = datetime.datetime.now()
				
					archivo.name = str(request.user.usuario.empresa.id)+'-'+str(request.user.usuario.id)+'-'+str(t.year)+str(t.month)+str(t.day)+str(t.hour)+str(t.minute)+str(t.second)+str(file_extension)
					destino = archivo


					serializer.save(solicitud_id =  request.DATA['solicitud_id'] , soporte = destino )

					# SE REGISTRA LA CARGA DEL SOPORTE
					logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='multa.SolicitudSoporte',id_manipulado=serializer.data['id'])
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:

					print(serializer.errors)
					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "nombre" in serializer.errors:
						mensaje = serializer.errors["nombre"][0]+" En el campo nombre"
					elif "anulado" in serializer.errors:
						mensaje = serializer.errors["anulado"][0]+" En el campo anulado"
					else:
						mensaje = 'datos requeridos no fueron recibidos'


					return Response({'message': mensaje ,'success':'fail',
			 			'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				print(e)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
			  		'data':''},status=status.HTTP_400_BAD_REQUEST)

#Api rest para SolicitudPronunciamiento
class SolicitudPronunciamientoSerializer(serializers.HyperlinkedModelSerializer):

	apelacion = SolicitudApelacionSerializer(read_only = True , allow_null = True)
	apelacion_id = serializers.PrimaryKeyRelatedField(write_only=True , queryset=SolicitudApelacion.objects.all())

	class Meta:
		model = SolicitudPronunciamiento
		fields=( 'id', 'apelacion', 'apelacion_id' , 'comentarios' , 'fecha_transacion' )

class SolicitudPronunciamientoLiteSerializer(serializers.HyperlinkedModelSerializer):

	class Meta:
		model = SolicitudPronunciamiento
		fields=( 'id', 'comentarios' , 'fecha_transacion' )

class SolicitudPronunciamientoViewSet(viewsets.ModelViewSet):
	"""
	Retorna una lista de fondos de proyectos , 
	puede utilizar el parametro (dato) a traves del cual puede consultar todos los fondos.
	"""
	model = SolicitudPronunciamiento
	queryset = model.objects.all()
	serializer_class = SolicitudPronunciamientoSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(SolicitudPronunciamientoViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			apelacion_id = self.request.query_params.get('apelacion', None)
			empresaId = self.request.query_params.get('empresa', None)
			
			qset=(~Q(id=0))

			if apelacion_id:
				qset = qset &( Q(apelacion_id = apelacion_id) )
			
			queryset = self.model.objects.filter(qset)

			#utilizar la variable ignorePagination para quitar la paginacion
			ignorePagination= self.request.query_params.get('ignorePagination',None)
			if ignorePagination is None:
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					return self.get_paginated_response({'message':'','success':'ok',
					'data':serializer.data})

			serializer = SolicitudPronunciamientoLiteSerializer(queryset,many=True)

			return Response({'message':'','success':'ok',
				'data':serializer.data})				
		except Exception as e:
			print(e)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()
			try:
				serializer = SolicitudPronunciamientoSerializer(data=request.DATA,context={'request': request})

				if serializer.is_valid():			

					serializer.save(apelacion_id =  request.DATA['apelacion_id'] )
					# SE REGISTRA LA TRANSACION
					logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='multa.SolicitudPronunciamiento',id_manipulado=serializer.data['id'])
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:

					print(serializer.errors)
					if "non_field_errors" in serializer.errors:
						mensaje = serializer.errors['non_field_errors']
					elif "nombre" in serializer.errors:
						mensaje = serializer.errors["nombre"][0]+" En el campo nombre"
					elif "anulado" in serializer.errors:
						mensaje = serializer.errors["anulado"][0]+" En el campo anulado"
					else:
						mensaje = 'datos requeridos no fueron recibidos'


					return Response({'message': mensaje ,'success':'fail',
						'data':''},status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				print(e)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)


# CARTA DE RESPUESTA A LOS DESCARGOS PRESENTADOS EN LAS MULTAS
@transaction.atomic
def create_respuesta_descargo(request):
	if request.method == 'POST':
		try:
			solicitud_id = request.POST['id'] if 'id' in request.POST else None;
			fecha = request.POST['fechaEnvio'] if 'fechaEnvio' in request.POST else None;
			grupoSinin = request.POST['grupoSinin'] if 'grupoSinin' in request.POST else None;
			
			usuarioActual = request.POST['usuarioSolicitante_id'] if 'usuarioSolicitante_id' in request.POST else request.user.usuario.id;

			u = Usuario.objects.get(pk = usuarioActual)
			p = Persona.objects.get(pk = u.persona_id)
			empresaIdActual = u.empresa.id
			empresaConsecutivoDigitado = u.empresa.consecutivoDigitado	
			# print "mendoza"
			if empresaConsecutivoDigitado:
				consecutivo = request.POST['consecutivo'] if 'consecutivo' in request.POST else None;
			else:
				# SE  INCREMENTA EL CONSECUTIVO EN 1
				cC = CorrespondenciaConsecutivo.objects.get(prefijo_id = request.POST['prefijo_id'] , ano = int(fecha[:4]))	
				consecutivo = cC.numero	

			# VALIDA QUE NO EXISTA LA CORRESPONDENCIA CON EL NUMERO DE CONSECUTIVO EN EL MISMO AÑO
			correspondencias = CorrespondenciaEnviada.objects.filter(empresa_id = empresaIdActual , anoEnvio = int(fecha[:4]) , consecutivo = consecutivo)
			if correspondencias:
				mensaje='El numero de consecutivo ('+str(consecutivo)+') no puede estas repetido en la misma anualidad.'
				return JsonResponse({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)

			# print "if correspondencias"
			if correspondencias.count()==0:
				d = datetime.datetime.now()
				myListDestinatarioCopiar = request.POST['destinatarioCopia'] if 'destinatarioCopia' in request.POST else None;
				sid = transaction.savepoint()

				try:
					insert_list = []
					solicitud = Solicitud.objects.get(pk = solicitud_id , fk_Solicitud_SolicitudEmpresa__propietario = False)

					contrato_id = solicitud.contrato.id

					try:
						contrato = Contrato.objects.get(pk = contrato_id)
						convenio = ''

						if contrato.mcontrato is not None:
							convenio = ' - Convenio '+str(contrato.mcontrato.nombre)

						referencia = contrato.nombre+'  No '+str(contrato.numero)+convenio;

					except Contrato.DoesNotExist:
						mensaje='No existe un contrato asociado a la multa para hacer una respuesta de descargo.'
						return JsonResponse({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)	


					if  int(grupoSinin)>0:
						# print "grupo sinin mayor que cero"
						destinatario = int(request.POST['destinatario'])
								
						try:
							usuarioDestinatario = Usuario.objects.get(pk = destinatario)
							usuarioDestinatarioFuncionario = Funcionario.objects.get(persona_id = usuarioDestinatario.persona.id
																			 ,empresa_id = usuarioDestinatario.empresa.id)
							cargo = usuarioDestinatarioFuncionario.cargo.nombre
						except Funcionario.DoesNotExist:
							cargo = 'No registra'
						correspondencia = CorrespondenciaEnviada( empresa_id = empresaIdActual
										,asunto = 'RESPUESTA A DESCARGOS POR IMPOSICION DE MULTA.'	
										,referencia = referencia	
										,fechaEnvio = fecha
										,anoEnvio = int(fecha[:4])
										,ciudad_id = request.POST['ciudad_id']
										,prefijo_id = request.POST['prefijo_id']
										,firma_id = request.POST['firma_id']
										,anulado = 0
										,consecutivo = consecutivo
										,grupoSinin = 1
										,persona_destino = usuarioDestinatario.persona.nombres+' '+usuarioDestinatario.persona.apellidos
										,cargo_persona = cargo
										,empresa_destino = usuarioDestinatario.empresa.nombre
										,direccion = usuarioDestinatario.empresa.direccion	
										,municipioEmpresa_id = request.POST['municipioEmpresa_id']
										,usuarioSolicitante_id = usuarioActual
										 )
						correspondencia.save()
						# SE REGISTRA EL LOG DE LA CORRESPONDENCIA ENNVIADA
						logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='correspondencia.CorrespondenciaEnviada',id_manipulado=correspondencia.id)
						logs_model.save()
						# print "radicado"
						radicado = CorrespondenciaConsecutivo.objects.get(empresa_id = usuarioDestinatario.empresa.id
															 ,tipo_id = tipoRadicado
															 ,ano = d.year
															 )

						cR = CorrespondenciaRecibida(
											 radicado = radicado.numero
											,fechaRecibida = fecha
											, anoRecibida = int(fecha[:4]) 
											,remitente = p.nombres+' '+p.apellidos
											,asunto = request.POST['asunto']
											,privado = 1 # carta privada
											,correspondenciaEnviada_id = correspondencia.id
											,empresa_id = usuarioDestinatario.empresa.id
											,usuarioSolicitante_id = usuarioActual
											)
						cR.save()
						# print "hernandez"
						# SE REGISTRA EL LOG DE LA CORRESPONDENCIA RECIBIDA 
						logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='correspondencia_recibida.CorrespondenciaRecibida',id_manipulado= cR.id )
						logs_model.save()

						radicado.numero = radicado.numero+1
						radicado.save()

						cRa=CorrespondenciaRecibidaAsignada(
											correspondenciaRecibida_id = cR.id
											,usuario_id = destinatario
											,fechaAsignacion = d
											,estado_id = 33# estado de la correspondencia por revisar
											,respuesta_id = None # false 
											,copia = 0 # boleano false porque no es copia
											)
						cRa.save()

						# SE REGISTRA EL LOG DE LA CORRESPONDENCIA A QUIEN SE LE ASIGNA 
						logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='correspondencia_recibida.CorrespondenciaRecibidaAsignada',id_manipulado= cRa.id )
						logs_model.save()		
						
					else:
						# print "grupo sinin menor que cero"
						correspondencia= CorrespondenciaEnviada( empresa_id = empresaIdActual
										,asunto = 'RESPUESTA A DESCARGOS POR IMPOSICION DE MULTA.'	
										,referencia = referencia	
										,fechaEnvio = fecha
										,anoEnvio = int(fecha[:4])
										,ciudad_id = request.POST['ciudad_id']
										,prefijo_id = request.POST['prefijo_id']
										,firma_id = request.POST['firma_id']
										,anulado = 0
										,consecutivo = consecutivo
										,grupoSinin = 0
										,persona_destino = request.POST["persona_destino"]
										,cargo_persona = request.POST["cargo_persona"]
										,empresa_destino = request.POST["empresa_destino"]
										,direccion = request.POST["direccion"]
										,municipioEmpresa_id = request.POST['municipioEmpresa_id']
										,usuarioSolicitante_id = usuarioActual
										 )
						correspondencia.save()
						# SE REGISTRA EL LOG DE LA CORRESPONDENCIA ENNVIADA
						logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='correspondencia.CorrespondenciaEnviada',id_manipulado=correspondencia.id)
						logs_model.save()
					# print "mendoza se registro la correspondencia"
					if not empresaConsecutivoDigitado:
						cC.numero = cC.numero+1
						cC.save()			
					if myListDestinatarioCopiar:
						correspondencia.usuario.add(*myListDestinatarioCopiar)

						# SE REGISTRA EL LOG DE LAS COPIAS QUE SE HACEN A LOS USUARIOS
						logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='correspondencia.CorrespondenciaEnviada.usuario',id_manipulado=correspondencia.id)
						logs_model.save()

					# INSERTAR PROYECTO

					try:			
						solicitud.correspondenciadescargo_id = correspondencia.id
						solicitud.save()

						solicitud_historial = SolicitudHistorial(solicitud_id = solicitud.id
													,estado_id = request.POST["estado_id"]
													,usuario_id = request.user.usuario.id
													,comentarios = unicode('Respuesta al descargo presentado', 'utf-8')  )
						solicitud_historial.save()

						transaction.savepoint_commit(sid)

						return JsonResponse({'message':'El registro ha sido guardado exitosamente','success':'ok',
							'data': {'correspondenciadescargo_id':correspondencia.id} },status=status.HTTP_201_CREATED)
					except Solicitud.DoesNotExist as e:
						transaction.savepoint_rollback(sid)						
						mensaje='No existe la multa para hacer una respuesta de descargo.'
						return JsonResponse({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)

						
				except CorrespondenciaConsecutivo.DoesNotExist as e:
					transaction.savepoint_rollback(sid)						
					mensaje='No existe el (consecutivo ó radicado) para la fecha de envio solicitada.'
					return JsonResponse({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)

				except Exception as e:
					transaction.savepoint_rollback(sid)
					print(e)
					mensaje='El consecutivo no puede  estar repetido en la misma anualidad de envio.'
					return JsonResponse({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)


			else:
				# print(serializer.errors)
				mensaje = 'datos requeridos no fueron recibidos'

				return JsonResponse({'message':mensaje,'success':'fail','data':''},status=status.HTTP_400_BAD_REQUEST)
			
		except Exception as e:
			print(e)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# def generateFormatSolicitud2(request):
# 	if request.method == 'GET':
# 		try:

# 			solicitudId = request.GET['solicitud_id']

# 			solicitud = Solicitud.objects.get(pk = solicitudId)
# 			correspondencia = CorrespondenciaEnviada.objects.get(pk = solicitud.correspondenciasolicita.id)

# 			ruta = settings.STATICFILES_DIRS[0]
# 			newpath = ruta + '/papelera/'
# 			filename = "plantillas/multas/plantillaImposicion.docx"
# 			extension = filename[filename.rfind('.'):]
# 			nombre = 'empresa'+str(correspondencia.empresa.id)+str(correspondencia.consecutivo)+extension
# 			functions.descargarArchivoS3(str(filename), str(newpath) , nombre )	
# 			plantilla = settings.STATICFILES_DIRS[0] + '/papelera/'+nombre

# 			word = win32.gencache.EnsureDispatch('Word.Application')
# 			doc=word.Documents.Open(settings.STATICFILES_DIRS[0]+"/papelera/plantillaImposicion.docx")
# 			doc=word.Documents.Open("C:\\a2.docx")
# 			word.Visible = False

# 			word.ActiveDocument.Sections(1).Headers(win32.constants.wdHeaderFooterPrimary).Range.Find.Execute("consecutivo", False, False, False, False, False, True, 1, False, str(solicitud.contrato.mcontrato.nombre)+" No. "+str(solicitud.consecutivo), 2)
# 			doc.SaveAs("E:\\new_file.docx")
# 			doc.ActiveWindow.Close()
			
# 			# # worddoc = word.Documents.Open(r"C:\\a2.docx")
# 			# worddoc.SaveAs(r"E:\\new_file")
# 			# worddoc.ActiveWindow.Close()

# 		except Exception as e:
# 			print(e)
# 			pass	
			
# def generateFormatSolicitud(request):
# 	nombreArchivo=''
# 	if request.method == 'GET':
# 		nombre_modulo='multa.generateFormatSolicitud'
# 		try:			

# 			solicitudId = request.GET['solicitud_id']
# 			solicitud = Solicitud.objects.get(pk = solicitudId)
# 			correspondencia = CorrespondenciaEnviada.objects.get(pk = solicitud.correspondenciasolicita.id)

# 			ruta = settings.STATICFILES_DIRS[0]
# 			newpath = ruta + '/papelera/'
# 			filename = "plantillas/multas/plantillaImposicion.docx"
# 			extension = filename[filename.rfind('.'):]
# 			nombre = 'plantillaImposicion.docx'+extension
# 			functions.descargarArchivoS3(str(filename), str(newpath) , nombre )	
# 			plantilla = settings.STATICFILES_DIRS[0] + '/papelera/'+nombre

# 			word = win32.gencache.EnsureDispatch('Word.Application')
# 			doc=word.Documents.Open(plantilla)
# 			word.Visible = False
# 			word.ActiveDocument.Sections(1).Headers(win32.constants.wdHeaderFooterPrimary).Range.Find.Execute("consecutivo", False, False, False, False, False, True, 1, False, str(solicitud.contrato.mcontrato.nombre)+" No. "+str(solicitud.consecutivo), 2)
# 			new_name_file = str(newpath)+str(solicitud.consecutivo)+extension
# 			doc.SaveAs(new_name_file)
# 			doc.ActiveWindow.Close()

# 			f = open(new_name_file,'rb')
# 			document = Document(f)

# 			# document.add_paragraph('section one ' * 100)
# 			sections = document.sections # Sections object
# 			# print('len(sections) = {}'.format(len(sections))) # 1
# 			section = sections[0] # Section object
# 			# section.page_width = Cm(30.0)
# 			# section.page_height = Cm(1.0)
# 			section.left_margin = Cm(2.54)
# 			section.right_margin = Cm(2.54)
# 			section.top_margin = Cm(2.54)


# 			# for section in sections:
# 			# 	print(section.start_type) # NEW_PAGE (2)

# 			styles = document.styles
# 			# table_styles = [s for s in styles ]
# 			# for style in table_styles:
# 			# 	print(style.name)

# 			style = document.styles['Normal']
# 			font = style.font

# 			font.name = 'Arial'
# 			font.size = Pt(11)

# 			table = document.add_table(rows=1, cols=7 , style= styles['Table Grid'].name)
# 			col = table.columns[0] 
# 			# col.width=Inches(6.996189)
# 			col.width=Inches(3.5)

# 			hdr_cells = table.rows[0].cells	
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('FECHA DE DILIGENCIAMIENTO:')
# 			run.bold = True

# 			parrafo = hdr_cells[1].add_paragraph()
# 			run = parrafo.add_run('DIA')

# 			parrafo = hdr_cells[2].add_paragraph()
# 			run = parrafo.add_run(str((solicitud.fechaDiligencia).strftime("%d")))

# 			parrafo = hdr_cells[3].add_paragraph()
# 			run = parrafo.add_run('MES:')

# 			parrafo = hdr_cells[4].add_paragraph()
# 			run = parrafo.add_run(str((solicitud.fechaDiligencia).strftime("%m")))


# 			parrafo = hdr_cells[5].add_paragraph()
# 			run = parrafo.add_run(unicode('AÑO', 'utf-8'))

# 			parrafo = hdr_cells[6].add_paragraph()
# 			run = parrafo.add_run(str((solicitud.fechaDiligencia).strftime("%Y")))
# 			col = table.columns[6] 
# 			col.width=Inches(1)

# 			table = document.add_table(rows=1, cols=2 , style= styles['Table Grid'].name)
# 			col = table.columns[0] 
# 			col.width=Inches(3.5)

# 			col = table.columns[1] 
# 			col.width=Inches(5.64)

# 			hdr_cells = table.rows[0].cells	
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(unicode('CONTRATO:', 'utf-8'))
# 			run.bold = True

# 			parrafo = hdr_cells[1].add_paragraph()
# 			run = parrafo.add_run(solicitud.contrato.numero)

# 			hdr_cells = table.add_row().cells
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('CONTRATANTE:')
# 			run.bold = True

# 			parrafo = hdr_cells[1].add_paragraph()
# 			run = parrafo.add_run(solicitud.contrato.contratante.nombre)

# 			hdr_cells = table.add_row().cells
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('CONTRATISTA:')
# 			run.bold = True


# 			parrafo = hdr_cells[1].add_paragraph()
# 			run = parrafo.add_run(solicitud.contrato.contratista.nombre)

# 			hdr_cells = table.add_row().cells
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('NOMBRE DEL PROYECTO:')
# 			run.bold = True


# 			proyecto = Proyecto.objects.get(contrato = solicitud.contrato.id)

# 			parrafo = hdr_cells[1].add_paragraph()
# 			run = parrafo.add_run(proyecto.nombre)
# 			# run = parrafo.add_run('proyecto')

# 			hdr_cells = table.add_row().cells
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('POTESTAD PARA IMPOSICION DE MULTAS:')
# 			run.bold = True
# 			parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 			parrafo = hdr_cells[0].add_paragraph()
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(unicode('En el marco del  presente Contrato '+str(solicitud.contrato.contratante.nombre)+'. Ostenta la facultad  para aplicar multas al contratista, frente al evidente incumplimiento de las obligaciones contractuales, siempre que esta sanción sea impuesta teniendo en cuenta los procedimientos previstos para tal fin.', 'utf-8' ))
# 			parrafo = hdr_cells[0].add_paragraph()
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('CLAUSULA DECIMA TERCERA- MULTAS Y PENAS-')
# 			run.bold = True

# 			parrafo = hdr_cells[0].add_paragraph()
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('MULTAS: ')
# 			run.bold = True

# 			parrafo = hdr_cells[0].add_paragraph()
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(unicode('Sin perjuicio de la aplicación de la cláusula penal pecuniaria contenida en el presente contrato, en caso de incumplimiento del CONTRATISTA con las obligaciones a su cargo '+str(solicitud.contrato.contratante.nombre)+' PODRA IMPONERLE a su juicio y sin que por ello se entienda extinguida la obligación principal a cargo del CONTRATISTA, las multas.', 'utf-8' ))
# 			run.add_break()
# 			run.add_break()
# 			run.add_break()

# 			a, b = hdr_cells[:2]
# 			A = a.merge(b)

# 			hdr_cells = table.add_row().cells
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('HECHOS GENERADORES DE INCUMPLIMIENTO:')
# 			run.bold = True
# 			run.add_break()
# 			run.add_break()
# 			parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(correspondencia.contenido)
# 			parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 			a, b = hdr_cells[:2]
# 			A = a.merge(b)

# 			hdr_cells = table.add_row().cells
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('CLAUSULAS  CONTRACTUALES  AFECTADAS:')
# 			run.bold = True
# 			parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 			run.add_break()
# 			run.add_break()
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(correspondencia.clausula_afectada)
# 			run.add_break()
# 			run.add_break()	
# 			parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY	

# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(unicode('TASACIÓN DE  LAS MULTAS.', 'utf-8' ))
# 			run.bold = True
# 			parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 			a, b = hdr_cells[:2]
# 			A = a.merge(b)

# 			parrafo = hdr_cells[0].add_paragraph()
# 			table2 = hdr_cells[0].add_table(rows=1, cols=2 )
# 			table2.alignment = WD_TABLE_ALIGNMENT.CENTER
# 			table2.style = 'Table Grid'

			
# 			hdr_cells2 = table2.rows[0].cells
# 			parrafo = hdr_cells2[0].add_paragraph()
# 			run = parrafo.add_run(unicode('EVENTO', 'utf-8'))
# 			run.bold = True
# 			hdr_cells2[0].width=Inches(2.7)
# 			hdr_cells2[1].width=Inches(2.7)
# 			parrafo = hdr_cells2[1].add_paragraph()
# 			run = parrafo.add_run(unicode('VALOR DE LA MULTA', 'utf-8'))
# 			run.bold = True
# 			parrafo = hdr_cells2[0].add_paragraph()

# 			eventos = SolicitudEvento.objects.filter(solicitud_id = solicitud.id)

# 			for item in eventos:
# 				row_cells = table2.add_row().cells
# 				parrafo = row_cells[0].add_paragraph()
# 				run = parrafo.add_run(item.evento.nombre)

# 				valor_evento = format_decimal(item.evento.valor,  locale='es')
# 				parrafo = row_cells[1].add_paragraph()
# 				run = parrafo.add_run("$"+str(valor_evento))

# 				row_cells[0].width=Inches(2.7)
# 				row_cells[1].width=Inches(2.7)


# 			parrafo = hdr_cells[0].add_paragraph()
# 			table3 = hdr_cells[0].add_table(rows=3, cols=2 )
# 			table3.alignment = WD_TABLE_ALIGNMENT.CENTER
# 			table3.style = 'Table Grid'

# 			hdr_cells2 = table3.rows[0].cells
# 			parrafo = hdr_cells2[0].add_paragraph()
# 			run = parrafo.add_run(unicode('PRIMER HECHO.', 'utf-8'))
# 			run.bold = True
# 			parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 			a, b = hdr_cells2[:2]
# 			A = a.merge(b)
# 			hdr_cells2[0].width=Inches(5)

# 			hdr_cells2 = table3.rows[1].cells
# 			parrafo = hdr_cells2[0].add_paragraph()
# 			run = parrafo.add_run(unicode('SANCIÓN', 'utf-8'))
# 			run.bold = True
# 			hdr_cells2[0].width=Inches(2.7)
# 			hdr_cells2[1].width=Inches(2.7)
# 			parrafo = hdr_cells2[1].add_paragraph()

# 			numero = number_to_letters(eventos.count())

# 			if eventos.count()==1:
# 				frase = 'Se relaciona '
# 			else:
# 				frase = 'Se relacionan '

# 			run = parrafo.add_run(unicode(frase+numero+'('+str(eventos.count())+') incumplimiento.', 'utf-8'))

# 			hdr_cells2 = table3.rows[2].cells
# 			parrafo = hdr_cells2[0].add_paragraph()
# 			run = parrafo.add_run(unicode('TASACIÓN:', 'utf-8'))
# 			run.bold = True
# 			hdr_cells2[0].width=Inches(2)
# 			hdr_cells2[1].width=Inches(2)
# 			parrafo = hdr_cells2[1].add_paragraph()

# 			valor_imposicion = format_decimal(solicitud.valorImpuesto,  locale='es')
# 			run = parrafo.add_run('$ '+valor_imposicion)
		
# 			valor_imposicion_letra = number_to_letters(solicitud.valorImpuesto)
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run('')
# 			run.add_break()
# 			run.add_break()
# 			run = parrafo.add_run('El valor total  por concepto de MULTAS, es por la suma de ')
# 			run = parrafo.add_run('$'+valor_imposicion+'. '+valor_imposicion_letra.upper()+' PESOS ML.')
# 			run.bold = True
# 			run.add_break()
# 			run.add_break()
			

# 			pruebas = SolicitudSoporte.objects.filter(anulado=False , solicitud_id = solicitud.id)

# 			if pruebas:
# 				parrafo = hdr_cells[0].add_paragraph()
# 				run = parrafo.add_run('PRUEBAS:')
# 				run.bold = True
# 				run.add_break()
# 				parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 				for item in pruebas:
# 					parrafo = hdr_cells[0].add_paragraph()
# 					run = parrafo.add_run('- '+item.nombre)

# 				run = parrafo.add_run('')
# 				run.add_break()
# 				run.add_break()
# 				run.add_break()

# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(unicode('De acuerdo a lo dispuesto en el procedimiento para la imposición de multas, usted contara cinco (5) días hábiles para presentar los descargos, de no hacerlo esta multa quedara en firme de manera inmediata, en dicho caso no existe obligatoriedad por parte de '+str(solicitud.contrato.contratante.nombre)+'. De ratificar la decisión.', 'utf-8' ))
# 			parrafo = hdr_cells[0].add_paragraph()
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(unicode('En caso de presentar los descargos la decisión quedara sujeta al pronunciamiento de '+str(solicitud.contrato.contratante.nombre)+'. Quien evaluara la procedencia de la impugnación.', 'utf-8' ))
# 			run.add_break()
# 			run.add_break()
# 			run.add_break()
# 			run.add_break()
# 			run.add_break()
# 			run.add_break()
# 			run = parrafo.add_run(solicitud.firmaImposicion.persona.nombres+' '+solicitud.firmaImposicion.persona.apellidos)
# 			run.bold = True
# 			run.add_break()
# 			run = parrafo.add_run(unicode('Proyectos de Electrificación .', 'utf-8' ))
# 			parrafo = hdr_cells[0].add_paragraph()
# 			run = parrafo.add_run(solicitud.contrato.contratante.nombre)
# 			parrafo = hdr_cells[0].add_paragraph()
# 			parrafo = hdr_cells[0].add_paragraph()

# 			solicitud_empresa = SolicitudEmpresa.objects.filter(solicitud_id = solicitud.id).exclude(empresa_id=solicitud.firmaImposicion.empresa.id)
			
# 			for item in solicitud_empresa:
# 				run = parrafo.add_run(str('C.C. ')+item.empresa.nombre)


# 			# a, b = hdr_cells[:2]
# 			# A = a.merge(b)

# 			nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\FormatoNotificacion.docx'
# 			document.save(nombreArchivo)

# 			# except Exception as e:
# 			# 	print(e)
# 			# 	f.close()
				


# 			chunk_size = 108192

# 			nombreDescarga = 'FormatoNotificacion.docx'
# 			response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
# 				content_type=mimetypes.guess_type(nombreArchivo)[0])
# 			response['Content-Length'] = os.path.getsize(nombreArchivo)
# 			response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
			
# 			return response 
# 		except Proyecto.DoesNotExist:			
# 			return JsonResponse({'message':'No existe un proyecto asociado al contrato','status':'fail','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# 		except Exception as e:
# 			functions.toLog(e,nombre_modulo)
# 			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

		

@login_required
def index(request):
	es=EstadoMulta()
	empresaActual = request.user.usuario.empresa.id

	# total solicitudes a generar
	qsetSolicitudEmpresa = SolicitudEmpresa.objects.filter(empresa_id = empresaActual, propietario = False )

	idSolicitud = []
	for i in qsetSolicitudEmpresa:
		idSolicitud.append(i.solicitud_id)
		# print i.solicitud_id

	queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = idSolicitud).values('solicitud_id').annotate(id=Max('id')).values("id")
	queryset = SolicitudHistorial.objects.filter(id__in = queryHistorial , estado_id = es.solicitada).count()

	# TOTAL DESCARGOS SIN PRONUNCIAR
	queryHistorial_descargo = SolicitudHistorial.objects.filter(solicitud__fk_Solicitud_SolicitudEmpresa__empresa_id = empresaActual
																, solicitud__fk_Solicitud_SolicitudEmpresa__propietario = True).values('solicitud_id').annotate(id=Max('id')).values("id")

	queryset_notificadas = SolicitudHistorial.objects.filter(id__in = queryHistorial_descargo , estado_id = es.notificada_contratista)

	idConsulta = []
	for item in queryset_notificadas:
		idConsulta.append(item.solicitud.id)

	queryset_descargo = SolicitudApelacion.objects.filter(solicitud__in = idConsulta 
														, fk_Apelacion_SolicitudPronunciamiento__apelacion__isnull = True)

	queryset_descargo = queryset_descargo.count()
	
	return render(request, 'multa/index.html',{'total_descargo_sin_pronunciar': int(queryset_descargo) ,'total_generar': int(queryset) ,'model':'multa','app':'multa'})

@login_required
def multaElaboradas(request):
	return render(request, 'multa/multaElaboradas.html',{'model':'solicitud','app':'multa'})

@login_required
def multaElaborar(request):
	qsetEstadoContratos = Estado.objects.filter(app = "contrato")
	qsetConjuntoEvento = ConjuntoEvento.objects.all()

	departamentos = Departamento.objects.filter().values()
	empresaActual = request.user.usuario.empresa.id
	prefijos = CorresPfijo.objects.filter(empresa_id = empresaActual)
	funcionarios_firman = Funcionario.objects.filter(empresa_id = empresaActual)
	empresas_acceso = EmpresaAcceso.objects.filter(empresa_id = empresaActual).values_list('empresa_ver_id')

	qset = ( Q(pk__in = empresas_acceso) | Q(pk = empresaActual)  )
	
	empresas = Empresa.objects.filter(qset)


	return render(request, 'multa/multaElaborar.html',{ 'departamentos': departamentos, 'prefijos': prefijos , 'empresas':empresas ,'funcionarios_firman':funcionarios_firman , 'conjunto_eventos' : qsetConjuntoEvento , 'estados_contratos':qsetEstadoContratos , 'model':'solicitud','app':'multa'})

@login_required
def multaSolicitadas(request):
	empresaActual = request.user.usuario.empresa.id
	qsetFuncionariosFirman = Funcionario.objects.filter(empresa_id = empresaActual)
	return render(request, 'multa/multaSolicitadas.html',{'funcionarios_firman' : qsetFuncionariosFirman , 'model':'solicitud','app':'multa'})

@login_required
def multas(request):
	return render(request, 'multa/multas.html',{'model':'solicitud','app':'multa'})

@login_required
def multaConfirmadas(request):
	return render(request, 'multa/multaConfirmadas.html',{'model':'solicitud','app':'multa'})

@login_required
def multaPresentarDescargo(request , id = None):
	empresaActual = request.user.usuario.empresa.id
	qsetSolicitud = Solicitud.objects.get(pk = id ) 
	return render(request, 'multa/multaPresentarDescargo.html',{'estado' : qsetSolicitud.estado().estado.nombre, 'solicitud' : qsetSolicitud , 'model':'solicitudapelacion','app':'multa'})

@login_required
def multaRespuestaDescargo(request , id = None):

	solicitud = Solicitud.objects.get(pk = id)

	departamentos = Departamento.objects.filter().values()
	empresaActual = request.user.usuario.empresa.id
	prefijos = CorresPfijo.objects.filter(empresa_id = empresaActual)
	funcionarios_firman = Funcionario.objects.filter(empresa_id = empresaActual)
	empresas = EmpresaAcceso.objects.filter(empresa_id = empresaActual)

	
	estados_posibles = Estados_posibles.objects.filter(actual_id = solicitud.estado().estado.id)

	return render(request, 'multa/multaRespuestaDescargo.html'
							,{'departamentos': departamentos, 'prefijos':prefijos ,'funcionarios':funcionarios_firman 
							,'estados_posibles':estados_posibles
							,'empresas':empresas
							,'solicitud':solicitud
							,'model':'solicitudapelacion','app':'multa'}
							)

@login_required
def descargos(request):
	return render(request, 'multa/descargos.html',{'model':'solicitudapelacion','app':'multa'})

@login_required
def eventos(request):
	return render(request, 'multa/eventos.html',{'model':'evento','app':'multa'})

@login_required
def multaHistorial(request , id = None):
	empresaActual = request.user.usuario.empresa.id
	qsetSolicitud = Solicitud.objects.get(pk = id ) 
	qsetSolicitudHistorial = SolicitudHistorial.objects.filter(solicitud_id = id)
	return render(request, 'multa/multaHistorial.html',{  'estado' : qsetSolicitud.estado().estado.nombre, 'solicitud' : qsetSolicitud ,'solicitud_historial' : qsetSolicitudHistorial , 'model':'solicitud','app':'multa'})

def createWordSolicitud(request):

	nombreArchivo=''
	if request.method == 'GET':
		correspondencia_id = request.GET['correspondencia_id']
		# documento = CorrespondenciaEnviada.objects.get(id=1034)
		documento = CorrespondenciaEnviada.objects.get(id=correspondencia_id)

		try:
			plantilla = CorrespondenciaPlantilla.objects.get(empresa_id = documento.empresa.id)
		except CorrespondenciaPlantilla.DoesNotExist:
			return JsonResponse({'message':'La empresa no ha cargado la plantilla para las correspondencias','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		

		ruta = settings.STATICFILES_DIRS[0]
		newpath = ruta + '/papelera/'

		# newpath = r'static/papelera/'
		# filename = "plantillas/correspondenciaEnviada/plantilla1.docx"
		filename = str(plantilla.soporte)
		extension = filename[filename.rfind('.'):]
		nombre = 'empresa'+str(documento.empresa.id)+str(documento.consecutivo)+extension
		functions.descargarArchivoS3(str(filename), str(newpath) , nombre )	
		
		plantilla = settings.STATICFILES_DIRS[0] + '/papelera/'+nombre			
		solicitud = Solicitud.objects.get( correspondenciasolicita_id = correspondencia_id )


		try:
			f = open(plantilla,'rb')
			doc = Document(f)

			styles = doc.styles
			# doc.add_paragraph('Intense quote', style= styles["Luis"].name)
			# print styles["Luis"].name

			table_styles = [s for s in styles ]

			# for style in table_styles:
			# 	print(style.name)

			#codigo para buscar los textos y reemplazarlos con los del objeto documento
			for parrafo in doc.paragraphs:
				texto = parrafo.text

				if texto == '<fechaenvio>':
					fecha_envio = format_date(documento.fechaEnvio, format='long' , locale='es')

					ciudad = documento.ciudad.nombre
					parrafo.text = str(ciudad.capitalize())+' , '+str(fecha_envio)
				if texto == '<consecutivo>':
					parrafo.text = str(documento.prefijo.nombre)+'-'+str(documento.consecutivo)
				if texto == '<referencia>':

					parrafo.text = ''	
					run = parrafo.add_run('REFERENCIA :')
					run.bold = True	
					run.add_tab()
					run = parrafo.add_run(documento.referencia)
					run.add_break()	
					run = parrafo.add_run()

				if texto.strip() == '<destinatario>':
					parrafo.text = ''
					run = parrafo.add_run(unicode('Señores', 'utf-8'))
					run.add_break()	
					run = parrafo.add_run(documento.empresa_destino)
					run.bold = True
					run.add_break()	
					run = parrafo.add_run('Atn: '+documento.persona_destino)
					run.add_break()	
					run = parrafo.add_run(documento.cargo_persona)
					run.bold = True
					run.add_break()	
					if documento.direccion:
						run = parrafo.add_run(documento.direccion)
						run.add_break()	
					if documento.telefono:
						run = parrafo.add_run('Tel.'+documento.telefono+'.')
						run.add_break()
					if documento.municipioEmpresa:
						parrafo.add_run(documento.municipioEmpresa.nombre+'.')

				if texto.strip() == '<asunto>':
					parrafo.text = ''	
					run = parrafo.add_run('ASUNTO :')
					run.bold = True	
					run.add_tab()
					run = parrafo.add_run(documento.asunto)
					run.add_break()	
					run = parrafo.add_run()

				if texto == '<contenido>':

					parrafo.text = ''

					prior_paragraph = parrafo.insert_paragraph_before('')
					run = prior_paragraph.add_run('Por medio de la presente ')
					run = prior_paragraph.add_run(documento.firma.empresa.nombre)
					run.bold = True
					run = prior_paragraph.add_run('. en funcion de Interventor del contrato : ')
					run = prior_paragraph.add_run(str(solicitud.contrato.numero)+str(' - '))
					run.bold = True
					run = prior_paragraph.add_run(solicitud.contrato.nombre)
					run = prior_paragraph.add_run(' , recomienda muy respetuosamente se imponga las multas correspondientes por los incumplimientos contractuales que se vienen presentando por parte del contratista de obra ')
					run = prior_paragraph.add_run(solicitud.contrato.contratista.nombre)
					run.bold = True
					run = prior_paragraph.add_run(unicode(' cuales se describe a continuación ', 'utf-8'))


					prior_paragraph = parrafo.insert_paragraph_before('')
					run = prior_paragraph.add_run('HECHOS GENERADORES DE INCUMPLIMIENTO')	
					run.bold = True			
					prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER	
					prior_paragraph = parrafo.insert_paragraph_before('')
					run = prior_paragraph.add_run(documento.contenido)

					
					if documento.clausula_afectadaHtml:

						prior_paragraph = parrafo.insert_paragraph_before('')
						run = prior_paragraph.add_run('CLAUSULAS AFECTADAS')	
						run.bold = True			
						prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER	

						prior_paragraph = parrafo.insert_paragraph_before('')
						run = prior_paragraph.add_run(documento.clausula_afectada)	

					

					prior_paragraph = parrafo.insert_paragraph_before('')
					# Set a cell background (shading) color to RGB D9D9D9.
					# shading_elm = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
					shading_elm2 = parse_xml(r'<w:top {} w:val="single" w:sz="4" w:space="0" w:color="auto" />'.format(nsdecls('w')))
		
					# TABLA DE LOS EVENTOS IMPUESTOS
					styles = doc.styles
					# table_styles = [s for s in styles ]
					# for style in table_styles:
					# 	print(style.name)

					# if styles['TableGrid'].name:
					# 	print "Existe"

					# else:

					# 	print "No existe"					

					
					table2 = doc.add_table(rows=1, cols=2 )
					table2.alignment = WD_TABLE_ALIGNMENT.CENTER					

					try:
						table2.style = 'TableGrid'
					except Exception as e:
						table2.style = 'Normal Table'
				

					hdr_cells2 = table2.rows[0].cells
					parrafo = hdr_cells2[0].add_paragraph()
					run = parrafo.add_run(unicode('EVENTO', 'utf-8'))
					run.bold = True
					hdr_cells2[0].width=Inches(2.7)
					hdr_cells2[1].width=Inches(2.7)

					# hdr_cells2[1]._tc.get_or_add_tcPr().append(shading_elm2)


					parrafo = hdr_cells2[1].add_paragraph()
					run = parrafo.add_run(unicode('VALOR DE LA MULTA', 'utf-8'))
					run.bold = True
					parrafo = hdr_cells2[0].add_paragraph()
					
					
					eventos = SolicitudEvento.objects.filter(solicitud_id = solicitud.id)

					for item in eventos:
						row_cells = table2.add_row().cells
						parrafo = row_cells[0].add_paragraph()
						run = parrafo.add_run(item.evento.nombre)

						valor_evento = format_decimal(item.evento.valor,  locale='es')
						parrafo = row_cells[1].add_paragraph()
						run = parrafo.add_run("$ "+str(valor_evento))

						row_cells[0].width=Inches(2.7)
						row_cells[1].width=Inches(2.7)

					addTableAfterParagraph(table2,prior_paragraph)


				if texto == '<autor>':

					parrafo.text = ''

					run = parrafo.add_run('Cordialmente,')
					run.add_break()	
					run.add_break()	
					run.add_break()	
					run = parrafo.add_run(unicode('_____________________________________', 'utf-8'))
					run.add_break()	
					run = parrafo.add_run(documento.firma.persona.nombres+' '+documento.firma.persona.apellidos)
					run.bold = True
					run.add_break()	
					run = parrafo.add_run(documento.firma.cargo.nombre)
					run.add_break()
					run = parrafo.add_run(documento.firma.empresa.nombre)
					run.add_break()		

				# print parrafo.text
			nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\c' + str(documento.id) + '.docx'
			doc.save(nombreArchivo )
			f.close
		except Exception as e:

			# SI OCURRE UNA EXCEPCION
			print(e)			
			document = Document()
			document.add_heading('Error', 0)
			p = document.add_paragraph(str('se presentaron erores al crear la carta.')+str(e))

			document.add_page_break()
			nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\error.docx'
			document.save(nombreArchivo)

			chunk_size = 108192

			nombreDescarga = 'documento.docx'
			response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
				content_type=mimetypes.guess_type(nombreArchivo)[0])
			response['Content-Length'] = os.path.getsize(nombreArchivo)
			response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
			
			return response 


	chunk_size = 108192

	nombreDescarga = 'documento'+str(documento.prefijo.nombre)+'-'+str(documento.consecutivo)+'.docx'
	response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
		content_type=mimetypes.guess_type(nombreArchivo)[0])
	response['Content-Length'] = os.path.getsize(nombreArchivo)
	response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
	
	return response 

def generate_format_respuesta_descargo(request):
	if request.method == 'GET':		

		try:
			solicitud_id = request.GET['solicitud_id']

			solicitud = Solicitud.objects.get(pk = solicitud_id)
			documento = CorrespondenciaEnviada.objects.get(id=solicitud.correspondenciadescargo.id)
			

			try:
				plantilla = CorrespondenciaPlantilla.objects.get(empresa_id = documento.empresa.id)
			except CorrespondenciaPlantilla.DoesNotExist:
				return JsonResponse({'message':'La empresa no ha cargado la plantilla para las correspondencias','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		

			ruta = settings.STATICFILES_DIRS[0]
			newpath = ruta + '/papelera/'

			# newpath = r'static/papelera/'
			filename = str(plantilla.soporte)
			extension = filename[filename.rfind('.'):]
			nombre = 'empresa'+str(documento.empresa.id)+str(documento.consecutivo)+extension
			functions.descargarArchivoS3(str(filename), str(newpath) , nombre )	
			
			plantilla = settings.STATICFILES_DIRS[0] + '/papelera/'+nombre


			f = open(plantilla,'rb')
			doc = Document(f)

			styles = doc.styles
			table_styles = [s for s in styles ]

			for style in table_styles:
				print(style.name)


			# STYLES


			# style fecha envio and consecutivo
			obj_styles = styles
			obj_charstyle = obj_styles.add_style('fechaenvioStyle', WD_STYLE_TYPE.CHARACTER)
			obj_font = obj_charstyle.font
			obj_font.size = Pt(11)
			obj_font.name = 'Arial Unicode MS'

			# style referencia and consecutivo
			obj_styles = styles
			obj_charstyle = obj_styles.add_style('referenciaStyle', WD_STYLE_TYPE.CHARACTER)
			obj_font = obj_charstyle.font
			obj_font.size = Pt(8)
			obj_font.name = 'Arial Unicode MS'

			# style asunto and consecutivo
			obj_styles = styles
			obj_charstyle = obj_styles.add_style('asuntoStyle', WD_STYLE_TYPE.CHARACTER)
			obj_font = obj_charstyle.font
			obj_font.size = Pt(11)
			obj_font.name = 'Arial Unicode MS'

			# style destinatario and consecutivo
			obj_styles = styles
			obj_charstyle = obj_styles.add_style('destinatarioStyle', WD_STYLE_TYPE.CHARACTER)
			obj_font = obj_charstyle.font
			obj_font.size = Pt(11)
			obj_font.name = 'Arial Unicode MS'

			# style contenido and consecutivo
			obj_styles = styles
			obj_charstyle = obj_styles.add_style('contenidoStyle', WD_STYLE_TYPE.CHARACTER)
			obj_font = obj_charstyle.font
			obj_font.size = Pt(11)
			obj_font.name = 'Arial Unicode MS'


			# FIN STYLES

			#codigo para buscar los textos y reemplazarlos con los del objeto documento
			for parrafo in doc.paragraphs:
				texto = parrafo.text

				if texto == '<fechaenvio>':
					fecha_envio = format_date(documento.fechaEnvio, format='long' , locale='es')
					parrafo.text = ''

					ciudad = str(documento.ciudad.nombre)
					run = parrafo.add_run(ciudad.capitalize()+' , '+str(fecha_envio) , style = 'fechaenvioStyle' )

				if texto == '<consecutivo>':
					parrafo.text = ''
					run = parrafo.add_run('CONSECUTIVO:'+str(documento.prefijo.nombre)+'-'+str(documento.consecutivo) , style = 'referenciaStyle')

				if texto == '<referencia>':
					parrafo.text = ''
					run = parrafo.add_run('REFERENCIA:'+documento.referencia , style = 'referenciaStyle')
					
				if texto == '<destinatario>':
					parrafo.text = ''
					run = parrafo.add_run(unicode('Señores', 'utf-8') , style = 'destinatarioStyle' )
					run.add_break()	
					run = parrafo.add_run(documento.empresa_destino , style = 'destinatarioStyle' )
					run.bold = True
					run.add_break()	
					run = parrafo.add_run('Atn: '+documento.persona_destino , style = 'destinatarioStyle')
					run.add_break()	
					run = parrafo.add_run(documento.cargo_persona , style = 'destinatarioStyle')
					run.add_break()	
					if documento.direccion:
						run = parrafo.add_run(documento.direccion , style = 'destinatarioStyle' )
						run.add_break()	
					if documento.telefono:
						run = parrafo.add_run('Tel.'+documento.telefono+'.' , style = 'destinatarioStyle')
						run.add_break()
					if documento.municipioEmpresa:
						parrafo.add_run(documento.municipioEmpresa.nombre+'.' , style = 'destinatarioStyle')


				if texto == '<asunto>':
					parrafo.text = ''	
					run = parrafo.add_run('ASUNTO :' , style = 'asuntoStyle')
					run.bold = True	
					run.add_tab()
					run = parrafo.add_run(documento.asunto , style = 'asuntoStyle')
					run.bold = True	
					run.add_break()	
					run.add_break()

					# parrafo de complemento del asunto 
					fecha_diligencia = format_date(solicitud.fechaDiligencia, format='long' , locale='es')

					run = parrafo.add_run(unicode('Mediante oficio de fecha '+str(fecha_diligencia)+' y Habiéndose cursadas las notificaciones de rigor como lo establece el contrato '+str(solicitud.contrato.numero)+', ', 'utf-8') , style = 'asuntoStyle')

					run = parrafo.add_run('ELECTRICARIBE' , style = 'asuntoStyle')
					run.bold = True

					run = parrafo.add_run(unicode(' aplicó multa por incumplimiento a ', 'utf-8' ) , style = 'asuntoStyle')
				
					run = parrafo.add_run(solicitud.contrato.contratista.nombre)	
					run.bold = True

					run = parrafo.add_run(unicode('. A su vez,  el contratista presentó descargos por la  imposición de la multa dentro de los términos previstos. En consecuencia se procede a estudiar los fundamentos de la impugnación.', 'utf-8') , style = 'asuntoStyle')

				
				if texto == '<contenido>':

					parrafo.text = ''

					prior_paragraph = parrafo.insert_paragraph_before('')
					run = prior_paragraph.add_run('CONSIDERACIONES  DE LA INTERVENTORIA' , style = 'contenidoStyle' )
					run.bold = True
					run.add_break()
					prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

					try:
						pronunciamiento_interventoria = SolicitudPronunciamiento.objects.get(apelacion__solicitud__id = solicitud.id)
						pronunciamiento_comentarios = pronunciamiento_interventoria.comentarios
					except SolicitudPronunciamiento.DoesNotExist:
						pronunciamiento_comentarios = ''
					

					prior_paragraph = parrafo.insert_paragraph_before('')
					run = prior_paragraph.add_run( pronunciamiento_comentarios , style = 'asuntoStyle')
					run.add_break()
					run.add_break()
					prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

					prior_paragraph = parrafo.insert_paragraph_before('')
					run = prior_paragraph.add_run('PARA RESOLVER SE ANALIZA' , style = 'contenidoStyle')
					run.bold = True
					run.add_break()
					prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
					# print "debug"
					prior_paragraph = parrafo.insert_paragraph_before('')
					run = prior_paragraph.add_run(solicitud.correspondenciadescargo.contenido , style = 'asuntoStyle')
					prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


					prior_paragraph = parrafo.insert_paragraph_before('' )
					run = prior_paragraph.add_run("En consecuencia de lo anterior, y teniendo en cuenta las consideraciones de las partes y el estudio objetivo de las mismas , se " , style = 'asuntoStyle' )
					run.add_text('confirmada ').bold = True
					run.add_text('el valor de la  Multa interpuesta al contratista ')
					prior_paragraph.add_run(solicitud.contrato.contratista.nombre).bold = True
					prior_paragraph.add_run(' , por la suma de ' , style = 'asuntoStyle')

					valor_imposicion = format_decimal(solicitud.valorImpuesto,  locale='es')

					valor_imposicion_letra = number_to_letters(solicitud.valorImpuesto)

					prior_paragraph.add_run('$ '+valor_imposicion +' '+valor_imposicion_letra+' PESOS ML.' )

					prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

				if texto == '<autor>':

					parrafo.text = ''

					run = parrafo.add_run(unicode('_____________________________________', 'utf-8') , style = 'asuntoStyle')
					run.add_break()	
					run = parrafo.add_run(documento.firma.persona.nombres+' '+documento.firma.persona.apellidos , style = 'asuntoStyle')
					run.bold = True
					run.add_break()	
					run = parrafo.add_run(unicode('Proyectos de Electrificación.', 'utf-8') , style = 'asuntoStyle')
					run.add_break()
					run = parrafo.add_run(unicode('Electricaribe. S.A. ESP.', 'utf-8') , style = 'asuntoStyle')
					run.add_break()	

			# print parrafo.text
			nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\c' + str(documento.id) + '.docx'
			doc.save(nombreArchivo )
			f.close

			chunk_size = 108192

			nombreDescarga = 'documento'+str(documento.prefijo.nombre)+'-'+str(documento.consecutivo)+'.docx'
			response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
				content_type=mimetypes.guess_type(nombreArchivo)[0])
			response['Content-Length'] = os.path.getsize(nombreArchivo)
			response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
			
			return response 

		except Exception as e:
			print(e)
			
			document = Document()

			document.add_heading('Error', 0)

			p = document.add_paragraph(str('se presentaron erores al crear la carta.')+str(e))


			document.add_page_break()
			nombreArchivo = settings.STATICFILES_DIRS[0] + '\papelera\error.docx'
			document.save(nombreArchivo)

			chunk_size = 108192

			nombreDescarga = 'documento.docx'
			response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
				content_type=mimetypes.guess_type(nombreArchivo)[0])
			response['Content-Length'] = os.path.getsize(nombreArchivo)
			response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
			
			return response 


def generate_format_OF(request):
	if request.method == 'GET':	
		try:

			empresa_id = request.user.usuario.empresa.id
			solicitud_id = request.GET['solicitud_id']

			solicitud = Solicitud.objects.get(pk = solicitud_id)

			t = datetime.datetime.now()

			ruta = settings.STATICFILES_DIRS[0]
			newpath = ruta + '/papelera/'

			# newpath = r'static/papelera/'
			filename = "plantillas/multas/formatoMultaOF.xlsx"
			extension = filename[filename.rfind('.'):]

			nombre = 'empresa'+str(empresa_id)+'-'+str(request.user.usuario.id)+str(t.hour)+str(t.minute)+str(t.second)+extension
			functions.descargarArchivoS3(str(filename), str(newpath) , nombre )	
			
			plantilla = settings.STATICFILES_DIRS[0] + '/papelera/'+nombre

			wb = load_workbook(filename = plantilla , read_only=False)
			ws = wb.worksheets[0]

			#Estilo para los bordes del archivo
			medium_border = Border(
			left=Side(style='medium',color='000501'),
			right=Side(style='medium',color='000501'),
			top=Side(style='medium',color='000501'),
			bottom=Side(style='medium',color='000501'))

			medium_border_top = Border(
			left=Side(style='medium',color='000501'),
			right=Side(style='medium',color='000501'),
			bottom=Side(style='medium',color='000501'))

			medium_border_bottom = Border(
			left=Side(style='medium',color='000501'),
			right=Side(style='medium',color='000501'),
			top=Side(style='medium',color='000501'))


			medium_only_border_bottom = Border(
			bottom=Side(style='medium',color='000501'))

			medium_only_border_top = Border(
			top=Side(style='medium',color='000501'))

			medium_only_border_left = Border(
			left=Side(style='medium',color='000501'))

			medium_only_border_left_bottom = Border(
			left=Side(style='medium',color='000501'),
			bottom=Side(style='medium',color='000501'))

			medium_only_border_right = Border(
			right=Side(style='medium',color='000501'))

			medium_only_border_right_bottom = Border(
			right=Side(style='medium',color='000501'),
			bottom=Side(style='medium',color='000501'))

			medium_only_border_right_top = Border(
			right=Side(style='medium',color='000501'),
			top=Side(style='medium',color='000501'))

			thin_border = Border(
			left=Side(style='thin',color='000501'),
			right=Side(style='thin',color='000501'),
			top=Side(style='thin',color='000501'),
			bottom=Side(style='thin',color='000501'))

			thin_only_border_bottom = Border(
			bottom=Side(style='thin',color='000501'))

			thin_only_border_right = Border(
			right=Side(style='thin',color='000501'))

			thin_only_border_left = Border(
			left=Side(style='thin',color='000501'))

			hair_only_border_bottom_top = Border(
			top=Side(style='hair',color='000501'),
			bottom=Side(style='hair',color='000501'))

			hair_only_border_top = Border(
			top=Side(style='hair',color='000501'))

			hair_only_border_right = Border(
			right=Side(style='hair',color='000501'))

			hair_only_border_bottom = Border(
			bottom=Side(style='hair',color='000501'))

			hair_only_border_top_right_bottom = Border(
			right=Side(style='hair',color='000501'),
			bottom=Side(style='hair',color='000501'))
			
			hair_thin_only_border_top_right = Border(
			top=Side(style='hair',color='000501'),
			left=Side(style='medium',color='000501'))

			hair_thin_only_border_left_right = Border(
			right=Side(style='hair',color='000501'),
			left=Side(style='thin',color='000501'))



			#borde bottom general de la plantillan
			for x in xrange(2,18):
				ws.cell(row=2, column=x).border = medium_border


			for x in xrange(2,5):
				ws.cell(row=4, column=x).border = medium_border
				ws.cell(row=5, column=x).border = medium_border
				ws.cell(row=6, column=x).border = medium_border
				ws.cell(row=7, column=x).border = medium_border
				ws.cell(row=12, column=x).border = medium_border_top	

			ws.cell(row=8, column=2).border = medium_border.left
			ws.cell(row=8, column=5).border = medium_border.right
			ws.cell(row=9, column=2).border = medium_border.left
			ws.cell(row=9, column=5).border = medium_border.right
			ws.cell(row=10, column=5).border = medium_border.right
			ws.cell(row=11, column=5).border = medium_border.right

			ws.cell(row=9, column=2).border = hair_thin_only_border_top_right

			for x in xrange(3,5):
				ws.cell(row=11, column=x).border = hair_only_border_bottom_top

			for x in xrange(7,18):
				ws.cell(row=4, column=x).border = medium_border_bottom
				ws.cell(row=12, column=x).border = medium_border_top

			ws.cell(row=5, column=7).border = medium_border.left			
			ws.cell(row=7, column=7).border = medium_border.left

			ws.cell(row=6, column=14).border = hair_only_border_right
			ws.cell(row=8, column=14).border = hair_only_border_top_right_bottom

			ws.cell(row=5, column=18).border = medium_border.right
			ws.cell(row=6, column=18).border = medium_border.right
			ws.cell(row=7, column=18).border = medium_border.right
			ws.cell(row=8, column=18).border = medium_border.right
			ws.cell(row=9, column=18).border = medium_border.right
			ws.cell(row=10, column=18).border = medium_border.right
			ws.cell(row=11, column=18).border = medium_border.right

			for x in xrange(7,15):
				ws.cell(row=5, column=x).border = hair_only_border_bottom_top
				ws.cell(row=7, column=x).border = hair_only_border_bottom_top
				ws.cell(row=11, column=x).border = hair_only_border_bottom_top	

			ws.cell(row=5, column=17).border = hair_only_border_bottom_top
			ws.cell(row=7, column=17).border = hair_only_border_bottom_top

			for x in xrange(7,12):
				ws.cell(row=9, column=x).border = hair_only_border_bottom_top

			ws.cell(row=10, column=x).border = hair_only_border_right

			for x in xrange(13,18):
				ws.cell(row=9, column=x).border = hair_only_border_bottom_top

			for x in xrange(13,18):
				ws.cell(row=11, column=x).border = hair_only_border_bottom_top
				

			for x in xrange(3,17):
				ws.cell(row=14, column=x).border = medium_only_border_top


			for x in xrange(2,18):
				ws.cell(row=22, column=x).border = medium_only_border_bottom


			ws.cell(row=14, column=17).border = medium_only_border_right_top
			ws.cell(row=22, column=2).border = medium_only_border_left_bottom
			ws.cell(row=22, column=17).border = medium_only_border_right_bottom
			ws.cell(row=21, column=17).border = medium_only_border_right
			ws.cell(row=20, column=17).border = medium_only_border_right
			ws.cell(row=19, column=17).border = medium_only_border_right
			ws.cell(row=18, column=17).border = medium_only_border_right
			ws.cell(row=17, column=17).border = medium_only_border_right
			ws.cell(row=16, column=17).border = medium_only_border_right
			ws.cell(row=15, column=17).border = medium_only_border_right


			ws.cell(row=24, column=17).border = medium_border
			ws.cell(row=26, column=17).border = medium_border
			ws.cell(row=28, column=17).border = medium_border

			for x in xrange(4,9):
				ws.cell(row=24, column=x).border = thin_border

			for x in xrange(2,11):
				ws.cell(row=28, column=x).border = thin_only_border_bottom

			for x in xrange(2,7):	
				ws.cell(row=33, column=x).border = thin_border			
				ws.cell(row=39, column=x).border = thin_border

			for x in xrange(9,18):
				ws.cell(row=33, column=x).border = thin_border
				ws.cell(row=38, column=x).border = thin_border
				ws.cell(row=39, column=x).border = thin_border


			ws.cell(row=35, column=2).border = thin_only_border_left
			ws.cell(row=36, column=2).border = thin_only_border_left
			ws.cell(row=38, column=2).border = thin_only_border_left

			ws.cell(row=34, column=6).border = thin_only_border_right
			ws.cell(row=35, column=6).border = thin_only_border_right
			ws.cell(row=36, column=6).border = thin_only_border_right
			ws.cell(row=37, column=6).border = thin_only_border_right
			ws.cell(row=38, column=6).border = thin_only_border_right

			ws.cell(row=35, column=8).border = thin_only_border_left
			ws.cell(row=36, column=8).border = thin_only_border_left
			ws.cell(row=37, column=8).border = thin_only_border_left

			ws.cell(row=34, column=17).border = thin_only_border_right
			ws.cell(row=35, column=17).border = thin_only_border_right
			ws.cell(row=36, column=17).border = thin_only_border_right
			ws.cell(row=37, column=17).border = thin_only_border_right

			for x in xrange(5,7):
				ws.cell(row=41, column=x).border = thin_border

			for x in xrange(9,18):
				ws.cell(row=41, column=x).border = thin_border


			cursor = connection.cursor()	
			
			ws['F5'] = solicitud.contrato.contratista.nombre#NOMBRE CONTRATISTA DE OBRA									
			ws['F9'] = solicitud.contrato.contratista.nit#NIT 	
			ws['M9'] = solicitud.contrato.contratista.nombre#NOMBRE CONTRATISTA DE OBRA								 	

			es=EstadoMulta()

			tipo_contrato = tipoC()
						
			historial = SolicitudHistorial.objects.get(solicitud_id = solicitud.id , estado_id = es.confirmada)

			try:
	
				proyecto = Proyecto.objects.get(contrato = solicitud.contrato.id, contrato__tipo_contrato__id = tipo_contrato.contratoProyecto)
				
				nombre_proyecto = proyecto.nombre
				departamento_proyecto = proyecto.municipio.departamento.nombre
				municipio_proyecto = proyecto.municipio.nombre
				convenio = proyecto.mcontrato.nombre

			except Proyecto.DoesNotExist:
				nombre_proyecto = ''
				departamento_proyecto = ''
				municipio_proyecto = ''
				convenio = ''		


			ws['D16'] = 'Contrato '+str(solicitud.contrato.numero)+' - Convenio '+str(convenio)#Contrato xxxxxxx- Convenio XXXXXXXXXXXXXX 								
			ws['D17'] = nombre_proyecto#NOMBRE DEL PROYECTO								
			ws['D18'] = municipio_proyecto+' - '+departamento_proyecto#MUNICIPIO - DEPARTAMENTO
			ws['D19'] = historial.fecha.strftime("%Y-%m-%d") #FECHA CONFIRMACION DE LA MULTA								

			ws['O15'] = solicitud.valorImpuesto#VALOR DE LA MULTA IMPUESTA

			ws['B37'] = solicitud.codigoOF#CODIGO OF DE LA SOLICITUD

			wb.template = False
			response = HttpResponse(openpyxl.writer.excel.save_virtual_workbook(wb), content_type='application/vnd.ms-excel; charset=utf-8')
			response['Content-Disposition'] ='attachment; filename="solicitud_anticipos.xlsx"'
			return response
		except Exception as e:
			print(e)
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def addTableAfterParagraph(table, paragraph):
	tbl, p = table._tbl, paragraph._p
	p.addnext(tbl)


# exporta a excel multa
def exportReporteMulta(request):

	try:

		response = HttpResponse(content_type='application/vnd.ms-excel;charset=utf-8')
		response['Content-Disposition'] = 'attachment; filename="Reporte_proyecto.xls"'
		
		workbook = xlsxwriter.Workbook(response, {'in_memory': True})
		worksheet = workbook.add_worksheet('Multas')
		

		dato = request.GET['dato'] if 'dato' in request.GET else None;
		estado = request.GET['estado'] if 'estado' in request.GET else None;
		solicitudes_elaboradas = request.GET['solicitudes_elaboradas'] if 'solicitudes_elaboradas' in request.GET else None;
		solicitudes_solicitadas = request.GET['solicitudes_solicitadas'] if 'solicitudes_solicitadas' in request.GET else None;
		solicitudes_consulta = request.GET['solicitudes_consulta'] if 'solicitudes_consulta' in request.GET else None;
		propietario = request.GET['propietario'] if 'propietario' in request.GET else None;
		macro_contrato_id = request.GET['macro_contrato'] if 'macro_contrato' in request.GET else None;
		consecutivo = request.GET['consecutivo'] if 'consecutivo' in request.GET else None;
		fecha_desde = request.GET['fecha_desde'] if 'fecha_desde' in request.GET else None;
		fecha_hasta = request.GET['fecha_hasta'] if 'fecha_hasta' in request.GET else None;
		numero_contrato_obra = request.GET['numero_contrato_obra'] if 'numero_contrato_obra' in request.GET else None;
		contratista_id = request.GET['contratista'] if 'contratista' in request.GET else None;
		solicitante_id = request.GET['solicitante'] if 'solicitante' in request.GET else None;

		empresa_id = request.GET['empresa'] if 'empresa' in request.GET else request.user.usuario.empresa.id;
		solicitud_id = request.GET['solicitud'] if 'solicitud' in request.GET else None;

		# VARIABLES SOLICITUD
		correspondenciasolicita_id = request.GET['correspondenciasolicita'] if 'correspondenciasolicita' in request.GET else None;
		correspondenciadescargo_id = request.GET['correspondenciadescargo'] if 'correspondenciadescargo' in request.GET else None;
		firmaImposicion_id = request.GET['firmaImposicion'] if 'firmaImposicion' in request.GET else None;

		# VARIABLES HISTORIAL
		estado_id = request.GET['estado'] if 'estado' in request.GET else None;

		es=EstadoMulta()

		qset=(~Q(id=0))
		# VARIABLES SOLICITUD EMPRESA
		if dato:
			qset = qset & ( Q(solicitud__contrato__numero__icontains = dato) |
						Q(solicitud__consecutivo__icontains = dato) |
						Q(solicitud__contrato__nombre__icontains = dato))	
		if propietario:

			if solicitudes_consulta is None or solicitudes_consulta == "":
				qset = qset & ( Q(propietario = propietario) )
		
		if empresa_id:
			qset = qset & ( Q(empresa_id = empresa_id) )
		if solicitud_id:
			qset = qset & ( Q(solicitud_id = solicitud_id) )

		# VARIABLES SOLICITUD
		if correspondenciasolicita_id:
			qset = qset & ( Q(solicitud__correspondenciasolicita_id = correspondenciasolicita_id) )		

		if correspondenciadescargo_id:
			qset = qset & ( Q(solicitud__correspondenciadescargo_id = correspondenciadescargo_id) )
		if firmaImposicion_id:
			qset = qset & ( Q(solicitud__firmaImposicion_id = firmaImposicion_id) )

		if macro_contrato_id:
			qset = qset & ( Q(solicitud__contrato__mcontrato_id = macro_contrato_id) )
		if consecutivo:
			qset = qset & ( Q(solicitud__consecutivo__icontains = consecutivo) )
		if numero_contrato_obra:
			qset = qset &( Q(solicitud__contrato__numero__icontains = numero_contrato_obra) )
		if contratista_id:
			qset = qset &( Q(solicitud__contrato__contratista__id = contratista_id) )
		
		if fecha_desde or fecha_hasta:

			if (fecha_desde and (fecha_hasta is None) ):
				qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__gte= fecha_desde) 
								& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )
			
			if ( (fecha_desde is None) and fecha_hasta ):
				qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__lte=fecha_hasta)  
								& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )

			if (fecha_desde  and fecha_hasta ):	
				qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__fecha__range = (fecha_desde , fecha_hasta))
								& Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = es.solicitada) )

		queryset = SolicitudEmpresa.objects.filter(qset).values('solicitud_id').distinct()
		
		queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = queryset).values('solicitud_id').annotate(id=Max('id')).values("id")		

		if queryset:

			qset = ( Q(solicitud__fk_Solicitud_SolicitudHistorial__id__in = queryHistorial) )

			# VARIABLES HISTORIAL
			if estado_id:
				qset = qset & ( Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id = estado_id))	

			if solicitante_id:
				qset = qset & ( Q(empresa_id = solicitante_id) & Q(propietario = True) )
			else:
				qset = qset & ( Q(empresa_id = empresa_id) )	

			if solicitudes_consulta:
				qset = qset & (~Q(solicitud__fk_Solicitud_SolicitudHistorial__estado_id__in = [
						es.elaborada,
					]))
				queryset = SolicitudEmpresa.objects.filter(qset).distinct()
			else:
				queryset = SolicitudEmpresa.objects.filter(qset).distinct()

		# FORMATO PARA ENCABEZADOS	
		format1=workbook.add_format({'border':1,'font_size':12,'bold':True, 'bg_color':'#4a89dc','font_color':'white'})
		# FORMATO PARA CAMPOS DE TIPOS TEXTO
		format2=workbook.add_format({'border':1})
		# FORMATO DE FECHA
		format_date=workbook.add_format({'border':1,'num_format': 'yyyy-mm-dd '})
		format_date2 = workbook.add_format({'num_format': 'dd/mm/yy hh:mm'})
		# FORMATO MONEDA
		format_money=workbook.add_format({'border':1,'num_format': '$#,##0'})
		row=1
		col=0

		model = SolicitudEmpresa
		formato_fecha = "%Y-%m-%d"
		
		worksheet.set_column('A:D', 20)
		worksheet.write('A1', 'Consecutivo', format1)
		worksheet.write('B1', 'Fecha de solicitud', format1)
		worksheet.write('C1', 'Carta solicitud', format1)
		worksheet.write('D1', 'Multas Impuesta', format1)
		worksheet.write('E1', 'No contrato', format1)
		worksheet.write('F1', 'Nombre ', format1)
		worksheet.write('G1', 'Contratista', format1)
		worksheet.write('H1', 'Valor', format1)
		worksheet.write('I1', 'Estado', format1)

		worksheet.set_column('A:A', 11)
		worksheet.set_column('B:B', 20)
		worksheet.set_column('C:C', 16)
		worksheet.set_column('D:D', 15)
		worksheet.set_column('E:E', 18)
		worksheet.set_column('F:F', 30)
		worksheet.set_column('G:G', 18)
		worksheet.set_column('H:H', 18)
		worksheet.set_column('I:I', 22)

		for solicitudEmpresa in queryset:


			total_multa = 0
			idSolicitud = Solicitud.objects.filter(contrato_id = solicitudEmpresa.solicitud.contrato_id).values_list("id")
			queryHistorial = SolicitudHistorial.objects.filter(solicitud_id__in = idSolicitud).values('solicitud_id').annotate(id=Max('id')).values("id")
			if queryHistorial:
				qset = ( Q(pk__in = queryHistorial) )
				qset = qset & ( Q(estado_id__in = [es.contabilizada,es.pendiente_contabilizacion,es.confirmada] ))
				total_multa = SolicitudHistorial.objects.filter(qset).count()	


			estado_solicitud = ''
			datos = SolicitudHistorial.objects.filter(solicitud_id = solicitudEmpresa.solicitud_id).latest('id')
			if datos:
				estado_solicitud = str(datos.estado.nombre)
				worksheet.write(row, col+8, estado_solicitud ,format2)


			datos = SolicitudHistorial.objects.filter(solicitud_id = solicitudEmpresa.solicitud_id , estado_id = es.solicitada).latest('id')
			if datos:			
				dt_string = str(datos.fecha)
				new_dt = dt_string[:10]		
				worksheet.write(row, col+1, new_dt ,format2)

			consecutivo = ''
			if solicitudEmpresa.solicitud.consecutivo:
				consecutivo = solicitudEmpresa.solicitud.consecutivo

			carta_solicitud = ''
			if solicitudEmpresa.solicitud.correspondenciasolicita:
				carta_solicitud = str(solicitudEmpresa.solicitud.correspondenciasolicita.prefijo.nombre)+'-'+str(solicitudEmpresa.solicitud.correspondenciasolicita.consecutivo)

			worksheet.write(row, col, consecutivo ,format2)
			
			worksheet.write(row, col+2, carta_solicitud ,format2)
			worksheet.write(row, col+3, str(total_multa) ,format2)
			worksheet.write(row, col+4,solicitudEmpresa.solicitud.contrato.numero ,format2)
			worksheet.write(row, col+5,solicitudEmpresa.solicitud.contrato.nombre ,format_money)
			worksheet.write(row, col+6,solicitudEmpresa.solicitud.contrato.contratista.nombre ,format2)
			worksheet.write(row, col+7,solicitudEmpresa.solicitud.valorImpuesto ,format_money)
			
			row +=1	

		workbook.close()
		return response

	except Exception as e:
		functions.toLog(e,'multa.exportReporteMulta')
		return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@login_required
def VerSoporte(request):
	if request.method == 'GET':
		try:
			
			archivo = SolicitudHistorial.objects.get(pk=request.GET['id'])
			
			# filename = ""+str(archivo.soporte)+""
			# extension = filename[filename.rfind('.'):]
			# nombre = re.sub('[^A-Za-z0-9 ]+', '', archivo.nombre)
			return functions.exportarArchivoS3(str(archivo.soporte))

		except Exception as e:
			functions.toLog(e,'SolicitudHistorial.VerSoporte')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@login_required
def VerSoporteSolicitud(request):
	if request.method == 'GET':
		try:
			
			archivo = Solicitud.objects.get(pk=request.GET['id'])
			
			# filename = ""+str(archivo.soporte)+""
			# extension = filename[filename.rfind('.'):]
			# nombre = re.sub('[^A-Za-z0-9 ]+', '', archivo.nombre)
			return functions.exportarArchivoS3(str(archivo.soporte))

		except Exception as e:
			functions.toLog(e,'Solicitud.VerSoporteSolicitud')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@login_required
def VerSoporteSolicitudSoporte(request):
	if request.method == 'GET':
		try:
			
			archivo = SolicitudSoporte.objects.get(pk=request.GET['id'])
			return functions.exportarArchivoS3(str(archivo.soporte))

		except Exception as e:
			functions.toLog(e,'Solicitud.VerSoporteSolicitud')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)
		