from django.shortcuts import render#, redirect
from django.contrib.auth.decorators import login_required
from django.template import RequestContext
from django.conf import settings

# Para importaciones a excel y retornar Json a la vista
from django.http import HttpResponse, JsonResponse, StreamingHttpResponse

# Para consultas SQL
from django.db.models import Q#, Sum
from django.db import transaction#, connection
# from django.db.models.deletion import ProtectedError

from wsgiref.util import FileWrapper

from rest_framework import viewsets, serializers
from rest_framework.response import Response
from rest_framework.request import Request
from rest_framework import status
# from rest_framework.pagination import PageNumberPagination

import xlsxwriter
import json
import os
import mimetypes

from datetime import *
# import time
import datetime
from babel.dates import format_date

from coasmedas.functions import functions
from contrato.enumeration import tipoC

from .enumeration import estadoNoConformidad, tipoNoConformidad, valoracionNoConformidad
from .models import NoConformidad
from logs.models import Logs, Acciones
from proyecto.models import Proyecto
from usuario.models import Usuario
from estado.models import Estado
from contrato.models import Contrato
from empresa.models import Empresa 

from estado.views import EstadoLite2Serializer
from usuario.views import UsuarioLiteSerializer

from docx import Document
from docx.shared import Inches #, Pt
from tipo.models import Tipo
from tipo.views import TipoLiteSerializer
from rest_framework.decorators import api_view
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.style import WD_STYLE
# from docx.enum.text import WD_ALIGN_PARAGRAPH

#Api rest para No Conformidad
class ProyectoLiteSerializer(serializers.HyperlinkedModelSerializer):
    
	class Meta: 
		model = Proyecto
		fields=( 'id','nombre')

class NoConformidadSerializer(serializers.HyperlinkedModelSerializer):

	proyecto = ProyectoLiteSerializer(read_only=True)
	proyecto_id = serializers.PrimaryKeyRelatedField(write_only=True, queryset = Proyecto.objects.all())

	usuario = UsuarioLiteSerializer(read_only=True)
	usuario_id = serializers.PrimaryKeyRelatedField(write_only=True, queryset = Usuario.objects.all())

	estado = EstadoLite2Serializer(read_only=True)
	estado_id = serializers.PrimaryKeyRelatedField(write_only=True, queryset = Estado.objects.filter(app='No_conformidad'))

	detectada = UsuarioLiteSerializer(read_only=True)
	detectada_id = serializers.PrimaryKeyRelatedField(write_only=True, queryset = Usuario.objects.all())

	valoracion = TipoLiteSerializer(read_only=True)
	valoracion_id = serializers.PrimaryKeyRelatedField(write_only=True, queryset =  Tipo.objects.filter(app='no_conformidad_valoracion'))

	tipo = TipoLiteSerializer(read_only=True)
	tipo_id = serializers.PrimaryKeyRelatedField(write_only=True, queryset = Tipo.objects.filter(app='no_conformidad'))

	# proyecto = ProyectoLiteSerializer2(read_only = True, many=True)

	class Meta:
		model = NoConformidad
		fields=('id','proyecto','proyecto_id','usuario','usuario_id','estado','estado_id','detectada','detectada_id',
				'descripcion_no_corregida','descripcion_corregida', 'foto_no_corregida','foto_corregida','fecha_no_corregida','fecha_corregida',
				'foto_no_corregida2', 'foto_no_corregida3', 'foto_no_corregida2_publica', 'foto_no_corregida3_publica','foto_corregida2','foto_corregida2_publica', 'foto_corregida3','foto_corregida3_publica',
				'terminada','estructura','primer_correo','segundo_correo','tercer_correo', 'foto_no_corregida_publica', 'foto_corregida_publica', 
				'valoracion_id', 'valoracion', 'tipo', 'tipo_id')

class NoConformidadViewSet(viewsets.ModelViewSet):
	"""

	"""
	model=NoConformidad
	queryset = model.objects.all()
	serializer_class = NoConformidadSerializer
	nombre_modulo = 'No_Conformidad - NoConformidadViewSet'

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)

	def list(self, request, *args, **kwargs):
		try:
			queryset = super(NoConformidadViewSet, self).get_queryset()

			dato = self.request.query_params.get('dato', None)
			id_proyecto = self.request.query_params.get('id_proyecto', None)
			id_usuario = self.request.query_params.get('id_usuario', None)
			id_estado = self.request.query_params.get('id_estado').split(',') if self.request.query_params.get('id_estado') else None
			id_detectada = self.request.query_params.get('id_detectada').split(',') if self.request.query_params.get('id_detectada') else None
			estructura = self.request.query_params.get('estructura', None)
			fecha_desde = self.request.query_params.get('fecha_desde', None)
			fecha_hasta = self.request.query_params.get('fecha_hasta', None)

			sin_paginacion = self.request.query_params.get('sin_paginacion',None)
			id_empresa = request.user.usuario.empresa.id
			valoracion_id = self.request.query_params.get('valoracion_id', None)
			tipo_id = self.request.query_params.get('tipo_id', None)

			qset=(~Q(id=0))

			if dato:
				qset = qset &(
					Q(detectada__persona__nombres__icontains=dato)|
					Q(detectada__persona__apellidos__icontains=dato)|
					Q(proyecto__nombre__icontains=dato)|
					Q(descripcion_no_corregida__icontains=dato)|Q(descripcion_corregida__icontains=dato)|Q(estructura__icontains=dato)
					)

			if id_proyecto:
				qset = qset &(Q(proyecto__id=id_proyecto))

			if id_usuario:
				qset = qset &(Q(usuario__id=id_usuario))

			if id_estado:
				qset = qset &(Q(estado__id__in=id_estado))

			if id_detectada:
				qset = qset &(Q(detectada__id__in=id_detectada))

			if estructura:
				qset = qset &(Q(estructura__icontains=estructura))

			if fecha_desde:
				qset = qset &(Q(fecha_no_corregida__gte=fecha_desde))
			if fecha_hasta:
				qset = qset &(Q(fecha_no_corregida__lte=fecha_hasta))

			if id_empresa:
				qset = qset &(Q(proyecto__fk_proyecto_empresa_proyecto__empresa=id_empresa))

			if valoracion_id:
				qset = qset &(Q(valoracion__id=valoracion_id))

			if tipo_id:
				qset = qset &(Q(tipo__id=tipo_id))
	

			if qset is not None:
				queryset = self.model.objects.filter(qset).order_by('-id')

			if sin_paginacion is None:
				page = self.paginate_queryset(queryset)

				if page is not None:
					serializer = self.get_serializer(page,many=True)
					return self.get_paginated_response({'message':'','success':'ok','data':serializer.data})

			else:
				serializer = self.get_serializer(queryset,many=True)
				return Response({'message':'','success':'ok','data':serializer.data})
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()
			try:
				request.data['usuario_id'] = request.user.usuario.id
				request.data['estado_id'] = estadoNoConformidad.sin_corregir
				serializer = NoConformidadSerializer(data=request.data,context={'request': request})

				# estado_no_conformidad=estadoNoConformidad()
				# print ("texto")
				# 

				if serializer.is_valid():

					serializer.save(proyecto_id=request.data['proyecto_id'], usuario_id=request.user.usuario.id,
						estado_id=request.data['estado_id'], detectada_id=request.data['detectada_id'],
						tipo_id=request.data['tipo_id'],
						valoracion_id=request.data['valoracion_id'],
						foto_no_corregida=self.request.FILES.get('foto_no_corregida'),# if self.request.FILES.get('foto_no_corregida') is not None else '',
						foto_no_corregida2=self.request.FILES.get('foto_no_corregida2', None),
						foto_no_corregida3=self.request.FILES.get('foto_no_corregida3', None),
						foto_corregida=request.FILES['foto_corregida'] if request.FILES.get('foto_corregida') is not None else '',
						foto_corregida2=request.FILES['foto_corregida2'] if request.FILES.get('foto_corregida2') is not None else '',
						foto_corregida3=request.FILES['foto_corregida3'] if request.FILES.get('foto_corregida3') is not None else '')

					logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_crear,nombre_modelo='No_conformidad.NoConformidad',id_manipulado=serializer.data['id'])
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok','data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					# 
					return Response(serializer.errors,status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				transaction.savepoint_rollback(sid)
				functions.toLog(e,self.nombre_modulo)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error','data':''},status=status.HTTP_400_BAD_REQUEST)

	@transaction.atomic
	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			sid = transaction.savepoint()
			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				# 
				serializer = NoConformidadSerializer(instance,data=request.data,context={'request': request},partial=partial)

				if serializer.is_valid():

					serializer.save(proyecto_id=request.data['proyecto_id'], usuario_id=request.user.usuario.id,
						detectada_id=request.data['detectada_id'],
						tipo_id=request.data['tipo_id'],
						valoracion_id=request.data['valoracion_id'],
						estado_id=estadoNoConformidad.corregida if request.FILES.get('foto_corregida') is not None else instance.estado_id,
						foto_no_corregida=request.FILES['foto_no_corregida'] if request.FILES.get('foto_no_corregida') is not None else instance.foto_no_corregida,
						foto_no_corregida2=self.request.FILES.get('foto_no_corregida2', ''),
						foto_no_corregida3=self.request.FILES.get('foto_no_corregida3', ''),
						foto_corregida=request.FILES['foto_corregida'] if request.FILES.get('foto_corregida') is not None else instance.foto_corregida,
						foto_corregida2=request.FILES['foto_corregida2'] if request.FILES.get('foto_corregida2') is not None else '',
						foto_corregida3=request.FILES['foto_corregida3'] if request.FILES.get('foto_corregida3') is not None else '')

					logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_actualizar,nombre_modelo='No_conformidad.NoConformidad',id_manipulado=instance.id)
					logs_model.save()

					transaction.savepoint_commit(sid)
					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok','data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					# 
					return Response(serializer.errors,status=status.HTTP_400_BAD_REQUEST)
			except Exception as e:
				functions.toLog(e,self.nombre_modulo)
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error','data':''},status=status.HTTP_400_BAD_REQUEST)

	@transaction.atomic
	def destroy(self,request,*args,**kwargs):
		sid = transaction.savepoint()
		try:
			instance = self.get_object()
			# self.perform_destroy(instance)
			self.model.objects.get(pk=instance.id).delete()

			logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_borrar,nombre_modelo='No_conformidad.NoConformidad',id_manipulado=instance.id)
			logs_model.save()

			transaction.savepoint_commit(sid)
			return Response({'message':'El registro se ha eliminado correctamente','success':'ok','data':''},status=status.HTTP_201_CREATED)
		except Exception as e:
			functions.toLog(e,self.nombre_modulo)
			transaction.savepoint_rollback(sid)
			return Response({'message':'Se presentaron errores al procesar la solicitud','success':'error','data':''},status=status.HTTP_400_BAD_REQUEST)	
#Fin api rest para No Conformidad

@login_required
def no_conformidad(request):
	tipo_c=tipoC()

	# querysetTipos=Tipo.objects.filter(app='contrato')
	querysetEstado=Estado.objects.filter(app='No_conformidad')
	# querysetRubros=Rubro.objects.all()
	# id_usuario = request.user.usuario.id
	# queryset_empresa=Empresa.objects.filter(esContratante=True)
	valoraciones = Tipo.objects.filter(app='no_conformidad_valoracion')
	tipos = Tipo.objects.filter(app='no_conformidad')
	querysetMContrato=Contrato.objects.filter(tipo_contrato=tipo_c.m_contrato,empresacontrato__empresa=request.user.usuario.empresa.id,empresacontrato__participa=1,activo=1)
	return render(request, 'no_conformidad/no_conformidad.html',{'mcontrato':querysetMContrato,'estado':querysetEstado,'model':'noconformidad','app':'no_conformidad', 
																'tipos': tipos, 'valoraciones': valoraciones})


# Eliminar No Conformidad con una lista
@transaction.atomic
def destroyNoConformidad(request):
	if request.method == 'POST':
		sid = transaction.savepoint()
		try:
			lista=request.POST['_content']
			respuesta= json.loads(lista)
			myList = respuesta['lista']

			for item in myList:
				# 
				model_no_conformidad = NoConformidad.objects.get(pk=item['id'])
				model_no_conformidad.delete()

				logs_model=Logs(usuario_id=request.user.usuario.id,accion=Acciones.accion_borrar,nombre_modelo='No_conformidad.NoConformidad',id_manipulado=item['id'])
				logs_model.save()

			# transaction.commit()
			return JsonResponse({'message':'El registro se ha eliminado correctamente','success':'ok','data': ''})
			transaction.savepoint_commit(sid)
		except Exception as e:
			transaction.savepoint_rollback(sid)
			functions.toLog(e, 'NoConformidad - destroyNoConformidad')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# exporta a excel No Conformidad
def exportReporteNoConformidad(request):
	
	response = HttpResponse(content_type='application/vnd.ms-excel;charset=utf-8')
	response['Content-Disposition'] = 'attachment; filename="Reporte_No_Conformidad.xls"'
	
	workbook = xlsxwriter.Workbook(response, {'in_memory': True})
	worksheet = workbook.add_worksheet('No Conformidad')

	format1=workbook.add_format({'border':0,'font_size':12,'bold':True})
	format1.set_align('center')
	format1.set_align('vcenter')
	format2=workbook.add_format({'border':0})
	# format3=workbook.add_format({'border':0,'font_size':12})
	format5=workbook.add_format()
	format5.set_num_format('yyyy-mm-dd')

	worksheet.set_column('A:A', 10)
	worksheet.set_column('B:B', 20)
	worksheet.set_column('C:D', 15)
	worksheet.set_column('E:E', 35)
	worksheet.set_column('F:F', 13)
	worksheet.set_column('G:I', 30)
	worksheet.set_column('J:J', 20)

	dato = None
	id_proyecto = None
	id_estado = None
	estructura = None
	fecha_desde = None
	fecha_hasta = None
	id_empresa = request.user.usuario.empresa.id

	if request.GET['dato']:
		dato = request.GET['dato']

	if request.GET['id_proyecto']:
		id_proyecto = request.GET['id_proyecto']

	if request.GET['id_estado']:
		id_estado = request.GET['id_estado']

	if request.GET['estructura']:
		estructura = request.GET['estructura']

	if request.GET['fecha_desde']:
		fecha_desde = request.GET['fecha_desde']

	if request.GET['fecha_hasta']:
		fecha_hasta = request.GET['fecha_hasta']

	qset=(~Q(id=0))

	if dato:
		qset = qset &(
			Q(detectada__persona__nombres__icontains=dato)|
			Q(detectada__persona__apellidos__icontains=dato)|
			Q(proyecto__nombre__icontains=dato)|
			Q(descripcion_no_corregida__icontains=dato)|Q(descripcion_corregida__icontains=dato)|Q(estructura__icontains=dato)
			)

	if id_proyecto:
		qset = qset &(Q(proyecto__id=id_proyecto))

	if id_estado:
		qset = qset &(Q(estado__id__in=id_estado))

	if estructura:
		qset = qset &(Q(estructura__icontains=estructura))

	if fecha_desde:
		qset = qset &(Q(fecha_no_corregida__gte=fecha_desde))
	if fecha_hasta:
		qset = qset &(Q(fecha_no_corregida__lte=fecha_hasta))

	if id_empresa:
		qset = qset &(Q(proyecto__fk_proyecto_empresa_proyecto__empresa=id_empresa))

	# print ("texto")

	model=NoConformidad

	if qset is not None:
		queryset = model.objects.filter(qset).order_by('-id')

		worksheet.write('A1', 'Estado', format1)
		worksheet.write('B1', 'Fecha levantamiento', format1)
		worksheet.write('C1', 'Departamento', format1)
		worksheet.write('D1', 'Municipio', format1)
		worksheet.write('E1', 'Proyecto', format1)
		worksheet.write('F1', 'Fecha cierre', format1)
		worksheet.write('G1', 'Detectada', format1)
		worksheet.write('H1', 'Descripcion no corregida', format1)
		worksheet.write('I1', 'Descripcion corregida', format1)
		worksheet.write('J1', 'Estructura', format1)

	row=1
	col=0
	
	for noConformidad in queryset:
		# meses = ''
		# if detalle.cuenta is not None:
		# 	cuenta_nombre= "Girar de la cuenta de %s No. %s ' - ' %s" % (detalle.cuenta.nombre , detalle.cuenta.numero, detalle.cuenta.fiduciaria) if detalle.cuenta.nombre is not None else ''

		# if detalle.cuenta is not None:
		# 	cuenta_referencia= detalle.encabezado.referencia  if detalle.encabezado.referencia is not None else ''

		worksheet.write(row, col,str(noConformidad.estado.nombre),format2)
		worksheet.write(row, col+1,noConformidad.fecha_no_corregida,format5)
		worksheet.write(row, col+2,noConformidad.proyecto.municipio.departamento.nombre,format2)
		worksheet.write(row, col+3,noConformidad.proyecto.municipio.nombre,format2)
		worksheet.write(row, col+4,noConformidad.proyecto.nombre,format2)
		worksheet.write(row, col+5,noConformidad.fecha_corregida,format5)
		worksheet.write(row, col+6,noConformidad.detectada.persona.nombres+' '+noConformidad.detectada.persona.apellidos,format2)
		worksheet.write(row, col+7,noConformidad.descripcion_no_corregida,format2)
		worksheet.write(row, col+8,noConformidad.descripcion_corregida,format2)
		worksheet.write(row, col+9,noConformidad.estructura,format2)

		#worksheet.write(row, col+7,detalle.cuenta.nombre if detalle.cuenta is not None else '',format2)
		row +=1
	workbook.close()
	return response

# exporta a word No Conformidad
def exportReporteWordNoConformidad(request):
	import os.path as path
	try:				
		hoy = date.today()		
		# hoy = datetime.date.today().strftime("%d") + ' de ' +datetime.date.today().strftime("%B") +' del '+datetime.date.today().strftime("%Y")
		# formato_fecha = "%d-%B-%Y"
		# hoy = datetime.strptime(str(hoy), formato_fecha)
		fecha_envio = format_date(hoy, format='long' , locale='es')

		document = Document()

		dato = None
		id_proyecto = None
		id_estado = None
		estructura = None
		fecha_desde = None
		fecha_hasta = None
		id_empresa = request.user.usuario.empresa.id

		if request.GET['dato']:
			dato = request.GET['dato']
		if request.GET['id_proyecto']:
			id_proyecto = request.GET['id_proyecto']
		if request.GET['id_estado']:
			id_estado = request.GET['id_estado'].split(',')
		if request.GET['estructura']:
			estructura = request.GET['estructura']
		if request.GET['fecha_desde']:
			fecha_desde = request.GET['fecha_desde']
		if request.GET['fecha_hasta']:
			fecha_hasta = request.GET['fecha_hasta']

		qset=(~Q(id=0))

		if dato:
			qset = qset &(
				Q(detectada__persona__nombres__icontains=dato)|
				Q(detectada__persona__apellidos__icontains=dato)|
				Q(proyecto__nombre__icontains=dato)|
				Q(descripcion_no_corregida__icontains=dato)|Q(descripcion_corregida__icontains=dato)|Q(estructura__icontains=dato)
				)
		if id_proyecto:
			qset = qset &(Q(proyecto__id=id_proyecto))
		if id_estado:
			qset = qset &(Q(estado__id__in=id_estado))
		if estructura:
			qset = qset &(Q(estructura__icontains=estructura))
		if fecha_desde:
			qset = qset &(Q(fecha_no_corregida__gte=fecha_desde))
		if fecha_hasta:
			qset = qset &(Q(fecha_no_corregida__lte=fecha_hasta))
		if id_empresa:
			qset = qset &(Q(proyecto__fk_proyecto_empresa_proyecto__empresa=id_empresa))

		# print ("texto")

		if qset is not None:
			queryset = NoConformidad.objects.filter(qset).order_by('-id')

		for noConformidad in queryset:
			
			document.add_heading(str(fecha_envio), 9)

			p = document.add_paragraph()
			p.add_run('Proyecto: ').bold = True
			p.add_run(noConformidad.proyecto.nombre)
			# p.add_run('italic.').italic = True

			# document.add_heading('Heading, level 1', level=1)
			# document.add_paragraph('Intense quote', style='IntenseQuote')
			# document.add_paragraph(settings.MEDIA_URL, style='ListBullet')

			document.add_paragraph(noConformidad.descripcion_no_corregida, style='ListNumber')
			p = document.add_paragraph()
			p.add_run('Estructura: ').bold = True
			p.add_run(noConformidad.estructura)

			# newpath = r'static/papelera/'
			ruta = settings.STATICFILES_DIRS[0]
			newpath = ruta + '/papelera/'
			filename = str(noConformidad.foto_no_corregida)
			extension = filename[filename.rfind('.'):]
			nombre = str(noConformidad.id)+extension
			
			functions.descargarArchivoS3(str(filename), str(newpath) , nombre )

			if noConformidad.foto_corregida:

				filename2 = str(noConformidad.foto_corregida)
				extension2 = filename2[filename2.rfind('.'):]
				nombre2 = str(noConformidad.id)+'_2'+str(extension2)

				functions.descargarArchivoS3(str(filename2), str(newpath) , nombre2 )
			# parrafo = document.add_paragraph('')
			# run = parrafo.add_run()

			# document.add_picture(str(newpath)+"/"+str(nombre), width=Inches(2.25), height = Inches(2.50))
			# document.add_picture(str(settings.MEDIA_URL)+str(noConformidad.foto_no_corregida), width=Inches(1.25))

			table = document.add_table(rows=1, cols=2)
			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = 'NC Detectada'
			hdr_cells[1].text = 'NC Corregida'
			# for item in recordset:
			row_cells = table.add_row().cells

			parrafo = row_cells[0].add_paragraph()
			run = parrafo.add_run()
			# row_cells[1].text = 'hola'
			# row_cells[0].add_picture('./static/images/default_avatar_male.jpg', width=Inches(1.25))
			
			if path.exists(str(newpath)+"/"+str(nombre)):
				run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(2.25), height = Inches(2.50))
			else:
				row_cells[0].text = 'Sin foto'						

			if path.exists(str(newpath)+"/"+str(nombre)):
				run.add_picture(str(newpath)+"/"+str(nombre), width=Inches(2.25), height = Inches(2.50))
			else:
				row_cells[0].text = 'Sin foto'

			if noConformidad.foto_corregida:
				parrafo = row_cells[1].add_paragraph()
				run = parrafo.add_run()
				if path.exists(str(newpath)+"/"+str(nombre2)):
					run.add_picture(str(newpath)+"/"+str(nombre2), width=Inches(2.25), height = Inches(2.50))								
			else:
				row_cells[1].text = 'Sin foto'

			table.style = 'LightShading-Accent1'
			# table.style = 'TableGrid'

			# Devuelve un parrafo recien agregado al final del documento y que contiene solo un salto de pagina
			document.add_page_break()

		document.save('demo.docx')

		chunk_size = 108192
		nombreDescarga = 'documento.docx'

		response = StreamingHttpResponse(FileWrapper(open('demo.docx','rb'),chunk_size),content_type=mimetypes.guess_type('demo.docx')[0])
		response['Content-Length'] = os.path.getsize('demo.docx')
		response['Content-Disposition'] = "attachment; filename=%s" % nombreDescarga
		
		return response
	except Exception as e:		
		functions.toLog(e, 'NoConformidad - exportReporteWordNoConformidad')		
		return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@login_required
def VerSoporte(request):
	if request.method == 'GET':
		try:
			
			tipo = request.GET['tipo']
			archivo = NoConformidad.objects.get(pk=request.GET['id'])
			
			if tipo == 'foto_corregida':
				return functions.exportarArchivoS3(str(archivo.foto_corregida))
			elif tipo == 'foto_corregida2':
				return functions.exportarArchivoS3(str(archivo.foto_corregida2))
			elif tipo == 'foto_corregida3':
				return functions.exportarArchivoS3(str(archivo.foto_corregida3))	
			elif tipo == 'foto_no_corregida':
				return functions.exportarArchivoS3(str(archivo.foto_no_corregida))
			elif tipo == 'foto_no_corregida2':
				return functions.exportarArchivoS3(str(archivo.foto_no_corregida2))
			elif tipo == 'foto_no_corregida3':
				return functions.exportarArchivoS3(str(archivo.foto_no_corregida3))
			
		except Exception as e:
			functions.toLog(e,'NoConformidad.VerSoporte')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@login_required
def graficas(request):
	return render(request, 'no_conformidad/dashboard.html',{'model':'no_conformidad','app':'no_conformidad'})



@api_view(['GET'])
def get_dataGraph(request):
	if request.method == 'GET':
		try:
			datos = []
			#cantidad de expedientes por estado
			
			#import pdb; pdb.set_trace()
			
			datagrafica = []			
			total_corregido = NoConformidad.objects.filter(estado__id=estadoNoConformidad.corregida, proyecto__fk_proyecto_empresa_proyecto__empresa=request.user.usuario.empresa.id).count()
			total_sincorregir = NoConformidad.objects.filter(estado__id=estadoNoConformidad.sin_corregir, proyecto__fk_proyecto_empresa_proyecto__empresa=request.user.usuario.empresa.id).count()

			total = total_corregido + total_sincorregir

			
			datagrafica.append(['Corregido',
				round((float(total_corregido) / float(total))*100,2)])

			datagrafica.append(['Sin corregir',
				round((float(total_sincorregir) / float(total))*100,2)])

			datos.append(
				{
					'grafica' : 'Graficas no conformidades',
					'datagrafica' : datagrafica
				}
			)

			#cantidad de predios por docuentacion			
			
			datagrafica2 = []
			total_tecnica = NoConformidad.objects.filter(tipo__id=tipoNoConformidad.tecnica, proyecto__fk_proyecto_empresa_proyecto__empresa=request.user.usuario.empresa.id).count()
			total_administrativa = NoConformidad.objects.filter(tipo__id=tipoNoConformidad.admin, proyecto__fk_proyecto_empresa_proyecto__empresa=request.user.usuario.empresa.id).count()
			total_otros = NoConformidad.objects.filter(tipo=None, proyecto__fk_proyecto_empresa_proyecto__empresa=request.user.usuario.empresa.id).count()

			total2 = total_tecnica + total_administrativa + total_otros

			datagrafica2.append({	
					'id':0,				
					'nombre': 'Tecnica',
					'porcentaje': round((float(total_tecnica) / float(total2))*100,2),
					'color':'primary',
					'cant':'Cant: '+str(total_tecnica),
				})

			datagrafica2.append({
					'id':1,				
					'nombre': 'Administrativa',
					'porcentaje': round((float(total_administrativa) / float(total2))*100,2),
					'color':'info',
					'cant':'Cant: '+str(total_administrativa),
				})

			datagrafica2.append({
					'id':2,			
					'nombre': 'Otros',
					'porcentaje': round((float(total_otros) / float(total2))*100,2),
					'color':'warning',
					'cant':'Cant: '+str(total_otros),
				})


			datos.append(
				{
					'grafica' : 'Graficas tipos',
					'datagrafica' : datagrafica2
				}
			)



			datagrafica3 = []
			total_mayor = NoConformidad.objects.filter(valoracion__id=valoracionNoConformidad.mayor, proyecto__fk_proyecto_empresa_proyecto__empresa=request.user.usuario.empresa.id).count()
			total_menor = NoConformidad.objects.filter(valoracion__id=valoracionNoConformidad.menor, proyecto__fk_proyecto_empresa_proyecto__empresa=request.user.usuario.empresa.id).count()
			total_otros = NoConformidad.objects.filter(valoracion=None, proyecto__fk_proyecto_empresa_proyecto__empresa=request.user.usuario.empresa.id).count()

			total3 = total_mayor + total_menor + total_otros

			datagrafica3.append({	
					'id':3,					
					'nombre': 'Mayor',
					'porcentaje': round((float(total_mayor) / float(total3))*100,2),
					'color':'primary',
					'cant':'Cant: '+str(total_mayor),
				})

			datagrafica3.append({
					'id':4,					
					'nombre': 'Menor',
					'porcentaje': round((float(total_menor) / float(total3))*100,2),
					'color':'info',
					'cant':'Cant: '+str(total_menor),
				})

			datagrafica3.append({
					'id':5,					
					'nombre': 'Otros',
					'porcentaje': round((float(total_otros) / float(total3))*100,2),
					'color':'warning',
					'cant':'Cant: '+str(total_otros),
				})

			datos.append(
				{
					'grafica' : 'Graficas valoracion',
					'datagrafica' : datagrafica3
				}
			)

			return Response({'message':'','success':'ok','data':datos})	
					
		except Exception as e:
			functions.toLog(e,'servidumbre.datosGraficas')
			return Response({'message':'Se presentaron errores al procesar la solicitud','success':'error','data':''})	